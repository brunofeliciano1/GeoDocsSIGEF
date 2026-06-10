import os
import re
import unicodedata


MESES_PT = {
    1: "janeiro",
    2: "fevereiro",
    3: "março",
    4: "abril",
    5: "maio",
    6: "junho",
    7: "julho",
    8: "agosto",
    9: "setembro",
    10: "outubro",
    11: "novembro",
    12: "dezembro",
}


def formatar_data_extenso(valor):
    """Converte qualquer representação de data para "dd de mês de aaaa".

    Suporta: QDate nativo, QDate como string, datetime.date, dd/mm/aaaa, dd/mm/aa.
    Retorna "" para valores nulos ou não reconhecidos.
    """
    if valor is None:
        return ""

    # QDate nativo — .month() é callable
    if hasattr(valor, "month") and callable(getattr(valor, "month", None)):
        try:
            dia = int(valor.day())
            mes = int(valor.month())
            ano = int(valor.year())
            nome = MESES_PT.get(mes, "")
            return ("%02d de %s de %d" % (dia, nome, ano)) if nome else ""
        except Exception:
            pass

    # datetime.date / datetime.datetime — .month não é callable
    if hasattr(valor, "month") and not callable(getattr(valor, "month", None)):
        try:
            dia, mes, ano = int(valor.day), int(valor.month), int(valor.year)
            nome = MESES_PT.get(mes, "")
            return ("%02d de %s de %d" % (dia, nome, ano)) if nome else ""
        except Exception:
            pass

    texto = str(valor).strip()
    if not texto or texto.upper() in ("NULL", "NONE", ""):
        return ""

    # "PyQt5.QtCore.QDate(1985, 2, 24)"
    m = re.search(r"QDate\((\d{4}),\s*(\d{1,2}),\s*(\d{1,2})\)", texto)
    if m:
        ano, mes, dia = int(m.group(1)), int(m.group(2)), int(m.group(3))
        nome = MESES_PT.get(mes, "")
        return ("%02d de %s de %d" % (dia, nome, ano)) if nome else ""

    # "dd/mm/aaaa" ou "dd/mm/aa" (com ou sem horário)
    m = re.search(r"(\d{1,2})/(\d{1,2})/(\d{2,4})", texto)
    if m:
        dia, mes, ano = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if ano < 100:
            ano += 2000
        nome = MESES_PT.get(mes, "")
        return ("%02d de %s de %d" % (dia, nome, ano)) if nome else ""

    return texto


# Partículas de município que ficam em minúsculo no nome próprio
_PARTICULAS_MUNICIPIO = frozenset({"de", "da", "do", "das", "dos", "e"})


def formatar_municipio_para_data(municipio):
    """Converte nome de município para nome próprio (title case, partículas em minúsculo).

    Exemplos:
        'SANTANA DO ACARAÚ'      → 'Santana do Acaraú'
        'SÃO GONÇALO DO AMARANTE' → 'São Gonçalo do Amarante'
        'JUAZEIRO DO NORTE'      → 'Juazeiro do Norte'
    """
    if not municipio:
        return municipio or ""
    palavras = municipio.strip().split()
    resultado = []
    for i, p in enumerate(palavras):
        lower = p.lower()
        if i > 0 and lower in _PARTICULAS_MUNICIPIO:
            resultado.append(lower)
        else:
            resultado.append(p[0].upper() + p[1:].lower() if p else "")
    return " ".join(resultado)


def formatar_trt_sem_duplicar_uf(trt, uf=""):
    """Remove UF duplicada do final do código TRT/ART.

    O modelo usa '{{TRT}} - {{UF}}'. Se o valor de TRT já traz uma sigla
    de estado no final, esta função retira esse sufixo para que o modelo
    acrescente a UF apenas uma vez.
    """
    if not trt:
        return trt or ""
    trt = trt.strip()
    # Remove " - XX" ou "-XX" no final (qualquer sigla de 2 letras maiúsculas)
    trt = re.sub(r"\s*[-–]\s*[A-Z]{2}\s*$", "", trt, flags=re.IGNORECASE)
    return trt.strip()


PLACEHOLDER_PATTERN = re.compile(r"(\{\{\s*([^{}]+?)\s*\}\}|<<\s*([^<>]+?)\s*>>|\$\{\s*([^{}]+?)\s*\})")


def normalize_key(value):
    value = unicodedata.normalize("NFKD", str(value).strip().lower())
    value = "".join(char for char in value if not unicodedata.combining(char))
    value = re.sub(r"[^a-z0-9]+", "_", value)
    return value.strip("_")


def public_data(data):
    return {key: value for key, value in data.items() if not key.startswith("__")}


def ensure_docx_path(path):
    base, extension = os.path.splitext(path)
    if extension.lower() == ".docx":
        return path
    if extension:
        return base + ".docx"
    return path + ".docx"


def all_document_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for paragraph in table_paragraphs(table):
            yield paragraph
    for section in doc.sections:
        for container in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                for paragraph in table_paragraphs(table):
                    yield paragraph


def table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                for paragraph in table_paragraphs(nested_table):
                    yield paragraph


def clear_paragraph_content(paragraph):
    p_pr = paragraph._p.pPr
    paragraph._p.clear_content()
    if p_pr is not None and paragraph._p.pPr is None:
        paragraph._p.insert(0, p_pr)


def add_labeled_value(paragraph, label, value, label_bold, value_bold, signature_font=False):
    label_run = paragraph.add_run(label + " ")
    if signature_font:
        set_signature_run_font(label_run, label_bold)
    else:
        label_run.bold = label_bold
    value_run = paragraph.add_run(value)
    if signature_font:
        set_signature_run_font(value_run, value_bold)
    else:
        value_run.bold = value_bold


def set_signature_run_font(run, bold=None):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Arial"
    run.font.size = Pt(12)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is not None:
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_fonts.set(qn("w:cs"), "Arial")


def signature_line_for_owner(owner):
    reference = "Proprietário(a): %s" % owner.get("nome", "")
    length = max(45, min(len(reference), 70))
    return "_" * length


def add_blank_signature_lines(paragraph, count):
    for _index in range(count):
        paragraph.add_run().add_break()


def paragraph_is_single_placeholder(text, expected_key):
    match = PLACEHOLDER_PATTERN.fullmatch(text.strip())
    if not match:
        return False
    raw_key = match.group(2) or match.group(3) or match.group(4)
    return normalize_key(raw_key) == expected_key


def value_for_placeholder(match, data):
    raw_key = match.group(2) or match.group(3) or match.group(4)
    key = normalize_key(raw_key)
    return data.get(key, "")


def replace_match_across_runs(runs, run_ranges, start, end, replacement):
    affected = [
        (index, run_start, run_end)
        for index, (_run, run_start, run_end) in enumerate(run_ranges)
        if run_start < end and run_end > start
    ]
    if not affected:
        return

    first_index, first_start, _first_end = affected[0]
    last_index, last_start, _last_end = affected[-1]

    first_run = runs[first_index]
    if first_index == last_index:
        local_start = start - first_start
        local_end = end - first_start
        first_run.text = first_run.text[:local_start] + replacement + first_run.text[local_end:]
        return

    first_run.text = first_run.text[:start - first_start] + replacement
    for middle_index, _middle_start, _middle_end in affected[1:-1]:
        runs[middle_index].text = ""
    last_run = runs[last_index]
    last_run.text = last_run.text[end - last_start:]


def replace_in_paragraphs(paragraphs, data):
    count = 0
    for paragraph in paragraphs:
        count += replace_placeholders_in_paragraph(paragraph, data)
    return count


def replace_in_tables(tables, data):
    count = 0
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                count += replace_in_paragraphs(cell.paragraphs, data)
                count += replace_in_tables(cell.tables, data)
    return count


def replace_placeholders_in_paragraph(paragraph, data):
    runs = paragraph.runs
    if not runs:
        return 0

    full_text = "".join(run.text for run in runs)
    matches = list(PLACEHOLDER_PATTERN.finditer(full_text))
    if not matches:
        return 0

    if paragraph_is_single_placeholder(full_text, "bloco_proprietarios"):
        write_owner_block_paragraph(paragraph, data.get("__owners", []))
        return 1

    if paragraph_is_single_placeholder(full_text, "bloco_assinaturas_proprietarios"):
        write_owner_signature_block_paragraph(paragraph, data.get("__owners", []))
        return 1

    run_ranges = []
    position = 0
    for run in runs:
        start = position
        end = start + len(run.text)
        run_ranges.append((run, start, end))
        position = end

    for match in reversed(matches):
        replacement = value_for_placeholder(match, data)
        replace_match_across_runs(runs, run_ranges, match.start(), match.end(), replacement)
    return len(matches)


def write_owner_block_paragraph(paragraph, owners):
    clear_paragraph_content(paragraph)
    paragraph.paragraph_format.keep_together = True

    for index, owner in enumerate(owners):
        if index:
            paragraph.add_run().add_break()

        add_labeled_value(
            paragraph,
            "Proprietário(a):",
            owner["nome"],
            label_bold=True,
            value_bold=False,
            signature_font=True,
        )
        paragraph.add_run().add_break()
        if owner.get("cpf"):
            add_labeled_value(paragraph, "CPF:", owner["cpf"], label_bold=True, value_bold=False, signature_font=True)
        elif owner.get("cnpj"):
            add_labeled_value(paragraph, "CNPJ:", owner["cnpj"], label_bold=True, value_bold=False, signature_font=True)


def write_owner_signature_block_paragraph(paragraph, owners):
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    clear_paragraph_content(paragraph)
    paragraph.paragraph_format.keep_together = True

    for index, owner in enumerate(owners):
        if index:
            add_blank_signature_lines(paragraph, 6)

        line_run = paragraph.add_run(signature_line_for_owner(owner))
        set_signature_run_font(line_run)
        paragraph.add_run().add_break()
        add_labeled_value(
            paragraph,
            "Proprietário(a):",
            owner["nome"],
            label_bold=False,
            value_bold=True,
            signature_font=True,
        )
        paragraph.add_run().add_break()
        if owner.get("cpf"):
            add_labeled_value(paragraph, "CPF:", owner["cpf"], label_bold=False, value_bold=True, signature_font=True)
        elif owner.get("cnpj"):
            add_labeled_value(paragraph, "CNPJ:", owner["cnpj"], label_bold=False, value_bold=True, signature_font=True)


def paragraph_has_image(paragraph):
    xml = paragraph._p.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def is_responsible_signature_paragraph(text_key):
    tokens = (
        "resp_tec",
        "bruno_feliciano_de_lima_alves",
        "tecnico_agricola",
        "agrimensura_cfta",
        "codigo_de_credenciamento",
    )
    return any(token in text_key for token in tokens)


def is_short_signature_line(text_key):
    if len(text_key) > 80:
        return False
    tokens = (
        "tecnico_agricola",
        "agrimensura",
        "cfta",
        "codigo_de_credenciamento",
        "incra",
        "trt",
        "uf",
    )
    return any(token in text_key for token in tokens)


def is_signature_separator(text):
    compact = re.sub(r"\s+", "", str(text))
    return bool(compact) and set(compact) <= {"_", "-", "—", "–"}


def find_signature_block_start(paragraphs, index):
    start = index
    while start > 0 and index - start < 5:
        previous = paragraphs[start - 1]
        previous_key = normalize_key(previous.text)
        if paragraph_has_image(previous) or is_signature_separator(previous.text) or not previous_key:
            start -= 1
            continue
        break
    return start


def find_signature_block_end(paragraphs, index):
    end = index
    while end + 1 < len(paragraphs) and end - index < 8:
        next_paragraph = paragraphs[end + 1]
        next_key = normalize_key(next_paragraph.text)
        if not next_key:
            break
        if is_responsible_signature_paragraph(next_key) or is_short_signature_line(next_key):
            end += 1
            continue
        break
    return end


def find_signature_block_end_after(paragraphs, start_index):
    end = min(start_index + 2, len(paragraphs) - 1)
    for index in range(start_index + 1, len(paragraphs)):
        if is_responsible_signature_paragraph(normalize_key(paragraphs[index].text)):
            return find_signature_block_end(paragraphs, index)
        if index - start_index > 14:
            break
        end = index
    return end


def mark_keep_block(paragraphs, start_index, end_index):
    for index in range(start_index, end_index + 1):
        paragraph = paragraphs[index]
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = index < end_index


def marcar_blocos_assinatura_tecnica(paragraphs):
    for index, paragraph in enumerate(paragraphs):
        if not is_responsible_signature_paragraph(normalize_key(paragraph.text)):
            continue
        start_index = find_signature_block_start(paragraphs, index)
        end_index = find_signature_block_end(paragraphs, index)
        mark_keep_block(paragraphs, start_index, end_index)


def substituir_marcadores_docx(doc, dados):
    count = 0
    count += replace_in_paragraphs(doc.paragraphs, dados)
    count += replace_in_tables(doc.tables, dados)

    for section in doc.sections:
        for container in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            count += replace_in_paragraphs(container.paragraphs, dados)
            count += replace_in_tables(container.tables, dados)
    return count


# ── Formatação de campos cadastrais ──────────────────────────────────────────

def formatar_ccir(valor):
    """Formata CCIR para 999.920.756.130-4 (13 dígitos numéricos).

    Exemplos:
        9999207561304        → 999.920.756.130-4
        999.920.756.130-4    → 999.920.756.130-4  (sem alteração)
    """
    if not valor:
        return str(valor or "").strip()
    digitos = re.sub(r"\D+", "", str(valor))
    if len(digitos) == 13:
        return "%s.%s.%s.%s-%s" % (
            digitos[0:3], digitos[3:6], digitos[6:9], digitos[9:12], digitos[12]
        )
    return str(valor).strip()


def formatar_cib(valor):
    """Formata CIB para X.XXX.XXX-X (8 caracteres alfanuméricos).

    Exemplos:
        9C173B29    → 9.C17.3B2-9
        98948997    → 9.894.899-7
        9.C17.3B2-9 → 9.C17.3B2-9  (sem alteração)
    """
    if not valor:
        return str(valor or "").strip()
    alfanum = re.sub(r"[^A-Z0-9]", "", str(valor).upper())
    if len(alfanum) == 8:
        return "%s.%s.%s-%s" % (alfanum[0], alfanum[1:4], alfanum[4:7], alfanum[7])
    return str(valor).strip()


def formatar_car(valor):
    """Formata CAR: agrupa a parte após o 2.º hífen em blocos de 4 separados por ponto.

    Formato de saída: UF-NNNNNNN-XXXX.XXXX.XXXX.XXXX.XXXX.XXXX.XXXX.XXXX
    """
    if not valor:
        return str(valor or "").strip()
    v = str(valor).strip().upper()
    m = re.match(r"^([A-Z]{2})-(\d{7})-([A-Z0-9.]+)$", v)
    if not m:
        return v
    uf = m.group(1)
    municipio = m.group(2)
    codigo = re.sub(r"\.", "", m.group(3))
    blocos = [codigo[i:i + 4] for i in range(0, len(codigo), 4)]
    return "%s-%s-%s" % (uf, municipio, ".".join(blocos))


def formatar_cns(valor):
    """Formata CNS para 00.000-0 (6 dígitos numéricos).

    Exemplos:
        019232   → 01.923-2
        01.923-2 → 01.923-2  (sem alteração)
    """
    if not valor:
        return str(valor or "").strip()
    digitos = re.sub(r"\D+", "", str(valor))
    if len(digitos) == 6:
        return "%s.%s-%s" % (digitos[0:2], digitos[2:5], digitos[5])
    return str(valor).strip()


# Palavras conhecidas que são apenas rótulos de campo (sem valor real).
# Usadas por remover_paragrafos_placeholder_vazio para detectar "CIB: {{cib}}"
# quando cib está vazio, onde sobra apenas "CIB" após remover o placeholder.
_PLACEHOLDER_LABEL_WORDS = frozenset({
    "cib", "car", "cns", "ccir", "trt", "cnpj", "cpf", "rgi",
    "matricula", "matrícula", "cartorio", "cartório", "comarca", "registro",
})

# Rótulos que, após substituição, ficam com apenas "LABEL:" e nenhum valor.
_ROTULO_VAZIO_RE = re.compile(
    r'^\s*(?:CIB|CAR|CNS|CCIR|TRT|Cart[oó]rio|Comarca|'
    r'Matr[íi]cula(?:\s+N[°o]?\.?)?|Data\s+da\s+Matr[íi]cula)'
    r'\s*:?\s*$',
    re.IGNORECASE | re.UNICODE,
)


def remover_paragrafos_placeholder_vazio(document, data):
    """Remove parágrafos onde todos os placeholders têm valor vazio e o restante
    é apenas rótulo/pontuação (ex.: 'CIB: {{cib}}' quando cib está vazio).

    Deve ser chamada ANTES de substituir_marcadores_docx.
    """
    para_remover = []
    for para in all_document_paragraphs(document):
        texto = para.text
        matches = list(PLACEHOLDER_PATTERN.finditer(texto))
        if not matches:
            continue
        todos_vazios = all(
            not data.get(normalize_key(m.group(2) or m.group(3) or m.group(4) or ""), "")
            for m in matches
        )
        if not todos_vazios:
            continue
        texto_sem_ph = PLACEHOLDER_PATTERN.sub("", texto).strip(" :,;.-/()°ºª")
        # Trata rótulos conhecidos (ex.: "CIB" de "CIB: {{cib}}") como vazios
        if texto_sem_ph.lower() in _PLACEHOLDER_LABEL_WORDS:
            texto_sem_ph = ""
        if not texto_sem_ph:
            para_remover.append(para)

    for para in para_remover:
        try:
            p_elem = para._p
            pai = p_elem.getparent()
            if pai is not None:
                pai.remove(p_elem)
        except Exception:
            pass


def remover_paragrafos_rotulo_vazio(document):
    """Remove parágrafos/células que após substituição ficaram com apenas um rótulo sem valor.

    Ex.: 'CIB: ', 'CAR: ', 'CCIR: ', 'CNS: ' → removidos.
    Deve ser chamada DEPOIS de substituir_marcadores_docx.
    """
    para_remover = []
    for para in all_document_paragraphs(document):
        if _ROTULO_VAZIO_RE.match(para.text):
            para_remover.append(para)
    for para in para_remover:
        try:
            p_elem = para._p
            pai = p_elem.getparent()
            if pai is not None:
                pai.remove(p_elem)
        except Exception:
            pass


# ── QR Code ────────────────────────────────────────────────────────────────────

def gerar_qrcode_imagem(link, output_path, size_px=1200, border_px=20):
    """Gera QR Code PNG de alta qualidade a partir de um link.

    Estratégia de qualidade:
    - ERROR_CORRECT_H: máxima correção de erros (30% de codewords restauráveis)
    - box_size=30: cada módulo do QR gerado com 30 px nativos (imagem natural ~1000px+)
    - border=4: 4 módulos de margem branca interna (padrão ISO)
    - resize para size_px com LANCZOS: reamostragem de alta qualidade
    - borda preta externa de border_px aplicada por último
    Inserir no Word em 4,17 cm × 4,17 cm resulta em ~730 DPI → QR nítido.
    Requer: qrcode[pil]  (pip install qrcode[pil])
    """
    try:
        import qrcode as _qrcode
        from PIL import Image as _Image
    except ImportError as exc:
        raise RuntimeError(
            "Dependências para QR Code ausentes. "
            "Instale com: pip install qrcode[pil] — %s" % exc
        )

    qr = _qrcode.QRCode(
        version=None,
        error_correction=_qrcode.constants.ERROR_CORRECT_H,
        box_size=30,
        border=4,
    )
    qr.add_data(link)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white").convert("RGB")

    # Redimensionar para a área interna (size_px menos duas bordas) com LANCZOS
    inner = max(50, size_px - 2 * border_px)
    img = img.resize((inner, inner), _Image.LANCZOS)

    # Colar no canvas preto para criar a borda externa
    if border_px > 0:
        canvas = _Image.new("RGB", (size_px, size_px), (0, 0, 0))
        canvas.paste(img, (border_px, border_px))
        img = canvas

    img.save(output_path, "PNG")
    return output_path


def _segmentos_texto_runs(runs, text_start, text_end):
    """Extrai [(texto, rPr_clone)] no intervalo [text_start, text_end) dos runs."""
    from copy import deepcopy
    resultado = []
    pos = 0
    for run in runs:
        rs, re = pos, pos + len(run.text)
        pos = re
        ss = max(rs, text_start)
        se = min(re, text_end)
        if ss >= se:
            continue
        t = run.text[ss - rs:se - rs]
        if t:
            rPr = run._r.rPr
            resultado.append((t, deepcopy(rPr) if rPr is not None else None))
    return resultado


def substituir_qrcode_por_imagem(document, caminho_imagem, largura_cm=2.65, altura_cm=2.65):
    """Substitui {{QR_CODE}} no documento pela imagem em caminho_imagem.

    - Sozinho no parágrafo → limpa, centraliza e insere imagem.
    - Com texto ao redor → insere inline preservando texto anterior/posterior.
    - Não encontrando o placeholder, não faz nada.
    Deve ser chamada ANTES de substituir_marcadores_docx.
    """
    try:
        from docx.shared import Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return

    from copy import deepcopy

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        match = next(
            (m for m in PLACEHOLDER_PATTERN.finditer(full_text)
             if normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "qr_code"),
            None,
        )
        if not match:
            continue

        texto_fora = PLACEHOLDER_PATTERN.sub("", full_text).strip()

        if not texto_fora:
            clear_paragraph_content(paragraph)
            try:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
            paragraph.add_run().add_picture(
                caminho_imagem, width=Cm(largura_cm), height=Cm(altura_cm)
            )
        else:
            antes = _segmentos_texto_runs(runs, 0, match.start())
            depois = _segmentos_texto_runs(runs, match.end(), len(full_text))
            clear_paragraph_content(paragraph)
            for texto, rPr in antes:
                r = paragraph.add_run(texto)
                if rPr is not None:
                    e = r._r
                    if e.rPr is not None:
                        e.remove(e.rPr)
                    e.insert(0, deepcopy(rPr))
            paragraph.add_run().add_picture(
                caminho_imagem, width=Cm(largura_cm), height=Cm(altura_cm)
            )
            for texto, rPr in depois:
                r = paragraph.add_run(texto)
                if rPr is not None:
                    e = r._r
                    if e.rPr is not None:
                        e.remove(e.rPr)
                    e.insert(0, deepcopy(rPr))
        break  # único {{QR_CODE}} por documento
