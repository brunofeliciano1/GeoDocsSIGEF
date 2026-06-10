import os
import re
import unicodedata
import tempfile
import shutil
from contextlib import contextmanager
from datetime import datetime
from qgis.PyQt.QtCore import Qt
from qgis.PyQt.QtGui import QIcon
from qgis.PyQt.QtWidgets import (
    QAction,
    QDialog,
    QFileDialog,
    QFormLayout,
    QGridLayout,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QComboBox,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from .modulos.ferramentas import ferramenta_limpar_campo_cabecalho, gerar_documentos_ferramenta_sigef
from .docx_utils import (
    formatar_trt_sem_duplicar_uf,
    formatar_municipio_para_data,
    formatar_ccir,
    formatar_cib,
    formatar_car,
    formatar_cns,
)
from .modulos.memorial import (
    fill_memorial_template,
    find_memorial_template,
    resolve_memorial_output_path,
)
from .modulos.planilha_calculo import (
    find_planilha_template,
    fill_planilha_calculo_template,
)
from .modulos.capa import (
    find_capa_template,
    fill_capa_template,
)
from .modulos.laudo_tecnico import (
    find_laudo_template,
    fill_laudo_tecnico_template,
)
from .modulos.declaracao_respeito_limites import (
    find_declaracao_template,
    fill_declaracao_template,
)
from .modulos.declaracao_confrontantes import (
    find_declaracao_confrontantes_template,
    fill_declaracao_confrontantes_template,
    nome_confinante_para_arquivo,
)
from .modulos.declaracao_confrontantes_lote import gerar_declaracoes_confrontantes_lote
from .modulos.requerimento_retificacao_area import (
    find_requerimento_retificacao_area_template,
    fill_requerimento_retificacao_area_template,
    sanitizar_nome_denominacao,
)
from .modulos.requerimento_desmembramento import (
    find_requerimento_desmembramento_template,
    fill_requerimento_desmembramento_template,
)
from .modulos.declaracao_dispensa_anuencia import (
    find_declaracao_dispensa_anuencia_template,
    fill_declaracao_dispensa_anuencia_template,
)
from .modulos.pessoas_utils import buscar_proprietarios_confinante as _buscar_pessoas_confinante


DOCUMENT_BUTTONS = [
    ("Converter PDF", "#1565C0"),
    ("Memorial", "#2E7D32"),
    ("Planilha de Cálculo", "#E65100"),
    ("Capa", "#7B1FA2"),
    ("Decl. Respeito de Limites", "#00695C"),
    ("Laudo Técnico", "#C62828"),
    ("Decl. Confrontante Indiv.", "#5D4037"),
    ("Decl. Confrontantes Lote", "#4527A0"),
    ("Req. Retif. de Área", "#283593"),
    ("Req. Desmembramento", "#AD1457"),
    ("Dispensa de Anuência", "#00838F"),
]


APP_STYLE = """
QDialog {
    background: #f4faf8;
    color: #12333b;
}
QTabWidget::pane {
    border: 1px solid #b8d8d2;
    border-radius: 8px;
    background: #ffffff;
}
QTabBar::tab {
    background: #dcefeb;
    color: #16404b;
    padding: 8px 12px;
    border-top-left-radius: 6px;
    border-top-right-radius: 6px;
    margin-right: 2px;
}
QTabBar::tab:selected {
    background: #2a8f8a;
    color: #ffffff;
}
QLabel#introLabel {
    background: #e9f6f2;
    border-left: 4px solid #2a8f8a;
    border-radius: 6px;
    padding: 10px;
    color: #12333b;
}
QLabel#sectionLabel {
    color: #0d5967;
    font-weight: 700;
    padding-top: 8px;
}
QLineEdit, QTextEdit {
    background: #ffffff;
    border: 1px solid #aaccc8;
    border-radius: 6px;
    padding: 6px;
    selection-background-color: #2a8f8a;
}
QLineEdit:focus, QTextEdit:focus {
    border: 1px solid #1976a3;
}
QPushButton {
    background: #e7f2f5;
    border: 1px solid #82b9c8;
    border-radius: 6px;
    padding: 7px 14px;
    color: #12333b;
}
QPushButton:hover {
    background: #d7ecf2;
}
QPushButton#primaryButton {
    background: #1f8f78;
    border: 1px solid #167161;
    color: #ffffff;
    font-weight: 700;
}
QPushButton#primaryButton:hover {
    background: #197a68;
}
"""


PLACEHOLDER_PATTERN = re.compile(r"(\{\{\s*([^{}]+?)\s*\}\}|<<\s*([^<>]+?)\s*>>|\$\{\s*([^{}]+?)\s*\})")


class DropFileLineEdit(QLineEdit):

    def __init__(self, extensions, parent=None):
        super().__init__(parent)
        self.extensions = tuple(extension.lower() for extension in extensions)
        self.setAcceptDrops(True)
        self.setPlaceholderText("Arraste o arquivo aqui ou clique em Procurar")

    def dragEnterEvent(self, event):
        if self._accepted_event(event):
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if self._accepted_event(event):
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event):
        urls = event.mimeData().urls()
        if not urls:
            return
        path = urls[0].toLocalFile()
        if self._accepted_path(path):
            self.setText(path)
            event.acceptProposedAction()
        else:
            event.ignore()

    def _accepted_event(self, event):
        urls = event.mimeData().urls()
        return bool(urls and self._accepted_path(urls[0].toLocalFile()))

    def _accepted_path(self, path):
        return bool(path and os.path.splitext(path)[1].lower() in self.extensions)


class GeoDocsSIGEF:

    def __init__(self, iface):
        self.iface = iface
        self.action = None
        self.dialog = None

    def initGui(self):
        icon_path = os.path.join(os.path.dirname(__file__), "icon.png")
        self.action = QAction(
            QIcon(icon_path),
            "GeoDocs SIGEF",
            self.iface.mainWindow()
        )
        self.action.triggered.connect(self.run)

        self.iface.addPluginToMenu("&GeoDocs SIGEF", self.action)
        self.iface.addToolBarIcon(self.action)

    def unload(self):
        self.iface.removePluginMenu("&GeoDocs SIGEF", self.action)
        self.iface.removeToolBarIcon(self.action)

    def run(self):
        if self.dialog is None:
            self.dialog = GeoDocsSIGEFDialog(self.iface.mainWindow(), iface=self.iface)
        self.dialog.show()
        self.dialog.raise_()
        self.dialog.activateWindow()


class GeoDocsSIGEFDialog(QDialog):

    def __init__(self, parent=None, iface=None):
        super().__init__(parent)
        self.iface = iface
        self.setWindowTitle("GeoDocs SIGEF")
        self.resize(820, 580)

        self.plugin_dir = os.path.dirname(os.path.abspath(__file__))
        self.template_path = find_memorial_template(self.plugin_dir)
        self.setStyleSheet(APP_STYLE)

        layout = QVBoxLayout(self)
        layout.setSpacing(10)

        # --- Entradas compartilhadas ---
        form = QFormLayout()
        form.setLabelAlignment(Qt.AlignRight)

        self.pdf_path = DropFileLineEdit((".pdf",))
        self.excel_path = DropFileLineEdit((".xlsx", ".xlsm"))
        self.output_path = QLineEdit()
        self.output_path.setText(default_output_directory())
        self.layer_combo = QComboBox()
        self.layer_ids = []
        self.property_id = QLineEdit()
        self.property_id.setPlaceholderText("ID do imóvel (ou selecione na camada)")

        form.addRow("PDF SIGEF:", self._file_picker_row(
            self.pdf_path, "Selecionar memorial SIGEF em PDF", "Arquivos PDF (*.pdf)"
        ))
        form.addRow("Planilha Excel (opcional):", self._file_picker_row(
            self.excel_path, "Selecionar planilha de confrontantes",
            "Planilhas Excel (*.xlsx *.xlsm)"
        ))
        form.addRow("Camada QGIS:", self._layer_picker_row())
        form.addRow("ID do imóvel:", self._property_search_row())
        form.addRow("Pasta de saída:", self._save_picker_row())
        layout.addLayout(form)

        # --- Botões de geração ---
        section = QLabel("Gerar documento")
        section.setObjectName("sectionLabel")
        layout.addWidget(section)

        grid_widget = QWidget()
        grid = QGridLayout(grid_widget)
        grid.setSpacing(8)

        handlers = [
            self._run_converter_pdf,
            self._run_memorial,
            self._run_planilha_calculo,
            self._run_capa,
            self._run_declaracao_respeito_limites,
            self._run_laudo_tecnico,
            self._run_declaracao_confrontantes,
            self._run_declaracao_confrontantes_lote,
            self._run_requerimento_retificacao_area,
            self._run_requerimento_desmembramento,
            self._run_declaracao_dispensa_anuencia,
        ] + [self._make_placeholder_handler(name) for name, _ in DOCUMENT_BUTTONS[11:]]

        for i, ((label, color), handler) in enumerate(zip(DOCUMENT_BUTTONS, handlers)):
            btn = QPushButton(label)
            btn.setMinimumHeight(50)
            btn.setStyleSheet(
                "QPushButton { background: %s; color: #ffffff; border: none;"
                " border-radius: 6px; padding: 8px 10px; font-weight: 700; font-size: 10pt; }"
                "QPushButton:hover { border: 2px solid rgba(255,255,255,180); }" % color
            )
            btn.clicked.connect(handler)
            grid.addWidget(btn, i // 3, i % 3)

        layout.addWidget(grid_widget)

        # --- Log ---
        log_label = QLabel("Log")
        log_label.setObjectName("sectionLabel")
        layout.addWidget(log_label)

        self.log = QTextEdit()
        self.log.setReadOnly(True)
        self.log.setPlaceholderText("O resultado das operações aparece aqui.")
        self.log.setMaximumHeight(100)
        layout.addWidget(self.log)

        self._refresh_layers()

    # ---- handlers dos botões ----

    def _make_placeholder_handler(self, name):
        def handler():
            QMessageBox.information(
                self, "GeoDocs SIGEF",
                "A função '%s' ainda não foi implementada." % name
            )
        return handler

    def _run_converter_pdf(self):
        try:
            pdf_path = self.pdf_path.text().strip()
            output_folder = self.output_path.text().strip() or default_output_directory()
            if not pdf_path or not os.path.exists(pdf_path):
                raise RuntimeError("Selecione um PDF SIGEF válido.")
            if not os.path.isdir(output_folder):
                os.makedirs(output_folder, exist_ok=True)
            result = gerar_documentos_ferramenta_sigef(pdf_path, output_folder)
            self.log.setPlainText(
                "Converter PDF concluído.\n\nVértices: {vertices}\n"
                "Memorial: {memorial}\nPlanilha: {planilha}\nAnuências: {anuencias}".format(**result)
            )
            QMessageBox.information(self, "GeoDocs SIGEF", "Documentos gerados com sucesso!")
        except Exception as exc:
            self._show_error(str(exc))

    def _run_memorial(self):
        output_path = _safe_output_path(resolve_memorial_output_path(self.output_path.text().strip()))
        try:
            data = self._collect_memorial_data()
            layer = self._selected_qgis_layer()
            replaced_count = fill_memorial_template(self.template_path, output_path, data, layer)
            self.log.setPlainText(
                "Memorial gerado.\n\nModelo: %s\nMarcadores substituídos: %s\nArquivo: %s"
                % (self.template_path, replaced_count, output_path)
            )
            if data.get("__aviso_confrontantes"):
                self.log.append("\nAviso: %s" % data["__aviso_confrontantes"])
            if data.get("__debug_confrontantes"):
                self.log.append("\n── Debug confrontantes ──\n%s" % data["__debug_confrontantes"])
            QMessageBox.information(self, "GeoDocs SIGEF", "Memorial gerado com sucesso!")
        except Exception as exc:
            self._show_error(str(exc))

    def _run_capa(self):
        output_folder = self.output_path.text().strip() or default_output_directory()
        output_path = _safe_output_path(os.path.join(output_folder, "CAPA.docx"))
        try:
            template_path = find_capa_template(self.plugin_dir)
            if not template_path:
                raise RuntimeError(
                    "Modelo capa.docx não encontrado na pasta models/."
                )
            data = self._collect_memorial_data()
            layer = self._selected_qgis_layer()
            fill_capa_template(template_path, output_path, data, layer)
            self.log.setPlainText("Capa gerada.\n\nArquivo: %s" % output_path)
            QMessageBox.information(self, "GeoDocs SIGEF", "Capa gerada com sucesso!")
        except Exception as exc:
            self._show_error(str(exc))

    def _run_declaracao_respeito_limites(self):
        output_folder = self.output_path.text().strip() or default_output_directory()
        output_path = _safe_output_path(os.path.join(output_folder, "DECLARACAO_RESPEITO_LIMITES.docx"))
        try:
            template_path = find_declaracao_template(self.plugin_dir)
            if not template_path:
                raise RuntimeError(
                    "Modelo declaracao_respeito_limites.docx não encontrado na pasta models/."
                )
            data = self._collect_memorial_data()
            layer = self._selected_qgis_layer()
            avisos = fill_declaracao_template(template_path, output_path, data, layer)
            msg = "Declaração de Respeito de Limites gerada.\n\nArquivo: %s" % output_path
            if avisos:
                msg += "\n\nAvisos:\n" + "\n".join(avisos)
            if data.get("__aviso_confrontantes"):
                msg += "\n\nAviso confrontantes: %s" % data["__aviso_confrontantes"]
            self.log.setPlainText(msg)
            if data.get("__debug_confrontantes"):
                self.log.append("\n── Debug confrontantes ──\n%s" % data["__debug_confrontantes"])
            QMessageBox.information(
                self, "GeoDocs SIGEF", "Declaração de Respeito de Limites gerada com sucesso!"
            )
        except Exception as exc:
            self._show_error(str(exc))

    def _run_laudo_tecnico(self):
        output_folder = self.output_path.text().strip() or default_output_directory()
        output_path = _safe_output_path(os.path.join(output_folder, "LAUDO_TECNICO.docx"))
        try:
            template_path = find_laudo_template(self.plugin_dir)
            if not template_path:
                raise RuntimeError(
                    "Modelo laudo_tecnico.docx não encontrado na pasta models/."
                )
            data = self._collect_memorial_data()
            layer = self._selected_qgis_layer()
            avisos = fill_laudo_tecnico_template(template_path, output_path, data, layer)
            msg = "Laudo Técnico gerado.\n\nArquivo: %s" % output_path
            if avisos:
                msg += "\n\nAvisos:\n" + "\n".join(avisos)
            self.log.setPlainText(msg)
            QMessageBox.information(self, "GeoDocs SIGEF", "Laudo Técnico gerado com sucesso!")
        except Exception as exc:
            self._show_error(str(exc))

    def _run_planilha_calculo(self):
        output_folder = self.output_path.text().strip() or default_output_directory()
        output_path = _safe_output_path(os.path.join(output_folder, "PLANILHA_DE_CALCULO.docx"))
        try:
            template_path = find_planilha_template(self.plugin_dir)
            if not template_path:
                raise RuntimeError(
                    "Modelo planilha_de_calculo.docx não encontrado na pasta models/."
                )
            data = self._collect_memorial_data()
            layer = self._selected_qgis_layer()
            vertices = fill_planilha_calculo_template(template_path, output_path, data, layer)
            self.log.setPlainText(
                "Planilha de Cálculo gerada.\n\nVértices: %s\nArquivo: %s" % (vertices, output_path)
            )
            QMessageBox.information(self, "GeoDocs SIGEF", "Planilha de Cálculo gerada com sucesso!")
        except Exception as exc:
            self._show_error(str(exc))

    def _run_declaracao_confrontantes_lote(self):
        output_folder = self.output_path.text().strip() or default_output_directory()
        try:
            template_path = find_declaracao_confrontantes_template(self.plugin_dir)
            if not template_path:
                raise RuntimeError(
                    "Modelo declaracao_de_confrontantes.docx não encontrado na pasta models/."
                )

            pdf_path = self.pdf_path.text().strip()
            if not pdf_path or not os.path.exists(pdf_path):
                raise RuntimeError("Selecione um memorial SIGEF em PDF válido.")

            data = self._collect_memorial_data()
            servico_id = data.get("__servico_id")
            if not servico_id:
                raise RuntimeError(
                    "Não foi possível identificar o ID do serviço. "
                    "Selecione uma feição na camada de serviços."
                )

            confinante_layer = obter_camada_confinantes(self.iface)
            if confinante_layer is None:
                raise RuntimeError(
                    "Camada Confinantes Principal não encontrada no projeto. "
                    "Adicione a camada e tente novamente."
                )

            pdf_data = read_pdf_data(pdf_path)
            pdf_segments = pdf_data.get("__pdf_segments", [])
            denominacao = data.get("denominacao", "")

            gerados, falhas = gerar_declaracoes_confrontantes_lote(
                template_path,
                output_folder,
                confinante_layer,
                servico_id,
                pdf_segments,
                pdf_data,
                denominacao=denominacao,
            )

            msg = "Foram geradas %d declarações de confrontantes." % len(gerados)
            if gerados:
                msg += "\n\nArquivos:\n" + "\n".join(gerados)
            if falhas:
                msg += "\n\n%d apresentaram erro:\n" % len(falhas)
                msg += "\n".join("• %s: %s" % (nome, err) for nome, err in falhas)
            self.log.setPlainText(msg)

            if falhas:
                QMessageBox.warning(
                    self,
                    "GeoDocs SIGEF",
                    "Foram geradas %d declarações.\n%d apresentaram erro."
                    % (len(gerados), len(falhas)),
                )
            else:
                QMessageBox.information(
                    self,
                    "GeoDocs SIGEF",
                    "Foram geradas %d declarações de confrontantes com sucesso!" % len(gerados),
                )
        except Exception as exc:
            self._show_error(str(exc))

    def _run_requerimento_retificacao_area(self):
        output_folder = self.output_path.text().strip() or default_output_directory()
        try:
            template_path = find_requerimento_retificacao_area_template(self.plugin_dir)
            if not template_path:
                raise RuntimeError(
                    "Modelo requerimento_retificacao_area.docx não encontrado na pasta models/."
                )

            data = self._collect_memorial_data()
            layer, feature = self._find_qgis_feature()

            denominacao = data.get("denominacao", "").strip() or "REQUERIMENTO"
            nome_sanitizado = sanitizar_nome_denominacao(denominacao)
            output_path = _safe_output_path(
                os.path.join(
                    output_folder,
                    "REQUERIMENTO_RETIFICACAO_AREA_%s.docx" % nome_sanitizado,
                )
            )

            avisos = fill_requerimento_retificacao_area_template(
                template_path, output_path, data, layer, feature
            )

            msg = "Requerimento de Retificação de Área gerado.\n\nArquivo: %s" % output_path
            if avisos:
                msg += "\n\nAvisos:\n" + "\n".join(avisos)
            self.log.setPlainText(msg)
            QMessageBox.information(
                self, "GeoDocs SIGEF", "Requerimento de Retificação de Área gerado com sucesso!"
            )
        except Exception as exc:
            self._show_error(str(exc))

    def _run_requerimento_desmembramento(self):
        output_folder = self.output_path.text().strip() or default_output_directory()
        try:
            template_path = find_requerimento_desmembramento_template(self.plugin_dir)
            if not template_path:
                raise RuntimeError(
                    "Modelo requerimento_desmembramento.docx não encontrado na pasta models/."
                )

            # ── Captura a seleção múltipla ANTES de _collect_memorial_data,
            # pois este método chama layer.removeSelection() internamente.
            layer = self._selected_qgis_layer()
            if layer is None:
                raise RuntimeError("Selecione a camada Serviços 2 na lista de camadas.")

            selected = list(layer.selectedFeatures())
            if not selected:
                raise RuntimeError(
                    "Selecione as áreas do desmembramento na camada Serviços 2."
                )

            # ── Valida desm_tipo antes de qualquer processamento pesado
            sem_tipo = [
                str(f.id())
                for f in selected
                if not feature_value(layer, f, "desm_tipo").strip()
            ]
            if sem_tipo:
                raise RuntimeError(
                    "Há áreas selecionadas sem o campo desm_tipo preenchido "
                    "(id: %s)." % ", ".join(sem_tipo)
                )

            remanescentes = [
                f for f in selected
                if feature_value(layer, f, "desm_tipo").strip().upper() == "REMANESCENTE"
            ]
            desmembradas = [
                f for f in selected
                if feature_value(layer, f, "desm_tipo").strip().upper() == "DESMEMBRADA"
            ]

            if not remanescentes:
                raise RuntimeError("Nenhuma Área Remanescente foi selecionada.")
            if len(remanescentes) > 1:
                raise RuntimeError("Mais de uma Área Remanescente foi selecionada.")
            if not desmembradas:
                raise RuntimeError("Nenhuma Área Desmembrada foi selecionada.")

            # ── Dados gerais: vêm do ID digitado no plugin (serviço base)
            # _collect_memorial_data altera a seleção, mas salvamos selected acima.
            data = self._collect_memorial_data()

            # Feature base = feature encontrada pelo ID digitado (não pela seleção)
            _, feature_base = self._find_qgis_feature()

            denominacao = data.get("denominacao", "").strip() or "REQUERIMENTO"
            nome_sanitizado = sanitizar_nome_denominacao(denominacao)
            output_path = _safe_output_path(
                os.path.join(
                    output_folder,
                    "REQUERIMENTO_DESMEMBRAMENTO_%s.docx" % nome_sanitizado,
                )
            )

            # feature_base fornece os dados gerais; selected_features fornece as áreas
            avisos = fill_requerimento_desmembramento_template(
                template_path, output_path, data, layer, feature_base,
                selected_features=selected
            )

            msg = "Requerimento de Desmembramento gerado.\n\nArquivo: %s" % output_path
            if avisos:
                msg += "\n\nAvisos:\n" + "\n".join(avisos)
            self.log.setPlainText(msg)
            QMessageBox.information(
                self, "GeoDocs SIGEF", "Requerimento de Desmembramento gerado com sucesso!"
            )
        except Exception as exc:
            self._show_error(str(exc))

    def _run_declaracao_dispensa_anuencia(self):
        output_folder = self.output_path.text().strip() or default_output_directory()
        try:
            template_path = find_declaracao_dispensa_anuencia_template(self.plugin_dir)
            if not template_path:
                raise RuntimeError(
                    "Modelo declaracao_dispensa_anuencia.docx não encontrado na pasta models/."
                )

            # Camada de confinantes e features selecionadas (capturadas ANTES de
            # _collect_memorial_data para garantir que a seleção não seja alterada)
            confinante_layer = obter_camada_confinantes(self.iface)
            if confinante_layer is None:
                raise RuntimeError(
                    "Camada Confinantes Principal não encontrada no projeto. "
                    "Certifique-se de que a camada está carregada."
                )
            confinante_features = list(confinante_layer.selectedFeatures())
            if not confinante_features:
                raise RuntimeError(
                    "Selecione ao menos um confrontante na camada Confinantes Principal "
                    "para gerar a Declaração de Dispensa de Anuência."
                )

            # Dados do serviço principal (imóvel base)
            data = self._collect_memorial_data()
            serv_layer, serv_feature = self._find_qgis_feature()

            denominacao = data.get("denominacao", "").strip() or "DISPENSA_ANUENCIA"
            nome_sanitizado = sanitizar_nome_denominacao(denominacao)
            output_path = _safe_output_path(
                os.path.join(
                    output_folder,
                    "DECLARACAO_DISPENSA_ANUENCIA_%s.docx" % nome_sanitizado,
                )
            )

            avisos = fill_declaracao_dispensa_anuencia_template(
                template_path, output_path, data, serv_layer, serv_feature,
                confinante_layer=confinante_layer,
                confinante_features=confinante_features,
            )

            msg = "Declaração de Dispensa de Anuência gerada.\n\nArquivo: %s" % output_path
            if avisos:
                msg += "\n\nAvisos:\n" + "\n".join(avisos)
            self.log.setPlainText(msg)
            QMessageBox.information(
                self, "GeoDocs SIGEF",
                "Declaração de Dispensa de Anuência gerada com sucesso!"
            )
        except Exception as exc:
            self._show_error(str(exc))

    def _run_declaracao_confrontantes(self):
        output_folder = self.output_path.text().strip() or default_output_directory()
        try:
            template_path = find_declaracao_confrontantes_template(self.plugin_dir)
            if not template_path:
                raise RuntimeError(
                    "Modelo declaracao_de_confrontantes.docx não encontrado na pasta models/."
                )

            pdf_path = self.pdf_path.text().strip()
            if not pdf_path or not os.path.exists(pdf_path):
                raise RuntimeError("Selecione um memorial SIGEF em PDF válido.")

            # Diagnóstico: mostra estado atual das camadas antes de tentar localizar o confinante
            diag_lines = []
            try:
                active = self.iface.activeLayer()
                diag_lines.append(
                    "Camada ativa: %s" % (active.name() if active else "(nenhuma)")
                )
                from qgis.core import QgsProject
                for lyr in QgsProject.instance().mapLayers().values():
                    sel_count = lyr.selectedFeatureCount()
                    if sel_count > 0:
                        diag_lines.append(
                            "Camada '%s': %d feição(ões) selecionada(s)" % (lyr.name(), sel_count)
                        )
            except Exception:
                pass

            layer, confinante_feature = obter_confinante_selecionado(self.iface)

            if confinante_feature is None:
                diag = "\n".join(diag_lines) if diag_lines else "(sem informação de diagnóstico)"
                raise RuntimeError(
                    "Nenhum confinante selecionado na camada Confinantes Principal.\n\n"
                    "Diagnóstico:\n%s\n\n"
                    "Selecione uma feição na camada antes de gerar a declaração." % diag
                )

            try:
                confinante_id_diag = str(confinante_feature["id"] or "").strip()
            except Exception:
                confinante_id_diag = "(campo id não encontrado)"
            diag_lines.append(
                "Confinante encontrado: camada '%s', id=%s" % (layer.name(), confinante_id_diag)
            )
            self.log.setPlainText("\n".join(diag_lines))

            try:
                codigo_raw = str(confinante_feature["codigo"]).strip()
            except Exception:
                codigo_raw = ""
            if not codigo_raw or codigo_raw.upper() in ("NULL", "NONE", ""):
                raise RuntimeError(
                    "O confinante selecionado não possui campo 'codigo' preenchido. "
                    "Verifique o cadastro na camada."
                )

            pdf_data = read_pdf_data(pdf_path)
            pdf_segments = pdf_data.get("__pdf_segments", [])

            nome_sanitizado = nome_confinante_para_arquivo(layer, confinante_feature)
            output_path = _safe_output_path(
                os.path.join(
                    output_folder,
                    "DECLARACAO_CONFRONTANTE_%s.docx" % nome_sanitizado,
                )
            )

            avisos = fill_declaracao_confrontantes_template(
                template_path, output_path, layer, confinante_feature, pdf_segments, pdf_data
            )

            msg = "Declaração de Confrontantes gerada.\n\nArquivo: %s" % output_path
            msg += "\n\nDiagnóstico:\n" + "\n".join(diag_lines)
            if avisos:
                msg += "\n\nAvisos:\n" + "\n".join(avisos)
            self.log.setPlainText(msg)
            QMessageBox.information(
                self, "GeoDocs SIGEF", "Declaração de Confrontantes gerada com sucesso!"
            )
        except Exception as exc:
            self._show_error(str(exc))

    # ---- utilitários de UI ----

    def _file_picker_row(self, line_edit, title, file_filter):
        row = QWidget()
        row_layout = QHBoxLayout(row)
        row_layout.setContentsMargins(0, 0, 0, 0)
        button = QPushButton("Procurar")
        button.clicked.connect(lambda: self._select_input_file(line_edit, title, file_filter))
        row_layout.addWidget(line_edit)
        row_layout.addWidget(button)
        return row

    def _save_picker_row(self):
        row = QWidget()
        row_layout = QHBoxLayout(row)
        row_layout.setContentsMargins(0, 0, 0, 0)
        button = QPushButton("Selecionar pasta")
        button.clicked.connect(self._select_output_file)
        row_layout.addWidget(self.output_path)
        row_layout.addWidget(button)
        return row

    def _layer_picker_row(self):
        row = QWidget()
        row_layout = QHBoxLayout(row)
        row_layout.setContentsMargins(0, 0, 0, 0)
        refresh_button = QPushButton("Atualizar")
        refresh_button.clicked.connect(self._refresh_layers)
        row_layout.addWidget(self.layer_combo)
        row_layout.addWidget(refresh_button)
        return row

    def _property_search_row(self):
        row = QWidget()
        row_layout = QHBoxLayout(row)
        row_layout.setContentsMargins(0, 0, 0, 0)
        search_button = QPushButton("Selecionar imóvel")
        search_button.clicked.connect(self._select_qgis_feature)
        row_layout.addWidget(self.property_id)
        row_layout.addWidget(search_button)
        return row

    def _select_input_file(self, line_edit, title, file_filter):
        path, _ = QFileDialog.getOpenFileName(self, title, "", file_filter)
        if path:
            line_edit.setText(path)

    def _select_output_file(self):
        path = QFileDialog.getExistingDirectory(
            self,
            "Selecionar pasta de saída",
            self.output_path.text().strip() or default_output_directory()
        )
        if path:
            self.output_path.setText(path)

    def _show_error(self, message):
        self.log.setPlainText(message)
        QMessageBox.warning(self, "GeoDocs SIGEF", message)

    # ---- lógica de dados ----

    def _collect_memorial_data(self):
        self.template_path = find_memorial_template(self.plugin_dir)
        if not self.template_path:
            raise RuntimeError(
                "Modelo base não encontrado. Coloque memorial.docx na pasta models do plugin."
            )

        pdf_path = self.pdf_path.text().strip()
        if not pdf_path or not os.path.exists(pdf_path):
            raise RuntimeError("Selecione um memorial SIGEF em PDF válido.")

        data = {}
        pdf_data = read_pdf_data(pdf_path)

        excel_path = self.excel_path.text().strip() if hasattr(self, "excel_path") else ""
        if excel_path:
            if not os.path.exists(excel_path):
                raise RuntimeError("A planilha Excel informada não foi encontrada.")
            excel_segments = read_excel_confrontacoes_segments(excel_path)
            if not excel_segments:
                raise RuntimeError(
                    "Não foi possível localizar as colunas Código, Vante e Confrontação na planilha Excel."
                )
            merged_segments = merge_confrontacoes_pdf_excel(
                pdf_data.get("__pdf_segments", []),
                excel_segments
            )
            pdf_data["__pdf_segments"] = merged_segments
            pdf_data["descricao_perimetro"] = build_descricao_perimetro_from_segments(
                pdf_data, merged_segments
            )

        merge_data(data, pdf_data, overwrite=True)
        remove_qgis_only_fields(data)

        layer, feature = self._find_qgis_feature()
        confinante_layer = obter_camada_confinantes(self.iface)
        if layer and feature:
            layer.removeSelection()
            layer.select(feature.id())
            qgis_data = qgis_feature_to_data(layer, feature, [])
            merge_data(data, qgis_data, overwrite=True)

        # ── Enriquecer confrontantes na descrição do perímetro ────────────────
        pdf_segments = pdf_data.get("__pdf_segments", [])
        _servico_id = data.get("__servico_id")
        if pdf_segments and confinante_layer and _servico_id is not None:
            _enriched = enriquecer_confrontantes_segments_com_camada(
                pdf_segments, confinante_layer, _servico_id
            )
            # Atualiza data["__pdf_segments"] para que a Planilha de Cálculo
            # também receba os confinantes enriquecidos da camada
            data["__pdf_segments"] = _enriched
            data["descricao_perimetro"] = build_descricao_perimetro_from_segments(
                pdf_data, _enriched
            )

        # ── Confrontantes por lado — ÚNICA função autorizada ──────────────────
        confs = resolver_confrontantes_por_lado_unico(
            layer if (layer and feature) else None,
            feature if (layer and feature) else None,
            confinante_layer,
            pdf_segments,
        )
        for _k, _v in confs.items():
            data[_k] = _v

        # Fallback: se QGIS não tem cartório/CNS, usa o extraído do cabeçalho do PDF
        if not data.get("cartorio") and data.get("__cartorio_pdf"):
            data["cartorio"] = data["__cartorio_pdf"]
        if not data.get("nome_cartorio") and data.get("__cartorio_pdf"):
            data["nome_cartorio"] = data["__cartorio_pdf"]
        if not data.get("cns") and data.get("__cns_pdf"):
            raw_cns = data["__cns_pdf"]
            data["cns"] = "(%s)" % raw_cns
            data["cns_sem_parenteses"] = raw_cns

        add_aliases(data)
        finalize_data_fields(data)

        if not data:
            raise RuntimeError(
                "Não foi possível extrair dados. Confira se o PDF possui texto selecionável "
                "e se a camada QGIS foi selecionada."
            )

        return data

    def _refresh_layers(self):
        current_layer_id = (
            self.layer_ids[self.layer_combo.currentIndex()]
            if self.layer_ids and self.layer_combo.currentIndex() >= 0
            else None
        )
        self.layer_combo.clear()
        self.layer_ids = []
        try:
            from qgis.core import QgsMapLayer, QgsProject
        except Exception:
            self.layer_combo.addItem("QGIS indisponível")
            return

        for layer in QgsProject.instance().mapLayers().values():
            if layer.type() == QgsMapLayer.VectorLayer:
                self.layer_combo.addItem(layer.name())
                self.layer_ids.append(layer.id())

        if not self.layer_ids:
            self.layer_combo.addItem("Nenhuma camada vetorial aberta")
            return

        if current_layer_id in self.layer_ids:
            self.layer_combo.setCurrentIndex(self.layer_ids.index(current_layer_id))

    def _selected_qgis_layer(self):
        if not self.layer_ids or self.layer_combo.currentIndex() < 0:
            return None
        try:
            from qgis.core import QgsProject
        except Exception:
            return None
        return QgsProject.instance().mapLayer(self.layer_ids[self.layer_combo.currentIndex()])

    def _select_qgis_feature(self):
        try:
            layer, feature = self._find_qgis_feature()
            if layer is None or feature is None:
                self._show_error("Imóvel não encontrado na camada selecionada.")
                return
            layer.removeSelection()
            layer.select(feature.id())
            self.log.setPlainText("Imóvel selecionado na camada: %s" % layer.name())
        except Exception as exc:
            self._show_error(str(exc))

    def _find_qgis_feature(self):
        layer = self._selected_qgis_layer()
        if layer is None:
            return None, None

        property_id = self.property_id.text().strip()
        selected = list(layer.selectedFeatures())
        if not property_id and selected:
            return layer, selected[0]
        if not property_id:
            return layer, None

        id_field = field_name_lookup(layer).get("id")
        if not id_field:
            raise RuntimeError("A camada selecionada não possui campo 'id'.")

        for feature in layer.getFeatures():
            if ids_match(feature[id_field], property_id):
                return layer, feature
        return layer, None

    def _collect_qgis_data(self, pdf_segments=None):
        try:
            from qgis.core import QgsMessageLog, Qgis as _Qi

            def _tlog(msg):
                QgsMessageLog.logMessage(msg, "GeoDocsSIGEF", _Qi.Info)
        except Exception:
            def _tlog(msg):
                pass

        layer, feature = self._find_qgis_feature()
        _tlog("[Trace] _collect_qgis_data: layer=%s | feature=%s"
              % (layer.name() if layer else "None",
                 "encontrada" if feature else "None"))

        if layer is None or feature is None:
            _tlog("[Trace] _collect_qgis_data: retornando {} — sem camada/feição")
            return {}

        layer.removeSelection()
        layer.select(feature.id())
        confinante_layer = obter_camada_confinantes(self.iface)
        _tlog("[Trace] _collect_qgis_data: confinante_layer=%s"
              % (confinante_layer.name() if confinante_layer else "None"))

        return qgis_feature_to_data(
            layer, feature, pdf_segments or [],
            confinante_layer=confinante_layer,
        )


class MemorialPreviewDialog(QDialog):

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Prévia do Memorial")
        self.resize(900, 680)
        self.setStyleSheet(APP_STYLE)

        layout = QVBoxLayout(self)
        self.preview = QTextEdit()
        self.preview.setReadOnly(True)
        layout.addWidget(self.preview)

        buttons = QHBoxLayout()
        close_button = QPushButton("Fechar")
        close_button.clicked.connect(self.close)
        buttons.addStretch()
        buttons.addWidget(close_button)
        layout.addLayout(buttons)

    def set_preview_text(self, text):
        self.preview.setPlainText(text)


def default_output_directory():
    downloads = os.path.join(os.path.expanduser("~"), "Downloads")
    return downloads if os.path.isdir(downloads) else os.path.expanduser("~")


def obter_confinante_selecionado(iface):
    """Localiza a camada Confinantes Principal e retorna (layer, feature) da feição selecionada.

    Estratégias em ordem:
    1. Camada ativa — se for confinante (nome ou fonte), usa a seleção dela.
    2. Busca pelo nome visual "Confinantes Principal" em todo o projeto.
    3. Busca pela fonte PostgreSQL que contenha 'servicos' e 'confinante'.

    Retorna (None, None) se nenhuma feição selecionada for encontrada.
    """
    try:
        from qgis.core import QgsProject

        def _is_confinante_layer(layer):
            try:
                nome = layer.name().strip().lower()
                source = layer.source().lower()
                return (
                    nome == "confinantes principal" or
                    ("servicos" in source and "confinante" in source)
                )
            except Exception:
                return False

        # 1. Camada ativa
        layer = iface.activeLayer()
        if layer and _is_confinante_layer(layer):
            selected = list(layer.selectedFeatures())
            if selected:
                return layer, selected[0]

        # 2. Nome visual exato
        for layer in QgsProject.instance().mapLayersByName("Confinantes Principal"):
            selected = list(layer.selectedFeatures())
            if selected:
                return layer, selected[0]

        # 3. Varredura pela fonte PostgreSQL
        for layer in QgsProject.instance().mapLayers().values():
            if not _is_confinante_layer(layer):
                continue
            selected = list(layer.selectedFeatures())
            if selected:
                return layer, selected[0]

    except Exception:
        pass

    return None, None


def obter_camada_confinantes(iface):
    """Localiza a camada Confinantes Principal sem exigir seleção.

    Usa as mesmas estratégias de obter_confinante_selecionado, porém
    retorna apenas a camada (independente de haver feature selecionada).
    """
    try:
        from qgis.core import QgsProject

        def _is_confinante_layer(layer):
            try:
                nome = layer.name().strip().lower()
                source = layer.source().lower()
                return (
                    nome == "confinantes principal" or
                    ("servicos" in source and "confinante" in source)
                )
            except Exception:
                return False

        layer = iface.activeLayer()
        if layer and _is_confinante_layer(layer):
            return layer

        layers = QgsProject.instance().mapLayersByName("Confinantes Principal")
        if layers:
            return layers[0]

        for layer in QgsProject.instance().mapLayers().values():
            if _is_confinante_layer(layer):
                return layer

    except Exception:
        pass

    return None


@contextmanager
def _arquivo_temp(path):
    """Copia arquivo para pasta temporária, yield do caminho da cópia e limpa ao sair.

    Garante que a biblioteca (pdfplumber, openpyxl…) nunca segure o arquivo
    original do usuário. Se a cópia falhar, faz yield do caminho original como
    fallback (sem deletar nada ao final).
    """
    temp_dir = None
    work_path = path
    try:
        temp_dir = tempfile.mkdtemp(prefix="geodocssigef_")
        work_path = os.path.join(temp_dir, os.path.basename(path))
        shutil.copy2(path, work_path)
    except Exception:
        temp_dir = None
        work_path = path
    try:
        yield work_path
    finally:
        if temp_dir is not None:
            shutil.rmtree(temp_dir, ignore_errors=True)


def _safe_output_path(path):
    """Se o arquivo já existir (possivelmente aberto no Word/Excel), adiciona
    timestamp ao nome para não sobrescrever e não gerar erro de acesso."""
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return "%s_%s%s" % (base, ts, ext)


def ensure_docx_path(path):
    base, extension = os.path.splitext(path)
    if extension.lower() == ".docx":
        return path
    if extension:
        return base + ".docx"
    return path + ".docx"


def read_excel_confrontacoes_segments(path):
    """Lê confrontações de uma planilha Excel e retorna segmentos mínimos.

    Relação usada: Código + Vante.
    Colunas esperadas: Código, Vante e Confrontação.
    Caso os cabeçalhos não sejam encontrados, usa o padrão gerado pela ferramenta:
    A = Código, E = Vante, H = Confrontação.
    """
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError(
            "Biblioteca openpyxl não encontrada. Instale openpyxl no Python do QGIS."
        ) from exc

    with _arquivo_temp(path) as temp_path:
        try:
            wb = load_workbook(temp_path, data_only=True, read_only=True)
        except Exception as exc:
            raise RuntimeError(
                "Não foi possível abrir a planilha Excel. "
                "Feche o arquivo no Excel e tente novamente."
            ) from exc

        try:
            ws = wb.active
            header_row = None
            column_map = {}

            max_header_scan = min(ws.max_row or 1, 25)
            for row_idx in range(1, max_header_scan + 1):
                values = [cell.value for cell in ws[row_idx]]
                normalized_headers = [normalize_key(value) for value in values]

                temp = {}
                for idx, header in enumerate(normalized_headers, start=1):
                    if header in ("codigo", "cod", "vertice_codigo") and "codigo" not in temp:
                        temp["codigo"] = idx
                    elif header in ("vante", "codigo_vante", "segmento_vante") and "vante" not in temp:
                        temp["vante"] = idx
                    elif header in ("confrontacao", "confrontacoes", "confrontante", "confinante", "confinantes") and "confrontacao" not in temp:
                        temp["confrontacao"] = idx

                if {"codigo", "vante", "confrontacao"}.issubset(temp):
                    header_row = row_idx
                    column_map = temp
                    break

            # Fallback para a planilha padrão gerada pela ferramenta:
            # A Código, E Vante, H Confrontação.
            if not column_map:
                header_row = 1
                column_map = {"codigo": 1, "vante": 5, "confrontacao": 8}

            segments = []
            start_row = (header_row or 1) + 1

            for row in ws.iter_rows(min_row=start_row, values_only=True):
                def get_col(name):
                    idx = column_map.get(name, 0)
                    if not idx or idx > len(row):
                        return ""
                    value = row[idx - 1]
                    return clean_value(value)

                codigo = get_col("codigo").upper()
                vante = get_col("vante").upper()
                confrontacao = get_col("confrontacao")

                if not codigo or not vante or not confrontacao:
                    continue

                # Ignora linhas de cabeçalho repetidas.
                if normalize_key(codigo) == "codigo" or normalize_key(vante) == "vante":
                    continue

                segments.append({
                    "codigo": codigo,
                    "vante": vante,
                    "confrontacao": normalize_confrontante(confrontacao),
                })

            return segments

        finally:
            try:
                wb.close()
            except Exception:
                pass


def merge_confrontacoes_pdf_excel(pdf_segments, excel_segments):
    """Substitui somente a confrontação dos segmentos do PDF pela planilha.

    A geometria, os vértices, azimutes, distâncias e a ordem continuam sendo do PDF.
    A correspondência é feita por Código + Vante.
    Se uma linha do PDF não existir na planilha, mantém a confrontação original do PDF.
    """
    if not pdf_segments or not excel_segments:
        return pdf_segments or []

    excel_index = {}
    for row in excel_segments:
        codigo = clean_value(row.get("codigo", "")).upper()
        vante = clean_value(row.get("vante", "")).upper()
        confrontacao = normalize_confrontante(row.get("confrontacao", ""))
        if codigo and vante and confrontacao:
            excel_index[(codigo, vante)] = confrontacao

    merged = []
    for segment in pdf_segments:
        item = dict(segment)
        codigo = clean_value(item.get("codigo", "")).upper()
        vante = clean_value(item.get("vante", "")).upper()
        confrontacao_excel = excel_index.get((codigo, vante))
        if confrontacao_excel:
            item["confrontacao"] = confrontacao_excel
        merged.append(item)

    return merged


def read_pdf_data(path):
    with _arquivo_temp(path) as temp_path:
        text = normalize_text(extract_pdf_text(temp_path))
    data = extract_key_value_pairs(text)
    remove_pdf_disallowed_fields(data)
    segments = extract_pdf_segments(text)
    data["__pdf_segments"] = segments
    extract_pdf_alias_fields(text, data)

    descricao = build_descricao_perimetro_from_segments(data, segments)
    if not descricao:
        descricao = extract_descricao_perimetro(text)
    data["descricao_perimetro"] = descricao

    data["__pdf_text"] = text
    finalize_data_fields(data)
    return data


def remove_pdf_disallowed_fields(data):
    blocked = []
    for key in data:
        normalized = normalize_key(key)
        if (
            normalized in ("cpf", "cnpj", "cartorio", "cns", "matricula", "matricula_do_imovel")
            or normalized.startswith("cpf_")
            or normalized.startswith("cnpj_")
            or normalized.startswith("proprietario")
        ):
            blocked.append(key)
    for key in blocked:
        data.pop(key, None)


def remove_qgis_only_fields(data):
    blocked = []
    for key in data:
        normalized = normalize_key(key)
        if (
            normalized in ("cpf", "cnpj", "cartorio", "cns", "bloco_proprietarios", "bloco_assinaturas_proprietarios")
            or normalized.startswith("cpf_")
            or normalized.startswith("cnpj_")
            or normalized.startswith("proprietario")
        ):
            blocked.append(key)
    for key in blocked:
        data.pop(key, None)


def extract_pdf_text(path):
    """Extrai texto do PDF. Ordem de tentativa: fitz > pdfplumber > pypdf.

    fitz (PyMuPDF) é prioridade por ser mais estável dentro do QGIS e não
    causar access violation. pdfplumber preserva melhor a tabela de vértices
    do SIGEF, mas pode crashar em alguns PDFs. pypdf é o fallback final.
    """
    text = _extract_pdf_fitz(path)
    if text.strip():
        return text
    text = _extract_pdf_pdfplumber(path)
    if text.strip():
        return text
    return _extract_pdf_pypdf(path)


def _extract_pdf_fitz(path):
    """Extrai texto com PyMuPDF (fitz) — mais estável, sem risco de access violation."""
    try:
        import fitz  # PyMuPDF
        texts = []
        with fitz.open(path) as doc:
            for page in doc:
                try:
                    texts.append(page.get_text("text") or "")
                except Exception:
                    texts.append("")
        return "\n".join(texts)
    except Exception:
        return ""


def _extract_pdf_pdfplumber(path):
    """Extrai texto com pdfplumber — preserva tabela SIGEF, mas pode crashar."""
    texts = []
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                try:
                    texts.append(page.extract_text() or "")
                except Exception:
                    texts.append("")
        return "\n".join(texts)
    except Exception:
        return ""


def _extract_pdf_pypdf(path):
    """Extrai texto com pypdf/PyPDF2 — fallback final."""
    try:
        try:
            from pypdf import PdfReader
        except ImportError:
            from PyPDF2 import PdfReader
        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        raise RuntimeError(
            "Nenhuma biblioteca para ler PDF encontrada. "
            "Instale fitz (PyMuPDF), pdfplumber ou pypdf no Python do QGIS."
        )
    except Exception:
        return ""


def extract_pdf_alias_fields(text, data):
    patterns = {
        "codigo_da_certificacao": r"CERTIFICA[ÇC][ÃA]O:\s*([0-9a-fA-F-]+)",
        "certificacao": r"CERTIFICA[ÇC][ÃA]O:\s*([^\r\n]+)",
        "data_certificacao": r"Data\s+Certifica[cç][aã]o:\s*(\d{2}/\d{2}/\d{2,4})",
        "denominacao": r"Denomina[çc][ãa]o:\s*(.+?)(?=\s+Natureza\s+da\s+[ÁA]rea:|$)",
        "codigo_incra": r"C[oó]digo\s+INCRA/SNCR:\s*([0-9.\-]+)",
        "municipio_uf": r"Munic[ií]pio(?:/UF)?:\s*([^\r\n]+)",
        # "matricula" extraída via extrair_matricula_principal() abaixo (somente cabeçalho)
        "uf": r"\bUF:\s*([A-Z]{2})\b",
        "trt": r"T\.R\.T\.:\s*([^\r\n]+)",
        "documento_rt": r"Documento\s+de\s+RT:\s*([^\r\n]+)",
        "area_ha": r"[ÁA]rea\s*\(ha\):\s*([0-9.,]+\s*(?:ha)?)",
        "area": r"[ÁA]rea\s*\(Sistema\s+Geod[eé]sico\s+Local\):\s*([0-9.,]+\s*ha)",
        "perimetro_m": r"Per[ií]metro\s*\(m\):\s*([0-9.,]+\s*(?:m)?)",
        "perimetro": r"Per[ií]metro\s*\(m\):\s*([0-9.,]+\s*m)",
        "local_data": r"([A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-Za-zÀ-ÿ\s'.-]+-[A-Z]{2},\s*\d{1,2}\s+de\s+[A-Za-zÀ-ÿçÇ]+\s+de\s+\d{4})",
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            set_value(data, key, match.group(1))

    # ── Matrícula: somente do cabeçalho (nunca de confrontantes) ─────────────
    matricula = extrair_matricula_principal(text)
    if matricula:
        data["matricula"] = matricula
        data["matricula_do_imovel"] = matricula
    # else: campo ausente → não define data["matricula"] (fica vazio)

    # ── Cartório/CNS: somente do cabeçalho → chaves privadas como fallback ───
    # Os valores serão promovidos em _collect_memorial_data se o QGIS não tiver dados.
    cc = extrair_cartorio_cns_principal(text)
    if cc["cartorio"]:
        data["__cartorio_pdf"] = cc["cartorio"]
    if cc["cns"]:
        data["__cns_pdf"] = cc["cns"]

    # Limpeza de campos que podem vir colados na mesma linha do PDF
    if data.get("denominacao"):
        data["denominacao"] = ferramenta_limpar_campo_cabecalho(
            data["denominacao"], "Natureza da Área"
        )

    if data.get("matricula"):
        data["matricula"] = ferramenta_limpar_campo_cabecalho(
            data["matricula"], "Código INCRA/SNCR"
        )

    if data.get("proprietario"):
        data["proprietario"] = ferramenta_limpar_campo_cabecalho(
            data["proprietario"], "Município/UF"
        )

    if data.get("municipio_uf"):
        for rotulo in (
            "Natureza da Área",
            "CPF",
            "CNPJ",
            "Código INCRA/SNCR",
            "Responsável Técnico",
        ):
            data["municipio_uf"] = ferramenta_limpar_campo_cabecalho(
                data["municipio_uf"], rotulo
            )

    if data.get("natureza_area"):
        for rotulo in (
            "CPF",
            "CNPJ",
            "Código INCRA/SNCR",
            "Responsável Técnico",
        ):
            data["natureza_area"] = ferramenta_limpar_campo_cabecalho(
                data["natureza_area"], rotulo
            )

    if data.get("codigo_incra"):
        data["codigo_incra"] = format_codigo_incra(data["codigo_incra"])
    if data.get("documento_rt") and not data.get("trt"):
        data["trt"] = extract_trt_code(data["documento_rt"])
    elif data.get("trt"):
        data["trt"] = extract_trt_code(data["trt"])
    if data.get("codigo_da_certificacao") and not data.get("certificacao"):
        data["certificacao"] = data["codigo_da_certificacao"]
    split_municipio_uf(data)
    ensure_location_fields(data)


def split_municipio_uf(data):
    """Preenche MUNICIPIO, COMARCA, UF e MUNICIPIO_UF de forma robusta.

    O SIGEF pode trazer a linha de várias formas, por exemplo:
    - Município/UF: Jardim-CE
    - Município/UF: Jardim - CE
    - Município/UF: Jardim / CE
    - Município: Jardim-CE

    Também pode acontecer de o leitor de PDF juntar campos na mesma linha.
    Por isso esta função procura o padrão CIDADE-UF dentro do texto, em vez de
    exigir que a string inteira seja exatamente CIDADE-UF.
    """
    raw_value = (
        data.get("municipio_uf", "")
        or data.get("municipio", "")
        or data.get("comarca", "")
        or ""
    )
    raw_value = clean_value(raw_value)
    if not raw_value:
        return

    # Normaliza separadores comuns.
    normalized = raw_value.replace("/", "-").replace("–", "-").replace("—", "-")
    normalized = re.sub(r"\s*-\s*", "-", normalized)

    # Remove trechos que às vezes vêm colados depois do Município/UF.
    normalized = re.split(
        r"\b(?:Respons[aá]vel|Natureza|CPF|CNPJ|C[oó]digo|Matr[ií]cula|Cart[oó]rio|Sistema|Documento|[ÁA]rea|Per[ií]metro)\b",
        normalized,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0].strip(" ,.;:-")

    # Captura a última UF válida no final do trecho ou antes de texto colado.
    match = re.search(r"([A-Za-zÀ-ÿ .']+?)-([A-Z]{2})(?:\b|$)", normalized, re.IGNORECASE)
    if not match:
        return

    city = clean_value(match.group(1)).upper().strip(" ,.;:-")
    uf = clean_value(match.group(2)).upper()
    if not city or not uf:
        return

    data["municipio"] = city
    data["comarca"] = city
    data["uf"] = uf
    data["municipio_uf"] = "%s-%s" % (city, uf)


def ensure_location_fields(data):
    """Garante campos de localização mesmo quando a UF veio somente no TRT."""
    if data.get("municipio") and not data.get("comarca"):
        data["comarca"] = data["municipio"]

    # Se UF não veio em Município/UF, tenta extrair de Documento de RT ou TRT.
    if not data.get("uf"):
        for key in ("documento_rt", "trt"):
            value = clean_value(data.get(key, ""))
            match = re.search(r"-\s*([A-Z]{2})\b", value, re.IGNORECASE)
            if match:
                data["uf"] = match.group(1).upper()
                break

    if data.get("municipio") and data.get("uf"):
        data["municipio_uf"] = "%s-%s" % (data["municipio"], data["uf"])
        if not data.get("comarca"):
            data["comarca"] = data["municipio"]


def finalize_data_fields(data):
    """Ajustes finais dos campos antes de substituir os marcadores no DOCX."""
    clean_matricula_field(data)
    normalize_cartorio_cns_fields(data)
    ensure_location_fields(data)
    _uf = data.get("uf", "")
    for _k in ("trt", "t_r_t"):
        if data.get(_k):
            data[_k] = formatar_trt_sem_duplicar_uf(extract_trt_code(data[_k]), _uf)
    if data.get("municipio"):
        data["municipio_data"] = formatar_municipio_para_data(data["municipio"])
    if data.get("ccir"):
        data["ccir"] = formatar_ccir(data["ccir"])
    if data.get("cib"):
        data["cib"] = formatar_cib(data["cib"])
    if data.get("car"):
        data["car"] = formatar_car(data["car"])


def clean_matricula_field(data):
    """Garante que a matrícula não venha colada com Código INCRA/SNCR ou outros rótulos."""
    value = clean_value(data.get("matricula", ""))
    if not value:
        return

    value = re.split(
        r"\b(?:C[oó]digo\s+INCRA/SNCR|C[oó]digo\s+INCRA|Natureza\s+da\s+[ÁA]rea|CPF|CNPJ|Munic[ií]pio|Respons[aá]vel|Forma[çc][ãa]o|Cart[oó]rio|Sistema|Documento|[ÁA]rea|Per[ií]metro)\b",
        value,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0]
    value = value.strip(" ,.;:-")

    # Em geral matrícula é um número/código curto. Se vier "981 ...", mantém só o primeiro token útil.
    match = re.match(r"^([A-Za-z0-9./-]+)", value)
    if match:
        value = match.group(1)

    data["matricula"] = value
    data["matricula_do_imovel"] = value


def extract_cns_value(value):
    """Extrai CNS no padrão 01.948-9 de qualquer texto."""
    value = clean_value(value)
    match = re.search(r"\b\d{2}\.\d{3}-\d\b", value)
    return match.group(0) if match else ""


# ── Extração segura do cabeçalho SIGEF ────────────────────────────────────────

def _extrair_cabecalho_pdf(texto):
    """Retorna somente a seção de cabeçalho do PDF SIGEF.

    Tudo que vem antes da tabela de vértices/confrontantes.
    """
    fim = re.search(
        r"DESCRI[ÇC][ÃA]O\s+(?:DA\s+PARCELA|DO\s+PER[ÍI]METRO)"
        r"|LIMITES\s+E\s+CONFRONTA[ÇC][ÕO]ES"
        r"|CONFRONTANTES\s+DO\s+IM[ÓO]VEL"
        r"|V[ÉE]RTICE\s+SEGMENTO\s+VANTE"
        r"|COORDENADAS\s+GEOD[EÉ]SICAS",
        texto,
        re.IGNORECASE,
    )
    if fim:
        return texto[:fim.start()]
    # Fallback: primeiro código de vértice no formato M-001-A / V-001-A
    vertice = re.search(r"\b[A-Z]-\d{3,}-[A-Z]\b", texto)
    if vertice:
        return texto[:vertice.start()]
    return texto


def extrair_matricula_principal(texto):
    """Extrai matrícula somente do cabeçalho 'Matrícula do imóvel:'.

    Retorna '' quando o campo está ausente ou vazio.
    Nunca usa dados de confrontantes.
    """
    cabecalho = _extrair_cabecalho_pdf(texto)
    # Captura tudo até o fim da linha — [ \t]* evita absorver o \n seguinte
    match = re.search(
        r"Matr[ií]cula\s+do\s+im[oó]vel:[ \t]*([^\r\n]*)",
        cabecalho,
        re.IGNORECASE,
    )
    if not match:
        return ""
    valor = match.group(1).strip()
    # Remove rótulos do próximo campo que ficaram colados na mesma linha
    valor = re.split(
        r"\b(?:Cart[oó]rio|C[oó]digo|Natureza|Respons|Data|CPF|CNPJ|"
        r"Munic[ií]pio|[ÁA]rea|Per[ií]metro)\b",
        valor,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0].strip(" ,.;:-")
    if not valor:
        return ""
    # Matrícula é um número/código curto; extrai o primeiro token alfanumérico
    m2 = re.match(r"^([A-Za-z0-9./-]+)", valor)
    return m2.group(1) if m2 else ""


def extrair_cartorio_cns_principal(texto):
    """Extrai cartório e CNS somente do cabeçalho 'Cartório (CNS):'.

    Retorna {"cartorio": str, "cns": str}.
    Ambos '' quando o cabeçalho não existir no PDF.
    Nunca usa CNS de confrontantes.
    """
    cabecalho = _extrair_cabecalho_pdf(texto)
    # Captura tudo até o fim da linha — [ \t]* evita absorver o \n seguinte
    match = re.search(
        r"Cart[oó]rio\s*\(CNS\)[ \t]*:[ \t]*([^\r\n]+)",
        cabecalho,
        re.IGNORECASE,
    )
    if not match:
        return {"cartorio": "", "cns": ""}
    valor = match.group(1).strip()
    if not valor:
        return {"cartorio": "", "cns": ""}
    # Remove rótulos colados
    valor = re.split(
        r"\b(?:Data\s+Certifica|T\.R\.T\.|Denomina|C[oó]digo|[ÁA]rea|Matr[ií]cula)\b",
        valor,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0].strip(" ,.;:-")
    if not valor:
        return {"cartorio": "", "cns": ""}
    # Extrai código CNS no formato XX.XXX-X (com ou sem parênteses)
    cns_match = re.search(r"\(?\b(\d{2}\.\d{3}-\d)\b\)?", valor)
    cns = cns_match.group(1) if cns_match else ""
    # Remove o CNS do nome do cartório
    nome_cartorio = re.sub(r"\s*\(?\b%s\b\)?\s*" % re.escape(cns), " ", valor).strip(" ,.;:-()")\
        if cns else valor
    return {"cartorio": nome_cartorio, "cns": cns}


def normalize_cartorio_cns_fields(data):
    """Formata CNS com parênteses e evita CNS duplicado no nome do cartório."""
    cartorio = clean_value(data.get("cartorio", ""))
    cns_raw = clean_value(data.get("cns", ""))

    cns = extract_cns_value(cns_raw) or extract_cns_value(cartorio)

    if cartorio and cns:
        # Remove CNS cru ou com parênteses de dentro do nome do cartório para não duplicar.
        cartorio = re.sub(r"\(?\b%s\b\)?" % re.escape(cns), "", cartorio).strip()
        cartorio = re.sub(r"\s{2,}", " ", cartorio).strip(" ,.;:-")

    if cartorio:
        data["cartorio"] = cartorio

    if cns:
        data["cns"] = "(%s)" % cns
        # Campo alternativo caso algum modelo use outro marcador.
        data["cns_sem_parenteses"] = cns


def extract_key_value_pairs(text):
    data = {}
    for line in text.splitlines():
        line = " ".join(line.split())
        if not line:
            continue
        for key, value in re.findall(r"([^:\n]{2,80}):\s*([^:]+?)(?=\s+[A-Za-zÀ-ÿ .()/]+:|$)", line):
            set_value(data, key, value)
    return data


def extract_descricao_perimetro(text):
    patterns = [
        r"(LIMITES\s+E\s+CONFRONTA[ÇC][ÕO]ES:\s*.*?)(?:\n\s*Todas as coordenadas|\n\s*Observa[çc][õo]es:|\n\s*CONFRONTANTES DO IM[ÓO]VEL|$)",
        r"(DESCRI[ÇC][ÃA]O DO PER[ÍI]METRO\s*.*?)(?:\n\s*Todas as coordenadas|\n\s*Observa[çc][õo]es:|\n\s*CONFRONTANTES DO IM[ÓO]VEL|$)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            return clean_value(match.group(1))
    return extract_descricao_perimetro_by_bounds(text)


def extract_descricao_perimetro_by_bounds(text):
    start_match = re.search(
        r"(LIMITES\s+E\s+CONFRONTA\S+ES:\s*|DESCRI\S+O\s+DO\s+PER\S+METRO\s*)",
        text,
        re.IGNORECASE,
    )
    if not start_match:
        return ""

    start = start_match.start()
    end_match = re.search(
        r"(?:Todas\s+as\s+coordenadas|Observa\S+es:|CONFRONTANTES\s+DO\s+IM\S+VEL|DECLARA\S+O)",
        text[start_match.end():],
        re.IGNORECASE,
    )
    end = start_match.end() + end_match.start() if end_match else len(text)
    return clean_value(text[start:end])


def extract_pdf_segments(text):
    # Primeiro tenta usar o parser robusto baseado no script antigo do QGIS.
    # Ele lê linha por linha e valida código, longitude, latitude, altitude,
    # vante, azimute, distância e confrontação, evitando cortes como "MARIA" -> "ARIA".
    segments = extract_pdf_segments_with_ferramenta_parser(text)
    if segments:
        return segments

    segments = extract_pdf_segments_from_narrative(text)
    if segments:
        return segments

    return extract_pdf_segments_from_table_like_text(text)


def extract_pdf_segments_with_ferramenta_parser(text):
    segments = []
    seen = set()

    for line in text.splitlines():
        line = " ".join(line.strip().split())
        if not line:
            continue

        try:
            if ferramenta_ignorar_linha(line):
                continue
            record = ferramenta_parse_linha_vertice(line)
        except NameError:
            # Caso a seção Ferramentas seja removida futuramente, cai no parser antigo.
            record = None

        if not record:
            continue

        segment = {
            "codigo": clean_value(record.get("codigo", "")).upper(),
            "longitude": clean_value(record.get("longitude", "")),
            "latitude": clean_value(record.get("latitude", "")),
            "altitude": clean_value(record.get("altitude", "")),
            "vante": clean_value(record.get("vante", "")).upper(),
            "azimute": clean_value(record.get("azimute", "")),
            "dist_m": clean_value(record.get("dist", "")),
            "confrontacao": normalize_confrontante(record.get("confrontacao", "")),
        }

        key = (
            segment["codigo"],
            segment["longitude"],
            segment["latitude"],
            segment["altitude"],
            segment["vante"],
            segment["azimute"],
            segment["dist_m"],
            segment["confrontacao"],
        )
        if key in seen:
            continue
        seen.add(key)
        segments.append(segment)

    return segments


def extract_pdf_segments_from_narrative(text):
    segment_pattern = re.compile(
        r"(?:deste\s+segue\s+)?confrontando\s+com\s+(?:a\s+propriedade\s+de\s+)?(?P<confrontacao>.*?),\s*"
        r"com\s+os\s+seguintes\s+azimutes\s+e\s+dist\S?ncias:\s*"
        r"(?P<azimute>[^;]+?)\s+e\s+(?P<distancia>[0-9.,]+)\s*m\s+"
        r"at\S?\s+o\s+v\S?rtice\s+(?P<vante>[A-Z0-9\-]+)"
        r"(?:,\s*de\s+coordenadas\s*\((?P<coords>.*?)\))?"
        r"(?P<after_coords>[^;]*)?",
        re.IGNORECASE | re.DOTALL,
    )
    segments = []
    current_vertex = extract_first_vertex(text)
    current_coords = extract_first_vertex_coordinates(text)

    for match in segment_pattern.finditer(text):
        target_coords = parse_lon_lat(match.group("coords") or "")
        segment = {
            "codigo": current_vertex,
            "longitude": current_coords.get("longitude", ""),
            "latitude": current_coords.get("latitude", ""),
            "altitude": current_coords.get("altitude", ""),
            "vante": clean_value(match.group("vante")).upper(),
            "azimute": clean_value(match.group("azimute")),
            "dist_m": clean_value(match.group("distancia")),
            "confrontacao": normalize_confrontante(match.group("confrontacao")),
            "vante_longitude": target_coords.get("longitude", ""),
            "vante_latitude": target_coords.get("latitude", ""),
            "vante_altitude": extract_altitude(match.group("after_coords") or ""),
        }
        segments.append(segment)
        current_vertex = segment["vante"]
        current_coords = {
            "longitude": segment["vante_longitude"],
            "latitude": segment["vante_latitude"],
            "altitude": segment["vante_altitude"],
        }
    return segments


def extract_pdf_segments_from_table_like_text(text):
    segments = extract_pdf_segments_from_flat_table(text)
    if segments:
        return segments

    line_patterns = [
        re.compile(
            r"(?P<codigo>[A-Z0-9]+-[MPV]-\d+).*?"
            r"(?P<longitude>-?\d{1,3}[^\d\s,.'-]+\d{1,2}'\d{1,2},\d+\"?).*?"
            r"(?P<latitude>-?\d{1,3}[^\d\s,.'-]+\d{1,2}'\d{1,2},\d+\"?)\s+"
            r"(?P<altitude>-?\d{1,5}[.,]\d+)\s+"
            r"(?P<vante>[A-Z0-9]+-[MPV]-\d+).*?"
            r"(?P<azimute>\d{1,3}[^\d\s,.'-]+\d{1,2}'[^ ]*).*?"
            r"(?P<distancia>\d+[.,]\d+)(?:\s*m\b)?\s*.*?"
            r"(?P<confrontacao>[A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9 .:/|_-]{4,})$",
            re.IGNORECASE,
        ),
        re.compile(
            r"(?P<codigo>[A-Z0-9]+-[MPV]-\d+).*?"
            r"(?P<longitude>-?\d{1,3}[^\d\s,.'-]+\d{1,2}'\d{1,2},\d+\"?).*?"
            r"(?P<latitude>-?\d{1,3}[^\d\s,.'-]+\d{1,2}'\d{1,2},\d+\"?).*?"
            r"(?P<vante>[A-Z0-9]+-[MPV]-\d+).*?"
            r"(?P<azimute>\d{1,3}[^\d\s,.'-]+\d{1,2}'[^ ]*).*?"
            r"(?P<distancia>\d+[.,]\d+)(?:\s*m\b)?\s*.*?"
            r"(?P<confrontacao>[A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9 .:/|_-]{4,})$",
            re.IGNORECASE,
        ),
    ]
    segments = []
    for line in text.splitlines():
        line = " ".join(line.split())
        match = None
        for line_pattern in line_patterns:
            match = line_pattern.search(line)
            if match:
                break
        if not match:
            continue
        segments.append({
            "codigo": clean_value(match.group("codigo")).upper(),
            "longitude": clean_value(match.group("longitude")),
            "latitude": clean_value(match.group("latitude")),
            "altitude": clean_value(match.groupdict().get("altitude", "")),
            "vante": clean_value(match.group("vante")).upper(),
            "azimute": clean_value(match.group("azimute")),
            "dist_m": clean_value(match.group("distancia")),
            "confrontacao": normalize_confrontante(match.group("confrontacao")),
        })
    return segments


def extract_pdf_segments_from_flat_table(text):
    flat_text = " ".join(text.split())
    degree = r"[^\d\s,.'-]+"
    coordinate = r"-?\d{1,3}%s\d{1,2}'\d{1,2},\d+\"?" % degree
    azimuth = r"\d{1,3}%s\d{1,2}'\d{1,2}(?:[.,]\d+)?\"?" % degree
    row_pattern = re.compile(
        r"(?P<codigo>[A-Z0-9]+-[MPV]-\d+)\s+"
        r"(?P<longitude>%s)\s+"
        r"(?P<latitude>%s)\s+"
        r"(?P<altitude>-?\d{1,5}[.,]\d+)\s+"
        r"(?P<vante>[A-Z0-9]+-[MPV]-\d+)\s+"
        r"(?P<azimute>%s)\s+"
        r"(?P<distancia>\d+[.,]\d+)\s*"
        r"(?P<confrontacao>.*?)"
        r"(?=\s+[A-Z0-9]+-[MPV]-\d+\s+%s|$)"
        % (coordinate, coordinate, azimuth, coordinate),
        re.IGNORECASE,
    )

    segments = []
    for match in row_pattern.finditer(flat_text):
        confrontacao = normalize_confrontante(match.group("confrontacao"))
        if not confrontacao:
            continue
        segments.append({
            "codigo": clean_value(match.group("codigo")).upper(),
            "longitude": clean_value(match.group("longitude")),
            "latitude": clean_value(match.group("latitude")),
            "altitude": clean_value(match.group("altitude")),
            "vante": clean_value(match.group("vante")).upper(),
            "azimute": clean_value(match.group("azimute")),
            "dist_m": clean_value(match.group("distancia")),
            "confrontacao": confrontacao,
        })
    return segments


def extract_first_vertex(text):
    match = re.search(r"v\S?rtice\s+([A-Z0-9\-]+)", text, re.IGNORECASE)
    return clean_value(match.group(1)).upper() if match else ""


def extract_first_vertex_coordinates(text):
    match = re.search(
        r"v\S?rtice\s+[A-Z0-9\-]+.*?coordenadas\s*\((?P<coords>.*?)\)(?P<after_coords>[^;]*)?",
        text,
        re.IGNORECASE | re.DOTALL,
    )
    if not match:
        return {}
    result = parse_lon_lat(match.group("coords"))
    result["altitude"] = extract_altitude(match.group("after_coords") or "")
    return result


def parse_lon_lat(text):
    result = {}
    lon = re.search(r"Longitude:\s*([^,;]+(?:,\d+\"?)?)", text, re.IGNORECASE)
    lat = re.search(r"Latitude:\s*([^,;]+(?:,\d+\"?)?)", text, re.IGNORECASE)
    if lon:
        result["longitude"] = clean_value(lon.group(1))
    if lat:
        result["latitude"] = clean_value(lat.group(1))
    return result


def extract_altitude(text):
    match = re.search(r"altitude\s*(?:de\s*)?(-?[0-9.,]+)\s*m?", str(text), re.IGNORECASE)
    return clean_value(match.group(1)) if match else ""


_PALAVRAS_STOPWORDS_COMP = frozenset({"DE", "DA", "DO", "DOS", "DAS", "E", "EM", "A", "O", "AS", "OS"})

_RE_ESPECIAL_PUBLICO = re.compile(
    r"\b(?:"
    r"RODOVIA|ESTRADA|RUA|AVENIDA"
    r"|AREA\s+DE\s+SERVIDAO|AREA\s+PUBLICA"
    r"|RIO|RIACHO|CORREGO|LAGOA|ACUDE|CANAL|LAGO"
    r"|LINHA\s+FERREA|FERROVIA"
    r"|LINHA\s+DE\s+TRANSMISSAO|LINHA\s+DE\s+DISTRIBUICAO"
    r"|REDE\s+ELETRICA"
    r"|AREA\s+DE\s+MARINHA|TERRENOS\s+DE\s+MARINHA|TERRAS\s+DE\s+MARINHA"
    r"|FAIXA\s+DE\s+DOMINIO|FAIXA\s+DE\s+PASSAGEM"
    r")\b",
    re.IGNORECASE,
)


def normalizar_nome_para_comparacao(nome):
    """Remove acentos, converte para maiúsculas, remove pontuação e colapsa espaços."""
    if not nome:
        return ""
    t = unicodedata.normalize("NFKD", str(nome).upper())
    t = "".join(c for c in t if not unicodedata.combining(c))
    t = re.sub(r"[^A-Z0-9\s]", " ", t)
    return re.sub(r"\s+", " ", t).strip()


def eh_confrontante_especial_publico(nome):
    """Retorna True se o nome parece ser infraestrutura, área pública ou recurso natural."""
    if not nome:
        return False
    norm = normalizar_nome_para_comparacao(nome)
    return bool(_RE_ESPECIAL_PUBLICO.search(norm))


def _preparar_confinantes_para_enriquecimento(confinante_layer, servico_id):
    """Busca todos os confinantes do serviço e seus respectivos dados."""
    confinantes = _buscar_confinantes_por_codigo(confinante_layer, servico_id)
    if not confinantes:
        return []

    resultado = []
    _id_field = field_name_lookup(confinante_layer).get("id")
    for feat in confinantes:
        cid = None
        if _id_field:
            try:
                raw = feat[_id_field]
                if raw is not None:
                    cid = int(float(str(raw).strip()))
            except Exception:
                pass

        proprietario = ""
        if cid is not None:
            try:
                pessoas = _buscar_pessoas_confinante(confinante_layer, cid)
                if pessoas:
                    proprietario = str(pessoas[0].get("nome", "") or "").strip().upper()
            except Exception:
                pass

        nome_prop_raw = feature_value(confinante_layer, feat, "nome_propriedade")
        nome_prop = nome_prop_raw.strip().upper() if nome_prop_raw else ""
        if nome_prop in ("NULL", "NONE", "0"):
            nome_prop = ""

        resultado.append({
            "feature":          feat,
            "nome_prop":        nome_prop,
            "proprietario":     proprietario,
            "nome_prop_norm":   normalizar_nome_para_comparacao(nome_prop),
            "proprietario_norm": normalizar_nome_para_comparacao(proprietario),
        })
    return resultado


def buscar_confinante_correspondente(confrontante_texto, confinantes_dados):
    """Encontra o confinante da camada que melhor corresponde ao texto do PDF.

    Retorna o item de confinantes_dados com maior similaridade (≥ 0.75) ou None.
    """
    conf_norm = normalizar_nome_para_comparacao(confrontante_texto)
    if not conf_norm or not confinantes_dados:
        return None

    def _palavras(s):
        return [w for w in s.split() if w not in _PALAVRAS_STOPWORDS_COMP and len(w) > 1]

    conf_palavras = set(_palavras(conf_norm))
    if not conf_palavras:
        return None

    best = None
    best_score = 0.0

    for item in confinantes_dados:
        prop_norm = item["proprietario_norm"]
        if not prop_norm:
            continue
        prop_palavras = set(_palavras(prop_norm))
        if not prop_palavras:
            continue
        common = conf_palavras & prop_palavras
        if not common:
            continue
        score = 2.0 * len(common) / (len(conf_palavras) + len(prop_palavras))
        if score > best_score and score >= 0.75:
            best_score = score
            best = item

    return best


def montar_texto_confinante_completo(confinante_layer, feature):
    """Compõe o texto completo do confinante para a descrição do perímetro.

    Formato com nome_propriedade:
        NOME_PROP DE PROPRIEDADE DE PROPRIETARIO, MATRÍCULA N° X, CNS: Y, CCIR: Z, CIB: W

    Formato sem nome_propriedade:
        PROPRIETARIO, MATRÍCULA N° X, ...

    Retorna tuple (texto, has_nome_prop).
    """
    nome_prop_raw = feature_value(confinante_layer, feature, "nome_propriedade")
    nome_prop = nome_prop_raw.strip().upper() if nome_prop_raw else ""
    if nome_prop in ("NULL", "NONE", "0"):
        nome_prop = ""

    proprietario = ""
    try:
        _id_field = field_name_lookup(confinante_layer).get("id")
        if _id_field:
            raw_id = feature[_id_field]
            if raw_id is not None:
                cid = int(float(str(raw_id).strip()))
                pessoas = _buscar_pessoas_confinante(confinante_layer, cid)
                if pessoas:
                    proprietario = str(pessoas[0].get("nome", "") or "").strip().upper()
    except Exception:
        pass

    if nome_prop and proprietario:
        nome_base = "%s DE PROPRIEDADE DE %s" % (nome_prop, proprietario)
        has_nome_prop = True
    elif nome_prop:
        nome_base = nome_prop
        has_nome_prop = True
    elif proprietario:
        nome_base = proprietario
        has_nome_prop = False
    else:
        nome_base = ""
        has_nome_prop = False
        for _campo in ("nome", "descricao"):
            _v = feature_value(confinante_layer, feature, _campo)
            if _v and _v.upper() not in ("NULL", "NONE"):
                nome_base = _v.strip().upper()
                break
        if not nome_base:
            return "", False

    partes_extras = []
    _matricula = feature_value(confinante_layer, feature, "matricula")
    if _matricula and _matricula.upper() not in ("NULL", "NONE", "0"):
        partes_extras.append("MATRÍCULA N° %s" % _matricula.strip().upper())

    _cns = feature_value(confinante_layer, feature, "cns")
    if _cns and _cns.upper() not in ("NULL", "NONE", "0"):
        partes_extras.append("CNS: %s" % formatar_cns(_cns))

    _ccir = feature_value(confinante_layer, feature, "ccir")
    if _ccir and _ccir.upper() not in ("NULL", "NONE", "0"):
        partes_extras.append("CCIR: %s" % formatar_ccir(_ccir))

    _cib_v = feature_value(confinante_layer, feature, "cib")
    if _cib_v and _cib_v.upper() not in ("NULL", "NONE", "0"):
        partes_extras.append("CIB: %s" % formatar_cib(_cib_v))

    texto = (nome_base + ", " + ", ".join(partes_extras)) if partes_extras else nome_base
    return texto, has_nome_prop


def enriquecer_confrontantes_segments_com_camada(segments, confinante_layer, servico_id):
    """Substitui o campo confrontacao de cada segmento pelo texto completo da camada.

    Para cada segmento:
    - Se o confrontante for especial/público, mantém o texto original.
    - Se houver correspondência na camada, substitui pelo texto completo
      e marca confrontacao_has_nome_prop.
    - Se não houver correspondência, mantém o texto original.
    """
    if not segments or not confinante_layer or servico_id is None:
        return segments

    confinantes_dados = _preparar_confinantes_para_enriquecimento(confinante_layer, servico_id)
    if not confinantes_dados:
        return segments

    resultado = []
    for seg in segments:
        copia = dict(seg)
        confrontacao = copia.get("confrontacao", "")

        if not confrontacao or eh_confrontante_especial_publico(confrontacao):
            copia.setdefault("confrontacao_has_nome_prop", False)
            resultado.append(copia)
            continue

        match = buscar_confinante_correspondente(confrontacao, confinantes_dados)
        if match:
            texto, has_nome_prop = montar_texto_confinante_completo(
                confinante_layer, match["feature"]
            )
            if texto:
                copia["confrontacao"] = texto
                copia["confrontacao_has_nome_prop"] = has_nome_prop
                resultado.append(copia)
                continue

        copia.setdefault("confrontacao_has_nome_prop", False)
        resultado.append(copia)

    return resultado


def build_descricao_perimetro_from_segments(cabecalho, segments):
    if not segments:
        return ""

    prepared_segments = prepare_segments_for_description(segments)
    first = prepared_segments[0]
    first_vertex = first.get("codigo", "")
    if not first_vertex:
        return ""

    first_coords = format_vertex_coordinates(first.get("longitude", ""), first.get("latitude", ""))
    meridian = meridiano_central_text(cabecalho, first)
    start_parts = [
        "LIMITES E CONFRONTAÇÕES: Inicia-se a descrição deste perímetro no vértice %s" % first_vertex
    ]
    if first_coords:
        start_parts.append(
            "georreferenciado no Sistema Geodésico Brasileiro, DATUM - SIRGAS2000, %s, "
            "de coordenadas %s%s"
            % (meridian, first_coords, format_altitude_suffix(first.get("altitude", "")))
        )

    description = ", ".join(start_parts) + ";"
    segment_texts = []
    for segment in prepared_segments:
        segment_text = build_segment_description(segment)
        if segment_text:
            segment_texts.append(segment_text)

    if segment_texts:
        description += " " + " ".join(segment_texts)
    if description.endswith(";"):
        description = description[:-1]

    description = clean_value(description)

    # Remove ponto final para permitir continuação no template
    description = description.rstrip(".")

    return description


def prepare_segments_for_description(segments):
    prepared = [dict(segment) for segment in segments]
    by_code = {
        clean_value(segment.get("codigo", "")).upper(): segment
        for segment in prepared
        if clean_value(segment.get("codigo", ""))
    }

    for index, segment in enumerate(prepared):
        target = by_code.get(clean_value(segment.get("vante", "")).upper())
        if target is None and prepared:
            target = prepared[(index + 1) % len(prepared)]

        if target:
            segment.setdefault("vante_longitude", target.get("longitude", ""))
            segment.setdefault("vante_latitude", target.get("latitude", ""))
            segment.setdefault("vante_altitude", target.get("altitude", ""))
            if not segment.get("vante_longitude"):
                segment["vante_longitude"] = target.get("longitude", "")
            if not segment.get("vante_latitude"):
                segment["vante_latitude"] = target.get("latitude", "")
            if not segment.get("vante_altitude"):
                segment["vante_altitude"] = target.get("altitude", "")
    return prepared


def build_segment_description(segment):
    vante = clean_value(segment.get("vante", "")).upper()
    if not vante:
        return ""

    confrontacao = normalize_confrontante(segment.get("confrontacao", ""))
    if confrontacao:
        if segment.get("confrontacao_has_nome_prop"):
            # Começa com nome_propriedade — não adicionar "a propriedade de"
            text = "deste segue confrontando com %s" % confrontacao
        else:
            text = "deste segue confrontando com a propriedade de %s" % confrontacao
    else:
        text = "deste segue"

    azimute = clean_value(segment.get("azimute", ""))
    distancia = format_distance(segment.get("dist_m", ""))
    if azimute and distancia:
        text += ", com os seguintes azimutes e distâncias: %s e %s" % (azimute, distancia)
    elif distancia:
        text += ", com distância de %s" % distancia
    elif azimute:
        text += ", com azimute de %s" % azimute

    text += " até o vértice %s" % vante
    coords = format_vertex_coordinates(segment.get("vante_longitude", ""), segment.get("vante_latitude", ""))
    if coords:
        text += ", de coordenadas %s%s" % (coords, format_altitude_suffix(segment.get("vante_altitude", "")))
    return text + ";"


def format_vertex_coordinates(longitude, latitude):
    longitude = clean_value(longitude)
    latitude = clean_value(latitude)
    if longitude and latitude:
        return "(Longitude:%s, Latitude:%s)" % (longitude, latitude)
    return ""


def format_altitude_suffix(value):
    altitude = format_measure(value, "m")
    return " e altitude %s" % altitude if altitude else ""


def format_distance(value):
    return format_measure(value, "m")


def format_measure(value, unit):
    value = clean_value(value)
    if not value:
        return ""
    if re.search(r"%s\b" % re.escape(unit), value, re.IGNORECASE):
        return value
    return "%s%s" % (value, unit)


def meridiano_central_text(cabecalho, segment):
    for key in ("meridiano_central", "mc"):
        value = clean_value(cabecalho.get(key, ""))
        if value:
            return value

    longitude = clean_value(segment.get("longitude", ""))
    match = re.search(r"(-?\d+)", longitude)
    if match:
        degrees = int(match.group(1))
        suffix = "W" if degrees < 0 else "E"
        return "MC-%s°%s" % (abs(degrees), suffix)
    return "MC-39°W"


def normalize_confrontante(value):
    value = clean_value(value)
    if "|" in value:
        value = value.split("|")[-1]
    value = re.split(
        r"\b(?:TODAS AS COORDENADAS|OBSERVAÇÕES|OBSERVACOES|CONFRONTANTES DO IMÓVEL|CONFRONTANTES DO IMOVEL|DECLARAÇÃO|DECLARACAO|CACTTUS HUB)\b",
        value,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0]
    value = re.sub(r"\s+", " ", value)
    value = re.sub(r"^(A\s+PROPRIEDADE\s+DE|PROPRIEDADE\s+DE)\s+", "", value, flags=re.IGNORECASE)
    return value.strip(" ,.;:-").upper()


def extract_trt_code(value):
    match = re.search(r"\b(BR\d+)\b", str(value), re.IGNORECASE)
    return match.group(1).upper() if match else clean_value(value)


def format_codigo_incra(value):
    digits = only_digits(value)
    if len(digits) == 13:
        return "%s.%s.%s.%s-%s" % (digits[0:3], digits[3:6], digits[6:9], digits[9:12], digits[12])
    return clean_value(value)


def format_cpf(value):
    digits = only_digits(value)
    if len(digits) == 11:
        return "%s.%s.%s-%s" % (digits[0:3], digits[3:6], digits[6:9], digits[9:11])
    return clean_value(value)


def format_cnpj(value):
    digits = only_digits(value)
    if len(digits) == 14:
        return "%s.%s.%s/%s-%s" % (digits[0:2], digits[2:5], digits[5:8], digits[8:12], digits[12:14])
    return clean_value(value)


def only_digits(value):
    return re.sub(r"\D+", "", str(value))


def is_valid_field_value(value):
    value = clean_value(value)
    if not value:
        return False
    if value.upper() == "NULL":
        return False
    if only_digits(value) == "0":
        return False
    return value != "0"


def add_aliases(data):
    aliases = {
        "certificacao": "codigo_da_certificacao",
        "codigo_da_certificacao": "certificacao",
        "codigo_incra_sncr": "codigo_incra",
        "matricula_do_imovel": "matricula",
        "area_ha": "area",
        "area": "area_ha",
        "perimetro_m": "perimetro",
        "perimetro": "perimetro_m",
        "t_r_t": "trt",
        "trt": "t_r_t",
    }
    for source, target in aliases.items():
        if data.get(source) and not data.get(target):
            data[target] = data[source]

    ensure_location_fields(data)


def merge_data(target, source, overwrite):
    for key, value in source.items():
        if key.startswith("__"):
            target[key] = value
        elif overwrite or key not in target or not target[key]:
            target[key] = value


def field_name_lookup(layer):
    return {normalize_key(field.name()): field.name() for field in layer.fields()}


def feature_value(layer, feature, field_key):
    field_name = field_name_lookup(layer).get(normalize_key(field_key))
    if not field_name:
        return ""
    return clean_value(feature[field_name])


def qgis_feature_to_data(layer, feature, pdf_segments=None, confinante_layer=None):
    data = {}
    for source, target in (
        ("nome_cartorio",      "cartorio"),
        ("cns",                "cns"),
        ("data_matricula",     "data_matricula"),
        ("cib",                "cib"),
        ("car",                "car"),
        ("ccir",               "ccir"),
        ("qrcode_pasta_drive", "qr_code"),
    ):
        value = feature_value(layer, feature, source)
        if is_valid_field_value(value):
            data[target] = value

    # Expõe o nome do cartório também sob a chave original para o placeholder
    # {{nome_cartorio}} usado no Laudo Técnico.
    nome_cartorio = data.get("cartorio", "")
    data["nome_cartorio"] = nome_cartorio  # "" quando não cadastrado na camada

    # Armazena o id do serviço para busca direta de proprietários no Laudo Técnico
    id_field = field_name_lookup(layer).get("id")
    if id_field:
        try:
            sid = feature[id_field]
            if sid:
                data["__servico_id"] = sid
        except Exception:
            pass

    owners = get_proprietarios_do_servico(layer, feature)
    for owner in owners:
        suffix = owner["suffix"]
        data["proprietario%s" % suffix] = owner["nome"]
        if owner.get("cpf"):
            data["cpf%s" % suffix] = owner["cpf"]
        if owner.get("cnpj"):
            data["cnpj%s" % suffix] = owner["cnpj"]

    if owners and not data.get("proprietario"):
        data["proprietario"] = owners[0]["nome"]
        if owners[0].get("cpf"):
            data["cpf"] = owners[0]["cpf"]
        if owners[0].get("cnpj"):
            data["cnpj"] = owners[0]["cnpj"]

    data["bloco_proprietarios"] = build_owner_block(owners)
    data["bloco_assinaturas_proprietarios"] = build_owner_signature_block(owners)
    data["__owners"] = owners

    servico_id = data.get("__servico_id")

    try:
        from qgis.core import QgsMessageLog, Qgis as _Qi
        def _tlog(msg):
            QgsMessageLog.logMessage(msg, "GeoDocsSIGEF", _Qi.Info)
    except Exception:
        def _tlog(msg):
            pass

    _tlog("[Trace] qgis_feature_to_data: servico_id=%s | confinante_layer=%s | pdf_segments=%d"
          % (servico_id,
             confinante_layer.name() if confinante_layer else "None",
             len(pdf_segments or [])))

    confrontantes = confrontantes_por_lado(
        feature.geometry(), pdf_segments or [],
        confinante_layer=confinante_layer, servico_id=servico_id,
    )

    _tlog("[Trace] qgis_feature_to_data: confrontantes_por_lado retornou fonte=%s metodo=%s"
          % (confrontantes.get("__confrontantes_fonte", "?"),
             confrontantes.get("__metodo_confrontantes", "?")))

    data.update(confrontantes)
    return data


def ids_match(feature_value_raw, typed_value):
    feature_value_clean = clean_value(feature_value_raw)
    typed_value_clean = clean_value(typed_value)
    if feature_value_clean == typed_value_clean:
        return True
    return only_digits(feature_value_clean) and only_digits(feature_value_clean) == only_digits(typed_value_clean)


def normalizar_id_relacao(valor):
    """Normaliza um valor de ID para comparação entre campos de tipos diferentes.

    Converte inteiros, decimais e strings para uma representação canônica sem
    casas decimais desnecessárias (ex.: 7123.0 → "7123").
    """
    if valor is None:
        return ""
    txt = str(valor).strip().replace(",", ".")
    if not txt or txt.upper() in ("NULL", "NONE"):
        return ""
    try:
        num = float(txt)
        if num.is_integer():
            return str(int(num))
    except Exception:
        pass
    if txt.endswith(".0"):
        txt = txt[:-2]
    return txt.strip()


def _buscar_confinantes_por_codigo(confinante_layer, servico_id):
    """Retorna lista de features da camada Confinantes Principal onde campo 'codigo' == servico_id."""
    if not confinante_layer or servico_id is None:
        return []
    try:
        from qgis.core import QgsMessageLog, Qgis as _Qi
        def _log(msg):
            QgsMessageLog.logMessage(msg, "GeoDocsSIGEF", _Qi.Info)
    except Exception:
        def _log(msg):
            pass

    try:
        lookup = field_name_lookup(confinante_layer)
        codigo_field = lookup.get("codigo")
        if not codigo_field:
            _log("[Confrontantes] _buscar: campo 'codigo' não encontrado na camada Confinantes Principal")
            return []

        sid_norm = normalizar_id_relacao(servico_id)
        _log("[Confrontantes] servico_id bruto        = %r" % servico_id)
        _log("[Confrontantes] servico_id normalizado   = %r" % sid_norm)

        result = []
        for feat in confinante_layer.getFeatures():
            try:
                raw_cod = feat[codigo_field]
                cod_norm = normalizar_id_relacao(raw_cod)
                _log("[Confrontantes] codigo bruto=%r  normalizado=%r  match=%s"
                     % (raw_cod, cod_norm, cod_norm == sid_norm))
                if cod_norm == sid_norm:
                    result.append(feat)
            except Exception:
                pass
        return result
    except Exception:
        return []


def _ler_nome_confinante_camada(layer, feature):
    """Lê o nome do confinante: prioridade nome_propriedade → nome → descricao."""
    for campo in ("nome_propriedade", "nome", "descricao"):
        v = feature_value(layer, feature, campo)
        if v and v.upper() not in ("NULL", "NONE"):
            return normalize_confrontante(v)
    return ""


def _ler_nome_confinante_com_proprietario(confinante_layer, feature):
    """Compõe o nome completo do confinante para bloco por lado (NORTE/LESTE/SUL/OESTE).

    Formato com nome_propriedade:
        NOME_PROPRIEDADE DE PROPRIEDADE DE PROPRIETARIO, MATRÍCULA N° X, CNS: Y, CCIR: Z, CIB: W
    Formato sem nome_propriedade:
        PROPRIETARIO, MATRÍCULA N° X, ...
    Confrontante especial (rod., rio, etc.): mantém o nome como está, sem "DE PROPRIEDADE DE".

    Retorna tuple (nome_propriedade, proprietario, nome_final).
    """
    # 1. nome_propriedade
    nome_prop_raw = feature_value(confinante_layer, feature, "nome_propriedade")
    nome_prop = nome_prop_raw.strip().upper() if nome_prop_raw else ""
    if nome_prop in ("NULL", "NONE", "0"):
        nome_prop = ""

    # 2. proprietário vinculado pelo campo id da feature
    proprietario = ""
    try:
        _id_field = field_name_lookup(confinante_layer).get("id")
        if _id_field:
            raw_id = feature[_id_field]
            if raw_id is not None:
                cid = int(float(str(raw_id).strip()))
                pessoas = _buscar_pessoas_confinante(confinante_layer, cid)
                if pessoas:
                    proprietario = str(pessoas[0].get("nome", "") or "").strip().upper()
    except Exception:
        pass

    # 3. compor nome_base
    if nome_prop and proprietario:
        if eh_confrontante_especial_publico(nome_prop):
            nome_base = nome_prop
        else:
            nome_base = "%s DE PROPRIEDADE DE %s" % (nome_prop, proprietario)
    elif nome_prop:
        nome_base = nome_prop
    elif proprietario:
        nome_base = proprietario
    else:
        nome_base = ""
        for _campo in ("nome", "descricao"):
            _v = feature_value(confinante_layer, feature, _campo)
            if _v and _v.upper() not in ("NULL", "NONE"):
                nome_base = normalize_confrontante(_v)
                break
        return nome_prop, proprietario, nome_base

    # 4. acrescentar campos opcionais formatados
    partes_extras = []
    _matricula = feature_value(confinante_layer, feature, "matricula")
    if _matricula and _matricula.upper() not in ("NULL", "NONE", "0"):
        partes_extras.append("MATRÍCULA N° %s" % _matricula.strip().upper())

    _cns = feature_value(confinante_layer, feature, "cns")
    if _cns and _cns.upper() not in ("NULL", "NONE", "0"):
        partes_extras.append("CNS: %s" % formatar_cns(_cns))

    _ccir = feature_value(confinante_layer, feature, "ccir")
    if _ccir and _ccir.upper() not in ("NULL", "NONE", "0"):
        partes_extras.append("CCIR: %s" % formatar_ccir(_ccir))

    _cib = feature_value(confinante_layer, feature, "cib")
    if _cib and _cib.upper() not in ("NULL", "NONE", "0"):
        partes_extras.append("CIB: %s" % formatar_cib(_cib))

    nome_final = (nome_base + ", " + ", ".join(partes_extras)) if partes_extras else nome_base
    return nome_prop, proprietario, nome_final


def _angulo_para_lado(angulo):
    """Converte ângulo em graus (atan2, eixo-x = leste) para lado cardinal.

    Setores de 90° centrados nos eixos:
        LESTE:  -45 < angulo <=  45
        NORTE:   45 < angulo <= 135
        OESTE:  135 < angulo <= 180  ou  -180 <= angulo <= -135
        SUL:   -135 < angulo <= -45
    """
    if -45.0 < angulo <= 45.0:
        return "leste"
    if 45.0 < angulo <= 135.0:
        return "norte"
    if angulo > 135.0 or angulo <= -135.0:
        return "oeste"
    return "sul"


def _dist_ponto_segmento(px, py, ax, ay, bx, by):
    """Distância mínima do ponto (px, py) ao segmento (ax,ay)-(bx,by).

    Projeta o ponto na reta do segmento; se a projeção cair fora do segmento,
    usa a distância ao vértice mais próximo.
    """
    import math as _m
    dx, dy = bx - ax, by - ay
    len_sq = dx * dx + dy * dy
    if len_sq == 0.0:
        return _m.sqrt((px - ax) ** 2 + (py - ay) ** 2)
    t = max(0.0, min(1.0, ((px - ax) * dx + (py - ay) * dy) / len_sq))
    return _m.sqrt((px - ax - t * dx) ** 2 + (py - ay - t * dy) ** 2)


def _extrair_segmentos_classificados(imovel_geometry):
    """Extrai os segmentos do anel externo do polígono e classifica cada um
    pelo lado da caixa envolvente (bbox) com base no ponto médio do segmento.

    Retorna lista de dicts:
        {ax, ay, bx, by, mid_x, mid_y, lado}
    Retorna [] se não for possível extrair.
    """
    try:
        points = exterior_ring_points(imovel_geometry)
        if len(points) < 2:
            return []
        xs = [p.x() for p in points]
        ys = [p.y() for p in points]
        min_x, max_x = min(xs), max(xs)
        min_y, max_y = min(ys), max(ys)
        segmentos = []
        n = len(points)
        for i in range(n):
            ax, ay = points[i].x(), points[i].y()
            bx, by = points[(i + 1) % n].x(), points[(i + 1) % n].y()
            mid_x = (ax + bx) / 2.0
            mid_y = (ay + by) / 2.0
            lado = classify_by_bbox(mid_x, mid_y, min_x, max_x, min_y, max_y)
            segmentos.append(
                {"ax": ax, "ay": ay, "bx": bx, "by": by,
                 "mid_x": mid_x, "mid_y": mid_y, "lado": lado}
            )
        return segmentos
    except Exception:
        return []


def _classificar_confinantes_pela_camada(imovel_geometry, confinante_layer, servico_id):
    """Classifica confrontantes pelos pontos da camada Confinantes Principal.

    Algoritmo principal (segmentos do perímetro):
    1. Extrai os segmentos do anel externo do polígono do imóvel.
    2. Classifica cada segmento pelo lado da caixa envolvente (bbox)
       usando o ponto médio do segmento.
    3. Para cada ponto confinante (codigo == servico_id), calcula a distância
       ao segmento mais próximo e atribui o lado daquele segmento.
    4. Agrupa nomes por lado sem repetições, mantendo ordem de leitura.

    Fallback (atan2):
    - Usado apenas se não for possível extrair os segmentos do imóvel.
    - Calcula o ângulo do ponto confinante em relação ao centroide do imóvel.

    Retorna dict confrontante_norte/leste/sul/oeste, ou None se nenhum
    ponto encontrado para este servico_id.
    Emite mensagens no painel 'GeoDocsSIGEF' do Log de Mensagens do QGIS.
    """
    import math as _math

    try:
        from qgis.core import QgsMessageLog, Qgis as _Qgis
        def _log(msg):
            QgsMessageLog.logMessage(msg, "GeoDocsSIGEF", _Qgis.Info)
    except Exception:
        def _log(msg):
            pass

    _log("[Confrontantes] ── início ──────────────────────────────────────")
    _log("[Confrontantes] Serviço ID: %s" % servico_id)

    confinantes = _buscar_confinantes_por_codigo(confinante_layer, servico_id)
    _log("[Confrontantes] Pontos encontrados na camada: %d" % len(confinantes))

    if not confinantes:
        _log("[Confrontantes] Nenhum ponto — usando fallback PDF.")
        return None

    # ── Extrai segmentos do perímetro ─────────────────────────────────────────
    segmentos = _extrair_segmentos_classificados(imovel_geometry)
    if segmentos:
        _log("[Confrontantes] Segmentos extraídos do perímetro: %d" % len(segmentos))
        _log("[Confrontantes] Método: segmento mais próximo (bbox)")
    else:
        # Fallback atan2: obtém centroide do imóvel
        _log("[Confrontantes] Segmentos indisponíveis — fallback atan2.")
        try:
            centroid_geom = imovel_geometry.centroid()
            if centroid_geom and not centroid_geom.isEmpty():
                _ref = centroid_geom.asPoint()
            else:
                _ref = imovel_geometry.pointOnSurface().asPoint()
            _rx, _ry = _ref.x(), _ref.y()
            _log("[Confrontantes] Referência atan2: (%.6f, %.6f)" % (_rx, _ry))
        except Exception:
            _log("[Confrontantes] Erro ao obter referência atan2 — fallback PDF.")
            return None

    grouped = {"norte": [], "leste": [], "sul": [], "oeste": []}

    for feat in confinantes:
        nome = _ler_nome_confinante_camada(confinante_layer, feat)
        geom = feat.geometry()
        if geom is None or geom.isEmpty():
            _log("[Confrontantes] Feature sem geometria — ignorada.")
            continue

        try:
            pt = geom.centroid().asPoint()
            px, py = pt.x(), pt.y()
        except Exception:
            _log("[Confrontantes] Centroide do confinante falhou para '%s' — ignorado." % nome)
            continue

        if segmentos:
            # ── Método principal: segmento mais próximo ───────────────────────
            min_dist = None
            lado = None
            melhor_seg = None
            for seg in segmentos:
                d = _dist_ponto_segmento(px, py, seg["ax"], seg["ay"], seg["bx"], seg["by"])
                if min_dist is None or d < min_dist:
                    min_dist = d
                    lado = seg["lado"]
                    melhor_seg = seg
            _log(
                "[Confrontantes] nome=%-40s | pt=(%.4f, %.4f)"
                " | seg_mid=(%.4f, %.4f) | dist=%.6f | lado=%s"
                % (nome, px, py,
                   melhor_seg["mid_x"], melhor_seg["mid_y"],
                   min_dist, lado.upper())
            )
        else:
            # ── Fallback atan2 ────────────────────────────────────────────────
            dx = px - _rx
            dy = py - _ry
            angulo = _math.degrees(_math.atan2(dy, dx))
            lado = _angulo_para_lado(angulo)
            _log(
                "[Confrontantes] nome=%-40s | atan2 | dx=%+.4f | dy=%+.4f"
                " | angulo=%+.2f° | lado=%s"
                % (nome, dx, dy, angulo, lado.upper())
            )

        if nome and nome not in grouped[lado]:
            grouped[lado].append(nome)

    norte = ", ".join(grouped["norte"])
    leste = ", ".join(grouped["leste"])
    sul = ", ".join(grouped["sul"])
    oeste = ", ".join(grouped["oeste"])

    result = {
        # Placeholders confrontante_* e confrontacao_* populados explicitamente
        # para garantir que ambas as variantes cheguem ao DOCX.
        "confrontante_norte":   norte,
        "confrontante_leste":   leste,
        "confrontante_sul":     sul,
        "confrontante_oeste":   oeste,
        "confrontacao_norte":   norte,
        "confrontacao_leste":   leste,
        "confrontacao_sul":     sul,
        "confrontacao_oeste":   oeste,
        "__confrontantes_fonte":  "camada_confinantes",
        "__metodo_confrontantes": "camada_segmento_mais_proximo",
    }

    _log("[Confrontantes] ── resultado ─────────────────────────────────────")
    _log("[Confrontantes] AO NORTE : %s" % (norte or "(vazio)"))
    _log("[Confrontantes] AO LESTE : %s" % (leste or "(vazio)"))
    _log("[Confrontantes] AO SUL   : %s" % (sul   or "(vazio)"))
    _log("[Confrontantes] AO OESTE : %s" % (oeste or "(vazio)"))
    _log("[Confrontantes] Método   : camada_segmento_mais_proximo")
    _log("[Confrontantes] ────────────────────────────────────────────────────")

    return result


def resolver_confrontantes_por_lado_unico(
    layer_servico, feature_servico, confinante_layer, pdf_segments
):
    """Função central e única para classificar confrontantes por lado (NORTE/LESTE/SUL/OESTE).

    Prioridade 1: camada Confinantes Principal (codigo == servico_id, segmento mais próximo).
    Prioridade 2: fallback PDF (bbox → geometria QGIS → azimute).

    Nunca mistura os dois métodos. Retorna dict completo com confrontante_*,
    confrontacao_*, __confrontantes_fonte, __metodo_confrontantes,
    __aviso_confrontantes e __debug_confrontantes.
    """
    _msgs = []

    def _d(msg):
        _msgs.append(msg)
        try:
            from qgis.core import QgsMessageLog, Qgis as _Qi
            QgsMessageLog.logMessage(msg, "GeoDocsSIGEF", _Qi.Info)
        except Exception:
            pass

    def _build(norte, leste, sul, oeste, fonte, metodo, aviso):
        return {
            "confrontante_norte":    norte,
            "confrontante_leste":    leste,
            "confrontante_sul":      sul,
            "confrontante_oeste":    oeste,
            "confrontacao_norte":    norte,
            "confrontacao_leste":    leste,
            "confrontacao_sul":      sul,
            "confrontacao_oeste":    oeste,
            "__confrontantes_fonte":  fonte,
            "__metodo_confrontantes": metodo,
            "__aviso_confrontantes":  aviso,
            "__debug_confrontantes":  "\n".join(_msgs),
        }

    _d("[Confrontantes] ══ INÍCIO resolver_confrontantes_por_lado_unico ══")

    # ── Extrair servico_id ────────────────────────────────────────────────────
    servico_id = None
    if layer_servico and feature_servico:
        _id_field = field_name_lookup(layer_servico).get("id")
        if _id_field:
            try:
                _sid = feature_servico[_id_field]
                if _sid is not None:
                    servico_id = _sid
            except Exception:
                pass
    _d("[Confrontantes] servico_id       = %s" % servico_id)
    _d("[Confrontantes] confinante_layer  = %s"
       % (confinante_layer.name() if confinante_layer else "None"))
    _d("[Confrontantes] pdf_segments      = %d segmentos" % len(pdf_segments or []))

    # ── Método 1: Camada Confinantes Principal ────────────────────────────────
    if confinante_layer is not None and servico_id is not None:
        confinantes = _buscar_confinantes_por_codigo(confinante_layer, servico_id)
        _d("[Confrontantes] pontos encontrados na camada: %d" % len(confinantes))

        if confinantes:
            imovel_geom = (
                feature_servico.geometry()
                if feature_servico and not feature_servico.geometry().isEmpty()
                else None
            )
            segmentos = _extrair_segmentos_classificados(imovel_geom) if imovel_geom else []
            _d("[Confrontantes] segmentos do perímetro: %d" % len(segmentos))

            if segmentos:
                grouped = {"norte": [], "leste": [], "sul": [], "oeste": []}
                for _feat in confinantes:
                    nome_prop, proprietario, nome = _ler_nome_confinante_com_proprietario(
                        confinante_layer, _feat
                    )
                    _d("[Confrontantes] nome_propriedade      = %s" % (nome_prop or "(vazio)"))
                    _d("[Confrontantes] proprietario_vinculado = %s" % (proprietario or "(vazio)"))
                    _d("[Confrontantes] nome_final             = %s" % (nome or "(vazio)"))
                    _geom = _feat.geometry()
                    if not _geom or _geom.isEmpty():
                        _d("[Confrontantes] feature sem geometria ignorada (%s)" % nome)
                        continue
                    try:
                        _pt = _geom.centroid().asPoint()
                        px, py = _pt.x(), _pt.y()
                    except Exception:
                        _d("[Confrontantes] centroide falhou para '%s'" % nome)
                        continue
                    min_dist = None
                    lado_escolhido = None
                    melhor_seg = None
                    for _seg in segmentos:
                        _dv = _dist_ponto_segmento(
                            px, py, _seg["ax"], _seg["ay"], _seg["bx"], _seg["by"]
                        )
                        if min_dist is None or _dv < min_dist:
                            min_dist = _dv
                            lado_escolhido = _seg["lado"]
                            melhor_seg = _seg
                    _d(
                        "[Confrontantes] nome=%-38s | pt=(%.4f,%.4f) | "
                        "seg_mid=(%.4f,%.4f) | dist=%.6f | lado=%s"
                        % (
                            nome, px, py,
                            melhor_seg["mid_x"], melhor_seg["mid_y"],
                            min_dist, lado_escolhido.upper(),
                        )
                    )
                    if nome and nome not in grouped[lado_escolhido]:
                        grouped[lado_escolhido].append(nome)

                norte = ", ".join(grouped["norte"])
                leste = ", ".join(grouped["leste"])
                sul = ", ".join(grouped["sul"])
                oeste = ", ".join(grouped["oeste"])
                _d("[Confrontantes] ── resultado camada ──────────────────────")
                _d("[Confrontantes] NORTE = %s" % (norte or "(vazio)"))
                _d("[Confrontantes] LESTE = %s" % (leste or "(vazio)"))
                _d("[Confrontantes] SUL   = %s" % (sul   or "(vazio)"))
                _d("[Confrontantes] OESTE = %s" % (oeste or "(vazio)"))
                return _build(
                    norte, leste, sul, oeste,
                    "camada_confinantes",
                    "camada_segmento_mais_proximo",
                    "Confrontantes classificados pela camada Confinantes Principal "
                    "(segmento mais próximo).",
                )

            _d("[Confrontantes] nenhum segmento extraído — fallback PDF")
        else:
            _d("[Confrontantes] nenhum ponto para codigo=%s — fallback PDF" % servico_id)
    else:
        _d("[Confrontantes] camada/servico_id indisponível — fallback PDF")

    # ── Fallback PDF ──────────────────────────────────────────────────────────
    _d("[Confrontantes] ── fallback PDF ──────────────────────────────────────")
    imovel_geom = (
        feature_servico.geometry()
        if (feature_servico
            and feature_servico.geometry()
            and not feature_servico.geometry().isEmpty())
        else None
    )

    if pdf_segments:
        coord_directions = classify_pdf_segments_by_coordinates(pdf_segments)
        if coord_directions:
            _r = group_confrontantes_by_directions(pdf_segments, coord_directions)
            if confrontantes_completos(_r):
                _d("[Confrontantes] PDF bbox — resultado completo")
                return _build(
                    _r.get("confrontante_norte", ""), _r.get("confrontante_leste", ""),
                    _r.get("confrontante_sul",   ""), _r.get("confrontante_oeste", ""),
                    "pdf", "pdf_bbox",
                    "Confrontantes classificados pelas coordenadas do PDF.",
                )

        if imovel_geom:
            qgis_d = classify_qgis_segments(imovel_geom)
            if can_match_segments_by_order(qgis_d, pdf_segments):
                _r = group_confrontantes_by_directions(pdf_segments, qgis_d)
                if confrontantes_completos(_r):
                    _d("[Confrontantes] PDF geometria QGIS — resultado completo")
                    return _build(
                        _r.get("confrontante_norte", ""), _r.get("confrontante_leste", ""),
                        _r.get("confrontante_sul",   ""), _r.get("confrontante_oeste", ""),
                        "pdf", "pdf_qgis_geom",
                        "Confrontantes classificados pela geometria QGIS.",
                    )

        az_d = classify_segments_by_azimute(pdf_segments)
        if len(az_d) == len(pdf_segments) and az_d:
            _r = group_confrontantes_by_directions(pdf_segments, az_d)
            if confrontantes_completos(_r):
                _d("[Confrontantes] PDF azimute — resultado completo")
                return _build(
                    _r.get("confrontante_norte", ""), _r.get("confrontante_leste", ""),
                    _r.get("confrontante_sul",   ""), _r.get("confrontante_oeste", ""),
                    "pdf", "pdf_azimute",
                    "Confrontantes classificados pelo azimute (fallback).",
                )

    _d("[Confrontantes] ── nenhum método retornou resultado completo ──────────")
    return _build("", "", "", "", "vazio", "nenhum",
                  "Não foi possível definir confrontantes por lado.")


def confrontantes_por_lado(geometry, pdf_segments, confinante_layer=None, servico_id=None):
    try:
        from qgis.core import QgsMessageLog, Qgis as _Qi
        def _tlog(msg):
            QgsMessageLog.logMessage(msg, "GeoDocsSIGEF", _Qi.Info)
    except Exception:
        def _tlog(msg):
            pass

    _tlog("[Trace] confrontantes_por_lado: confinante_layer=%s | servico_id=%s | pdf_segments=%d"
          % (confinante_layer.name() if confinante_layer else "None",
             servico_id, len(pdf_segments or [])))

    # ── Método 1 (prioridade): camada Confinantes Principal ──────────────────
    if confinante_layer is not None and servico_id is not None:
        camada_result = _classificar_confinantes_pela_camada(geometry, confinante_layer, servico_id)
        if camada_result is not None:
            _tlog("[Trace] confrontantes_por_lado: usando CAMADA (fonte=camada_confinantes)")
            if confrontantes_completos(camada_result):
                camada_result["__aviso_confrontantes"] = (
                    "Confrontantes classificados pela camada Confinantes Principal "
                    "(segmento mais próximo)."
                )
            else:
                camada_result["__aviso_confrontantes"] = (
                    "Confrontantes da camada Confinantes Principal "
                    "(alguns lados sem confinante cadastrado)."
                )
            return camada_result
        _tlog("[Trace] confrontantes_por_lado: camada retornou None — caindo para PDF")
    else:
        _tlog("[Trace] confrontantes_por_lado: camada não disponível — usando PDF direto")

    # ── Fallbacks baseados no PDF ─────────────────────────────────────────────
    pdf_result = None

    if pdf_segments:
        # Coordenadas do PDF (caixa envolvente)
        coord_directions = classify_pdf_segments_by_coordinates(pdf_segments)
        if coord_directions:
            pdf_result = group_confrontantes_by_directions(pdf_segments, coord_directions)

        # Geometria QGIS (centróide + ponto médio do anel externo)
        if not confrontantes_completos(pdf_result):
            qgis_directions = classify_qgis_segments(geometry)
            if can_match_segments_by_order(qgis_directions, pdf_segments):
                qgis_result = group_confrontantes_by_directions(pdf_segments, qgis_directions)
                if confrontantes_completos(qgis_result):
                    pdf_result = qgis_result

        # Azimute (último fallback)
        if not confrontantes_completos(pdf_result):
            az_directions = classify_segments_by_azimute(pdf_segments)
            if len(az_directions) == len(pdf_segments) and az_directions:
                az_result = group_confrontantes_by_directions(pdf_segments, az_directions)
                if confrontantes_completos(az_result):
                    az_result["__aviso_confrontantes"] = (
                        "Confrontantes classificados pelo azimute (fallback)."
                    )
                    pdf_result = az_result

    if pdf_result and confrontantes_completos(pdf_result):
        pdf_result["__confrontantes_fonte"] = "pdf"
        _tlog("[Trace] confrontantes_por_lado: retornando resultado PDF")
        return pdf_result

    _tlog("[Trace] confrontantes_por_lado: retornando confrontantes vazios")
    result = empty_confrontantes()
    result["__confrontantes_fonte"] = "pdf"
    result["__aviso_confrontantes"] = (
        "Não foi possível relacionar os segmentos do PDF com segurança. "
        "Confira a geometria do imóvel e o texto do memorial SIGEF."
    )
    return result


def confrontantes_por_lado_pdf(pdf_segments):
    if not pdf_segments:
        result = empty_confrontantes()
        result["__aviso_confrontantes"] = "Nenhum segmento com confrontante foi encontrado no PDF."
        return result

    # Método 1: coordenadas do PDF (reconstrói polígono)
    coord_directions = classify_pdf_segments_by_coordinates(pdf_segments)
    if coord_directions:
        result = group_confrontantes_by_directions(pdf_segments, coord_directions)
        if confrontantes_completos(result):
            result["__aviso_confrontantes"] = (
                "Os confrontantes foram classificados pelas coordenadas do PDF."
            )
            return result

    # Método 2: azimute (fallback)
    az_directions = classify_segments_by_azimute(pdf_segments)
    if len(az_directions) == len(pdf_segments) and az_directions:
        result = group_confrontantes_by_directions(pdf_segments, az_directions)
        if confrontantes_completos(result):
            result["__aviso_confrontantes"] = (
                "Confrontantes classificados pelo azimute do PDF (fallback)."
            )
            return result

    result = empty_confrontantes()
    result["__aviso_confrontantes"] = (
        "Não foi possível classificar os confrontantes pelas coordenadas nem pelo "
        "azimute do PDF."
    )
    return result


def confrontantes_completos(result):
    return any(result.get("confrontante_%s" % direction) for direction in ("norte", "leste", "sul", "oeste"))


def completar_confrontantes_vazios(result, fallback):
    if not confrontantes_completos(fallback):
        return result

    completed = dict(result)
    filled = False
    for direction in ("norte", "leste", "sul", "oeste"):
        key = "confrontante_%s" % direction
        if not completed.get(key) and fallback.get(key):
            completed[key] = fallback[key]
            filled = True

    if filled:
        completed["__aviso_confrontantes"] = (
            "Alguns confrontantes foram completados pelas coordenadas do PDF."
        )
    return completed


def empty_confrontantes():
    return {
        "confrontante_norte": "",
        "confrontante_leste": "",
        "confrontante_sul": "",
        "confrontante_oeste": "",
    }


def can_match_segments_by_order(qgis_directions, pdf_segments):
    if not qgis_directions:
        return False
    difference = abs(len(qgis_directions) - len(pdf_segments))
    return difference <= max(1, int(len(pdf_segments) * 0.1))


def group_confrontantes_by_directions(pdf_segments, directions):
    grouped = {"norte": [], "leste": [], "sul": [], "oeste": []}
    for index, segment in enumerate(pdf_segments):
        if index >= len(directions):
            break
        direction = directions[index]
        if direction is None:   # segmento sem coordenadas válidas — pula
            continue
        confrontante = normalize_confrontante(segment.get("confrontacao", ""))
        if direction in grouped and confrontante and confrontante not in grouped[direction]:
            grouped[direction].append(confrontante)

    return {
        "confrontante_norte": ", ".join(grouped["norte"]),
        "confrontante_leste": ", ".join(grouped["leste"]),
        "confrontante_sul": ", ".join(grouped["sul"]),
        "confrontante_oeste": ", ".join(grouped["oeste"]),
    }


def classify_qgis_segments(geometry):
    points = exterior_ring_points(geometry)
    if len(points) < 2:
        return []

    centroid = geometry.centroid().asPoint()
    directions = []
    for start, end in zip(points, points[1:] + points[:1]):
        midpoint_x = (start.x() + end.x()) / 2.0
        midpoint_y = (start.y() + end.y()) / 2.0
        directions.append(cardinal_direction(midpoint_x - centroid.x(), midpoint_y - centroid.y()))
    return directions


def exterior_ring_points(geometry):
    if geometry is None or geometry.isEmpty():
        return []

    try:
        if geometry.isMultipart():
            polygons = geometry.asMultiPolygon()
            if not polygons:
                return []
            ring = max((polygon[0] for polygon in polygons if polygon), key=ring_area_abs, default=[])
        else:
            polygon = geometry.asPolygon()
            ring = polygon[0] if polygon else []
    except Exception:
        return []

    points = list(ring)
    if len(points) > 1 and same_point(points[0], points[-1]):
        points = points[:-1]
    return points


def ring_area_abs(ring):
    if len(ring) < 3:
        return 0
    area = 0.0
    for start, end in zip(ring, ring[1:] + ring[:1]):
        area += start.x() * end.y() - end.x() * start.y()
    return abs(area) / 2.0


def same_point(a, b):
    return abs(a.x() - b.x()) < 1e-9 and abs(a.y() - b.y()) < 1e-9


def cardinal_direction(dx, dy):
    """Classifica direção cardinal pelo vetor (ponto_médio − centróide).

    dx = diferença de longitude  (positivo → leste)
    dy = diferença de latitude   (positivo → norte)

    Regra: o eixo dominante determina o lado.
        abs(dx) > abs(dy) → leste (dx > 0) ou oeste (dx < 0)
        abs(dy) >= abs(dx) → norte (dy > 0) ou sul  (dy < 0)
    """
    if abs(dx) > abs(dy):
        return "leste" if dx > 0 else "oeste"
    return "norte" if dy > 0 else "sul"


def classify_by_bbox(mx, my, min_x, max_x, min_y, max_y):
    """Classifica o lado da caixa envolvente mais próximo do ponto médio (mx, my).

    Para cada lado calcula a distância do ponto a esse lado e retorna o lado
    com menor distância — robusto para imóveis alongados/estreitos.

        dist_norte = |my − max_y|  (quão perto do topo/norte)
        dist_sul   = |my − min_y|  (quão perto do fundo/sul)
        dist_leste = |mx − max_x|  (quão perto da direita/leste)
        dist_oeste = |mx − min_x|  (quão perto da esquerda/oeste)

    Em caso de empate, a ordem de prioridade é: norte, sul, leste, oeste.
    """
    dists = {
        "norte": abs(my - max_y),
        "sul":   abs(my - min_y),
        "leste": abs(mx - max_x),
        "oeste": abs(mx - min_x),
    }
    return min(dists, key=lambda k: dists[k])


def cardinal_direction_from_azimute(az):
    """Classifica direção cardinal pelo azimute em graus decimais.

    Faixas (sentido do segmento no PDF SIGEF):
        315°-360° e 0°-45°  → norte
        45°-135°             → leste
        135°-225°            → sul
        225°-315°            → oeste
    """
    az = az % 360.0
    if az >= 315 or az < 45:
        return "norte"
    if az < 135:
        return "leste"
    if az < 225:
        return "sul"
    return "oeste"


def parse_azimute_decimal(azimute_str):
    """Converte azimute no formato 'DDD°MM'' (SIGEF) para graus decimais.

    Retorna float ou None se o formato não for reconhecido.
    """
    match = re.match(r"(\d+)\s*[°º]\s*(\d+)'", str(azimute_str).strip())
    if not match:
        return None
    return int(match.group(1)) + int(match.group(2)) / 60.0


def classify_segments_by_azimute(pdf_segments):
    """Classifica cada segmento pela direção cardinal do azimute.

    Retorna lista de strings ('norte'/'leste'/'sul'/'oeste') com o mesmo
    comprimento de pdf_segments, ou lista vazia se algum azimute não puder
    ser convertido.
    """
    directions = []
    for seg in pdf_segments:
        az = parse_azimute_decimal(seg.get("azimute", ""))
        if az is None:
            return []
        directions.append(cardinal_direction_from_azimute(az))
    return directions


def classify_pdf_segments_by_coordinates(pdf_segments):
    """Classifica cada segmento pelo lado da caixa envolvente (bounding box) do polígono.

    Para cada segmento calcula o ponto médio entre vértice inicial e vértice final,
    depois compara a distância desse ponto a cada lado da caixa envolvente:
        dist_norte = |my − max_y|
        dist_sul   = |my − min_y|
        dist_leste = |mx − max_x|
        dist_oeste = |mx − min_x|
    O segmento é classificado pelo lado com menor distância (classify_by_bbox).

    Este método é robusto para imóveis alongados/estreitos onde o centróide
    fica fora do segmento ou próximo a vários lados simultaneamente.

    Quando vante_longitude/vante_latitude não estão disponíveis (parser ferramenta),
    usa as coordenadas do próximo vértice da lista para fechar o polígono.

    Retorna lista do mesmo comprimento que pdf_segments (None para segmentos sem
    coordenadas válidas). Retorna [] se não houver pontos válidos suficientes.
    """
    start_points = []
    end_points = []
    for segment in pdf_segments:
        lon = dms_to_decimal(segment.get("longitude", ""))
        lat = dms_to_decimal(segment.get("latitude", ""))
        start_points.append((lon, lat) if (lon is not None and lat is not None) else None)

        end_lon = dms_to_decimal(segment.get("vante_longitude", ""))
        end_lat = dms_to_decimal(segment.get("vante_latitude", ""))
        end_points.append((end_lon, end_lat) if (end_lon is not None and end_lat is not None) else None)

    valid_points = [p for p in start_points if p is not None]
    if len(valid_points) < 2:
        valid_points = [p for p in end_points if p is not None]
    if len(valid_points) < 2:
        return []

    min_x = min(p[0] for p in valid_points)
    max_x = max(p[0] for p in valid_points)
    min_y = min(p[1] for p in valid_points)
    max_y = max(p[1] for p in valid_points)

    directions = []
    for index, start in enumerate(start_points):
        if start is None:
            directions.append(None)   # mantém alinhamento com pdf_segments
            continue
        end = end_points[index]
        if end is None:
            # Sem vante explícito: usa o próximo vértice de início (polígono fechado)
            next_points = [p for p in start_points[index + 1:] + start_points[:index]
                           if p is not None]
            end = next_points[0] if next_points else start
        midpoint_x = (start[0] + end[0]) / 2.0
        midpoint_y = (start[1] + end[1]) / 2.0
        directions.append(classify_by_bbox(midpoint_x, midpoint_y, min_x, max_x, min_y, max_y))

    return directions


def dms_to_decimal(value):
    value = str(value).strip()
    match = re.search(
        r"(-?\d+(?:[.,]\d+)?)\s*[^0-9,.'-]+\s*(\d+(?:[.,]\d+)?)'\s*(\d+(?:[.,]\d+)?)",
        value,
    )
    if not match:
        return None
    degrees = float(match.group(1).replace(",", "."))
    minutes = float(match.group(2).replace(",", "."))
    seconds = float(match.group(3).replace(",", "."))
    sign = -1 if degrees < 0 else 1
    return sign * (abs(degrees) + minutes / 60.0 + seconds / 3600.0)


def collect_qgis_owners(layer, feature):
    """Fallback legacy: lê proprietários dos campos diretos da camada (proprietario/cpf/cnpj)."""
    owners = []
    for index in range(0, 11):
        suffix = "" if index == 0 else "_%s" % index
        name = feature_value(layer, feature, "proprietario%s" % suffix)
        cpf = feature_value(layer, feature, "cpf%s" % suffix)
        cnpj = feature_value(layer, feature, "cnpj%s" % suffix)
        if not is_valid_field_value(name):
            continue

        owner = {"nome": name, "suffix": suffix}
        if is_valid_field_value(cpf):
            owner["cpf"] = format_cpf(cpf)
        if is_valid_field_value(cnpj):
            owner["cnpj"] = format_cnpj(cnpj)
        owners.append(owner)
    return owners


def get_proprietarios_do_servico(layer, feature):
    """Retorna proprietários via servicos.servico_pessoas + servicos.pessoas.

    Tenta obter os proprietários vinculados ao id do serviço usando a conexão
    PostgreSQL da camada. Se não conseguir (tabelas inexistentes, erro de conexão,
    camada não PostGIS), usa os campos legacy como fallback.
    """
    try:
        id_field = field_name_lookup(layer).get("id")
        if not id_field:
            return collect_qgis_owners(layer, feature)

        servico_id = feature[id_field]
        if not servico_id:
            return collect_qgis_owners(layer, feature)

        owners = _buscar_proprietarios_vinculados(layer, servico_id)
        return owners if owners else collect_qgis_owners(layer, feature)

    except Exception:
        return collect_qgis_owners(layer, feature)


def _buscar_proprietarios_vinculados(layer, servico_id):
    """Executa a consulta SQL nas tabelas de vínculo e retorna lista de owners.

    Tenta duas estratégias em sequência:
    1. psycopg2 direto (mais rápido, parâmetros seguros).
    2. QgsAbstractDatabaseProviderConnection (nativo QGIS 3.x).
    Retorna [] se ambas falharem.
    """
    SQL = """
        SELECT sp.ordem, p.id AS pessoa_id, p.nome, p.tipo_pessoa, p.cpf, p.cnpj,
               p.identidade, p.orgao, p.uf,
               p.data_nascimento, p.profissao, p.estado_civil,
               p.nacionalidade, p.regime_casamento,
               p.nome_da_mae, p.nome_do_pai,
               p.rua, p.numero, p.complemento, p.cep, p.bairro, p.cidade, p.estado,
               p.sexo
        FROM servicos.servico_pessoas sp
        JOIN servicos.pessoas p ON p.id = sp.pessoa_id
        WHERE sp.servico_id = %s
          AND COALESCE(sp.tipo_vinculo, 'PROPRIETARIO') = 'PROPRIETARIO'
        ORDER BY sp.ordem, sp.id
    """
    COLS = [
        "ordem", "pessoa_id", "nome", "tipo_pessoa", "cpf", "cnpj",
        "identidade", "orgao", "uf",
        "data_nascimento", "profissao", "estado_civil",
        "nacionalidade", "regime_casamento",
        "nome_da_mae", "nome_do_pai",
        "rua", "numero", "complemento", "cep", "bairro", "cidade", "estado",
        "sexo",
    ]

    # --- Estratégia 1: psycopg2 ---
    try:
        import psycopg2
        from qgis.core import QgsDataSourceUri

        uri = QgsDataSourceUri(layer.dataProvider().dataSourceUri())
        conn = psycopg2.connect(
            host=uri.host() or "localhost",
            port=int(uri.port()) if uri.port() else 5432,
            dbname=uri.database(),
            user=uri.username(),
            password=uri.password(),
        )
        cur = conn.cursor()
        cur.execute(SQL, (servico_id,))
        rows = cur.fetchall()
        desc_cols = [d[0] for d in cur.description]
        cur.close()
        conn.close()
        return _rows_to_owners(rows, desc_cols)
    except Exception:
        pass

    # --- Estratégia 2: QgsAbstractDatabaseProviderConnection (QGIS 3.x) ---
    try:
        from qgis.core import QgsProviderRegistry

        meta = QgsProviderRegistry.instance().providerMetadata("postgres")
        if meta:
            conn = meta.createConnection(layer.dataProvider().dataSourceUri(), {})
            if conn:
                sql_exec = SQL.replace("%s", str(int(servico_id)))
                rows = conn.executeSql(sql_exec)
                if rows:
                    return _rows_to_owners(rows, COLS)
    except Exception:
        pass

    return []


def _rows_to_owners(rows, cols):
    """Converte linhas SQL em lista de dicionários no formato esperado pelo plugin."""
    owners = []
    for i, row in enumerate(rows):
        record = dict(zip(cols, row))
        nome = clean_value(record.get("nome", ""))
        if not is_valid_field_value(nome):
            continue

        suffix = "" if i == 0 else "_%d" % i
        owner = {"nome": nome, "suffix": suffix}

        # pessoa_id e tipo_pessoa — usados pela declaração
        if record.get("pessoa_id"):
            owner["pessoa_id"] = record["pessoa_id"]
        tipo = clean_value(record.get("tipo_pessoa", ""))
        if is_valid_field_value(tipo):
            owner["tipo_pessoa"] = tipo

        cpf = clean_value(record.get("cpf", ""))
        cnpj = clean_value(record.get("cnpj", ""))
        if is_valid_field_value(cpf):
            owner["cpf"] = format_cpf(cpf)
        elif is_valid_field_value(cnpj):
            owner["cnpj"] = format_cnpj(cnpj)

        for field in (
            "identidade", "orgao", "uf", "data_nascimento", "profissao",
            "estado_civil", "nacionalidade", "regime_casamento",
            "nome_da_mae", "nome_do_pai",
            "rua", "numero", "complemento", "cep", "bairro", "cidade", "estado",
            "sexo",
        ):
            val = clean_value(record.get(field, ""))
            if is_valid_field_value(val):
                owner[field] = val

        owners.append(owner)
    return owners


def build_owner_block(owners):
    blocks = []
    for owner in owners:
        lines = ["Proprietário(a): %s" % owner["nome"]]
        if owner.get("cpf"):
            lines.append("CPF: %s" % owner["cpf"])
        elif owner.get("cnpj"):
            lines.append("CNPJ: %s" % owner["cnpj"])
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks)


def build_owner_signature_block(owners):
    blocks = []
    for owner in owners:
        document = ""
        if owner.get("cpf"):
            document = "CPF: %s" % owner["cpf"]
        elif owner.get("cnpj"):
            document = "CNPJ: %s" % owner["cnpj"]
        block = "________________________________________\nProprietário(a): %s" % owner["nome"]
        if document:
            block += "\n%s" % document
        blocks.append(block)
    return "\n\n".join(blocks)


def substituir_marcadores_docx(doc, dados):
    count = 0
    count += replace_in_paragraphs(doc.paragraphs, dados)
    count += replace_in_tables(doc.tables, dados)

    for section in doc.sections:
        for container in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            count += replace_in_paragraphs(container.paragraphs, dados)
            count += replace_in_tables(container.tables, dados)
    return count


def aplicar_regras_de_paginacao(doc):
    paragraphs = list(all_document_paragraphs(doc))
    marcar_blocos_assinatura_tecnica(paragraphs)
    marcar_bloco_declaracao(paragraphs)

    for paragraph in paragraphs:
        text_key = normalize_key(paragraph.text)
        if "proprietario_interessado" in text_key or "art_213" in text_key or "lei_n_6_015_73" in text_key:
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = True


def marcar_bloco_declaracao(paragraphs):
    for index, paragraph in enumerate(paragraphs):
        text_key = normalize_key(paragraph.text)
        if text_key != "declaracao" and not text_key.startswith("declaracao"):
            continue

        paragraph.paragraph_format.page_break_before = True
        end_index = find_signature_block_end_after(paragraphs, index)
        mark_keep_block(paragraphs, index, end_index)


def marcar_blocos_assinatura_tecnica(paragraphs):
    for index, paragraph in enumerate(paragraphs):
        if not is_responsible_signature_paragraph(normalize_key(paragraph.text)):
            continue
        start_index = find_signature_block_start(paragraphs, index)
        end_index = find_signature_block_end(paragraphs, index)
        mark_keep_block(paragraphs, start_index, end_index)


def find_signature_block_start(paragraphs, index):
    start = index
    while start > 0 and index - start < 5:
        previous = paragraphs[start - 1]
        previous_key = normalize_key(previous.text)
        if paragraph_has_image(previous) or is_signature_separator(previous.text) or not previous_key:
            start -= 1
            continue
        break
    return start


def find_signature_block_end(paragraphs, index):
    end = index
    while end + 1 < len(paragraphs) and end - index < 8:
        next_paragraph = paragraphs[end + 1]
        next_key = normalize_key(next_paragraph.text)
        if not next_key:
            break
        if is_responsible_signature_paragraph(next_key) or is_short_signature_line(next_key):
            end += 1
            continue
        break
    return end


def find_signature_block_end_after(paragraphs, start_index):
    end = min(start_index + 2, len(paragraphs) - 1)
    for index in range(start_index + 1, len(paragraphs)):
        if is_responsible_signature_paragraph(normalize_key(paragraphs[index].text)):
            return find_signature_block_end(paragraphs, index)
        if index - start_index > 14:
            break
        end = index
    return end


def mark_keep_block(paragraphs, start_index, end_index):
    for index in range(start_index, end_index + 1):
        paragraph = paragraphs[index]
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = index < end_index


def paragraph_has_image(paragraph):
    xml = paragraph._p.xml
    return "<w:drawing" in xml or "<w:pict" in xml


def is_signature_separator(text):
    compact = re.sub(r"\s+", "", str(text))
    return bool(compact) and set(compact) <= {"_", "-", "—", "–"}


def is_short_signature_line(text_key):
    if len(text_key) > 80:
        return False
    tokens = (
        "tecnico_agricola",
        "agrimensura",
        "cfta",
        "codigo_de_credenciamento",
        "incra",
        "trt",
        "uf",
    )
    return any(token in text_key for token in tokens)


def all_document_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for paragraph in table_paragraphs(table):
            yield paragraph
    for section in doc.sections:
        for container in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                for paragraph in table_paragraphs(table):
                    yield paragraph


def table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                for paragraph in table_paragraphs(nested_table):
                    yield paragraph


def is_responsible_signature_paragraph(text_key):
    tokens = (
        "resp_tec",
        "bruno_feliciano_de_lima_alves",
        "tecnico_agricola",
        "agrimensura_cfta",
        "codigo_de_credenciamento",
    )
    return any(token in text_key for token in tokens)


def replace_in_tables(tables, data):
    count = 0
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                count += replace_in_paragraphs(cell.paragraphs, data)
                count += replace_in_tables(cell.tables, data)
    return count


def replace_in_paragraphs(paragraphs, data):
    count = 0
    for paragraph in paragraphs:
        count += replace_placeholders_in_paragraph(paragraph, data)
    return count


def replace_placeholders_in_paragraph(paragraph, data):
    runs = paragraph.runs
    if not runs:
        return 0

    full_text = "".join(run.text for run in runs)
    matches = list(PLACEHOLDER_PATTERN.finditer(full_text))
    if not matches:
        return 0

    if paragraph_is_single_placeholder(full_text, "bloco_proprietarios"):
        write_owner_block_paragraph(paragraph, data.get("__owners", []))
        return 1

    if paragraph_is_single_placeholder(full_text, "bloco_assinaturas_proprietarios"):
        write_owner_signature_block_paragraph(paragraph, data.get("__owners", []))
        return 1

    run_ranges = []
    position = 0
    for run in runs:
        start = position
        end = start + len(run.text)
        run_ranges.append((run, start, end))
        position = end

    for match in reversed(matches):
        replacement = value_for_placeholder(match, data)
        replace_match_across_runs(runs, run_ranges, match.start(), match.end(), replacement)
    return len(matches)


def paragraph_is_single_placeholder(text, expected_key):
    match = PLACEHOLDER_PATTERN.fullmatch(text.strip())
    if not match:
        return False
    raw_key = match.group(2) or match.group(3) or match.group(4)
    return normalize_key(raw_key) == expected_key


def write_owner_block_paragraph(paragraph, owners):
    clear_paragraph_content(paragraph)
    paragraph.paragraph_format.keep_together = True

    for index, owner in enumerate(owners):
        if index:
            paragraph.add_run().add_break()

        add_labeled_value(
            paragraph,
            "Proprietário(a):",
            owner["nome"],
            label_bold=True,
            value_bold=False,
            signature_font=True,
        )
        paragraph.add_run().add_break()
        if owner.get("cpf"):
            add_labeled_value(paragraph, "CPF:", owner["cpf"], label_bold=True, value_bold=False, signature_font=True)
        elif owner.get("cnpj"):
            add_labeled_value(paragraph, "CNPJ:", owner["cnpj"], label_bold=True, value_bold=False, signature_font=True)


def write_owner_signature_block_paragraph(paragraph, owners):
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    clear_paragraph_content(paragraph)
    paragraph.paragraph_format.keep_together = True

    for index, owner in enumerate(owners):
        if index:
            add_blank_signature_lines(paragraph, 6)

        line_run = paragraph.add_run(signature_line_for_owner(owner))
        set_signature_run_font(line_run)
        paragraph.add_run().add_break()
        add_labeled_value(
            paragraph,
            "Proprietário(a):",
            owner["nome"],
            label_bold=False,
            value_bold=True,
            signature_font=True,
        )
        paragraph.add_run().add_break()
        if owner.get("cpf"):
            add_labeled_value(paragraph, "CPF:", owner["cpf"], label_bold=False, value_bold=True, signature_font=True)
        elif owner.get("cnpj"):
            add_labeled_value(paragraph, "CNPJ:", owner["cnpj"], label_bold=False, value_bold=True, signature_font=True)


def add_blank_signature_lines(paragraph, count):
    for _index in range(count):
        paragraph.add_run().add_break()


def add_labeled_value(paragraph, label, value, label_bold, value_bold, signature_font=False):
    label_run = paragraph.add_run(label + " ")
    if signature_font:
        set_signature_run_font(label_run, label_bold)
    else:
        label_run.bold = label_bold
    value_run = paragraph.add_run(value)
    if signature_font:
        set_signature_run_font(value_run, value_bold)
    else:
        value_run.bold = value_bold


def set_signature_run_font(run, bold=None):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Arial"
    run.font.size = Pt(12)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is not None:
        r_fonts.set(qn("w:ascii"), "Arial")
        r_fonts.set(qn("w:hAnsi"), "Arial")
        r_fonts.set(qn("w:cs"), "Arial")


def signature_line_for_owner(owner):
    reference = "Proprietário(a): %s" % owner.get("nome", "")
    length = max(45, min(len(reference), 70))
    return "_" * length


def clear_paragraph_content(paragraph):
    p_pr = paragraph._p.pPr
    paragraph._p.clear_content()
    if p_pr is not None and paragraph._p.pPr is None:
        paragraph._p.insert(0, p_pr)


def replace_match_across_runs(runs, run_ranges, start, end, replacement):
    affected = [
        (index, run_start, run_end)
        for index, (_run, run_start, run_end) in enumerate(run_ranges)
        if run_start < end and run_end > start
    ]
    if not affected:
        return

    first_index, first_start, _first_end = affected[0]
    last_index, last_start, _last_end = affected[-1]

    first_run = runs[first_index]
    if first_index == last_index:
        local_start = start - first_start
        local_end = end - first_start
        first_run.text = first_run.text[:local_start] + replacement + first_run.text[local_end:]
        return

    first_run.text = first_run.text[:start - first_start] + replacement
    for middle_index, _middle_start, _middle_end in affected[1:-1]:
        runs[middle_index].text = ""
    last_run = runs[last_index]
    last_run.text = last_run.text[end - last_start:]


def value_for_placeholder(match, data):
    raw_key = match.group(2) or match.group(3) or match.group(4)
    key = normalize_key(raw_key)
    return data.get(key, "")


def build_data_preview(data):
    lines = []
    for key in sorted(public_data(data)):
        value = data[key]
        if value:
            lines.append("%s: %s" % (key, value))
    return "\n".join(lines)


def public_data(data):
    return {key: value for key, value in data.items() if not key.startswith("__")}


def set_value(data, key, value):
    key = normalize_key(key)
    value = clean_value(value)
    if key and value:
        data[key] = value


def clean_value(value):
    value = normalize_text(str(value))
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def normalize_text(value):
    return "".join(char for char in str(value) if is_valid_xml_char(char))


def is_valid_xml_char(char):
    code = ord(char)
    return code in (9, 10, 13) or 32 <= code <= 0xD7FF or 0xE000 <= code <= 0xFFFD or 0x10000 <= code <= 0x10FFFF


def normalize_key(value):
    value = unicodedata.normalize("NFKD", str(value).strip().lower())
    value = "".join(char for char in value if not unicodedata.combining(char))
    value = re.sub(r"[^a-z0-9]+", "_", value)
    return value.strip("_")
