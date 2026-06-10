import os
import shutil
import copy

from docx.shared import Pt

from ..docx_utils import (
    ensure_docx_path,
    substituir_marcadores_docx,
    normalize_key,
    all_document_paragraphs,
    PLACEHOLDER_PATTERN,
    clear_paragraph_content,
    formatar_data_extenso,
)
from .pessoas_utils import (
    buscar_proprietarios_servico,
    enriquecer_proprietarios,
    montar_bloco_proprietarios_resumido_segs,
)

_ARIAL = "Arial"
_PT12  = Pt(12)


def find_laudo_template(plugin_dir):
    path = os.path.join(plugin_dir, "models", "laudo_tecnico.docx")
    return path if os.path.exists(path) else None


def fill_laudo_tecnico_template(template_path, output_path, data, layer=None):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Biblioteca python-docx não encontrada. Instale python-docx no Python do QGIS."
        ) from exc

    template_path = os.path.abspath(template_path)
    output_path   = os.path.abspath(ensure_docx_path(output_path))

    if template_path == output_path:
        raise RuntimeError("Escolha um arquivo de saída diferente do modelo original.")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copyfile(template_path, output_path)

    document   = Document(output_path)
    laudo_data = copy.copy(data)

    # ── Datas por extenso ────────────────────────────────────────
    cert_convertida = formatar_data_extenso(laudo_data.get("data_certificacao"))
    if cert_convertida:
        laudo_data["data_certificacao"] = cert_convertida

    mat_convertida = formatar_data_extenso(laudo_data.get("data_matricula"))
    if mat_convertida:
        laudo_data["data_matricula"] = mat_convertida

    # ── Fallbacks para campos obrigatórios ───────────────────────
    if not laudo_data.get("nome_cartorio"):
        laudo_data["nome_cartorio"] = ""
    if not laudo_data.get("data_matricula"):
        laudo_data["data_matricula"] = "[DATA DA MATRÍCULA NÃO INFORMADA]"

    # ── Proprietários: busca direta pelo servico_id ──────────────
    owners = []
    if layer:
        servico_id = laudo_data.get("__servico_id")
        if servico_id:
            try:
                owners = buscar_proprietarios_servico(layer, servico_id)
            except Exception:
                pass
    if not owners:
        owners = laudo_data.get("__owners", [])

    declarantes = enriquecer_proprietarios(owners, layer)
    avisos = []

    # ── Placeholders em negrito + Arial 12 (antes do substituir_marcadores) ──
    for placeholder_key in ("denominacao", "codigo_da_certificacao", "matricula"):
        valor = str(laudo_data.get(placeholder_key) or "").strip()
        if valor:
            _substituir_placeholder_bold_arial(document, placeholder_key, valor)

    # ── {{BLOCO_PROPRIETARIOS}} — formato resumido ───────────────
    _substituir_bloco_proprietarios_laudo(document, declarantes, avisos)

    # ── Demais placeholders (area, municipio, uf, trt, etc.) ─────
    substituir_marcadores_docx(document, laudo_data)

    document.save(output_path)
    return avisos


# ── Placeholder em negrito + Arial 12 ────────────────────────

def _substituir_placeholder_bold_arial(document, placeholder_key, valor):
    """Substitui {{placeholder_key}} inserindo o valor em negrito com Arial 12.

    O texto antes e depois do placeholder preserva a formatação original do template.
    """
    key_norm = normalize_key(placeholder_key)
    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        match = next(
            (m for m in PLACEHOLDER_PATTERN.finditer(full_text)
             if normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == key_norm),
            None,
        )
        if not match:
            continue
        segs_antes  = _segmentos_com_formato(runs, 0, match.start())
        segs_depois = _segmentos_com_formato(runs, match.end(), len(full_text))
        clear_paragraph_content(paragraph)
        for texto, rPr in segs_antes:
            _adicionar_run_formatado(paragraph, texto, rPr)
        r = paragraph.add_run(valor)
        r.bold = True
        r.font.name = _ARIAL
        r.font.size = _PT12
        for texto, rPr in segs_depois:
            _adicionar_run_formatado(paragraph, texto, rPr)


# ── Bloco de proprietários ────────────────────────────────────

def _substituir_bloco_proprietarios_laudo(document, declarantes, avisos_out):
    """Substitui {{BLOCO_PROPRIETARIOS}} com bloco resumido: nome + CPF/CNPJ + cônjuge.

    Nome e CPF/CNPJ em negrito; rótulo do cônjuge sem negrito. Arial 12 em todos os runs.
    O texto antes/depois preserva a formatação do template.
    """
    todos_segs = montar_bloco_proprietarios_resumido_segs(declarantes)

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue

        full_text = "".join(run.text for run in runs)
        matches = [
            m for m in PLACEHOLDER_PATTERN.finditer(full_text)
            if normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_proprietarios"
        ]
        if not matches:
            continue

        match = matches[0]

        segs_antes  = _segmentos_com_formato(runs, 0, match.start())
        segs_depois = _segmentos_com_formato(runs, match.end(), len(full_text))

        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = True

        for texto, rPr_clone in segs_antes:
            _adicionar_run_formatado(paragraph, texto, rPr_clone)

        for texto, bold in todos_segs:
            if texto:
                r = paragraph.add_run(texto)
                r.bold = bold
                r.font.name = _ARIAL
                r.font.size = _PT12

        for texto, rPr_clone in segs_depois:
            _adicionar_run_formatado(paragraph, texto, rPr_clone)


# ── Helpers de reconstrução de runs ──────────────────────────

def _segmentos_com_formato(runs, text_start, text_end):
    """Retorna [(texto, rPr_clone)] para o intervalo [text_start, text_end) dos runs."""
    from copy import deepcopy
    resultado = []
    pos = 0
    for run in runs:
        run_start = pos
        run_end   = pos + len(run.text)
        pos       = run_end

        seg_start = max(run_start, text_start)
        seg_end   = min(run_end,   text_end)
        if seg_start >= seg_end:
            continue

        texto = run.text[seg_start - run_start : seg_end - run_start]
        if not texto:
            continue

        rPr = run._r.rPr
        resultado.append((texto, deepcopy(rPr) if rPr is not None else None))
    return resultado


def _adicionar_run_formatado(paragraph, texto, rPr_clone):
    """Adiciona um run com rPr clonado (preserva fonte, tamanho, cor do template)."""
    if not texto:
        return
    run = paragraph.add_run(texto)
    if rPr_clone is None:
        return
    r_elem = run._r
    existing = r_elem.rPr
    if existing is not None:
        r_elem.remove(existing)
    r_elem.insert(0, rPr_clone)
