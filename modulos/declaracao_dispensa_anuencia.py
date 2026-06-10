import os
import re
import copy
import shutil
import unicodedata

from docx.shared import Pt

from ..docx_utils import (
    ensure_docx_path,
    substituir_marcadores_docx,
    normalize_key,
    all_document_paragraphs,
    PLACEHOLDER_PATTERN,
    clear_paragraph_content,
    marcar_blocos_assinatura_tecnica,
    formatar_trt_sem_duplicar_uf,
    formatar_municipio_para_data,
)
from .pessoas_utils import (
    enriquecer_proprietarios,
    buscar_proprietarios_servico,
    buscar_dados_servico,
    limpar_valor,
    montar_assinaturas_pessoas,
)
from .requerimento_retificacao_area import (
    _fval,
    _formatar_nome_proprio,
)
from .requerimento_desmembramento import (
    _parse_float,
    _formatar_area_ha,
    _formatar_perimetro_str,
)
from .declaracao_confrontantes import (
    _substituir_bloco_proprietarios_inline as _sub_bloco_proprietarios,
)
from .classificacao_confrontantes import classificar_confrontante_dispensa

_ARIAL = "Arial"
_PT12  = Pt(12)

# Primeiras palavras de nomes masculinos (artigo "Do")
_MASCULINOS = frozenset({
    "RIO", "RIACHO", "CORREGO", "ACUDE", "ARROIO",
    "LAGO", "CANAL", "IGARAPE", "RIBEIRAO",
    "RESERVATORIO", "TERRENO", "GASODUTO", "OLEODUTO",
    "PATRIMONIO", "PARQUE", "TERRITORIO", "CAMINHO",
    "MUNICIPIO", "ESTADO", "GOVERNO",
})


# ── Template ───────────────────────────────────────────────────────────────────

def find_declaracao_dispensa_anuencia_template(plugin_dir):
    path = os.path.join(plugin_dir, "models", "declaracao_dispensa_anuencia.docx")
    return path if os.path.exists(path) else None


# ── Artigo (Do / Da) ──────────────────────────────────────────────────────────

def _construir_artigo(nome):
    """Retorna 'Do' ou 'Da' conforme o gênero da primeira palavra do nome."""
    primeira = nome.strip().split()[0] if nome.strip() else ""
    norm = unicodedata.normalize("NFKD", primeira.upper())
    norm = "".join(c for c in norm if not unicodedata.combining(c))
    return "Do" if norm in _MASCULINOS else "Da"


# ── Leitura de nome do confrontante ──────────────────────────────────────────

def _ler_nome_confrontante(layer, feature):
    """Lê o nome do confrontante: prioridade nome_propriedade → nome → descricao."""
    for campo in ("nome_propriedade", "nome", "descricao"):
        v = _fval(layer, feature, campo)
        if v:
            return v.strip()
    return ""


# ── Construção dos itens de dispensa ─────────────────────────────────────────

def _montar_itens_dispensa(confinante_features, layer, municipio, uf):
    """Constrói lista de dicts com artigo, nome e texto complementar."""
    itens = []
    for feat in confinante_features:
        nome = _ler_nome_confrontante(layer, feat)
        if not nome:
            fid  = _fval(layer, feat, "id") or str(feat.id())
            nome = "Confrontante ID %s" % fid
        artigo    = _construir_artigo(nome)
        resultado = classificar_confrontante_dispensa(nome, municipio=municipio, uf=uf)
        itens.append({"artigo": artigo, "nome": nome, "texto": resultado["texto"]})
    return itens


def _montar_bullets_dispensa(itens):
    """Monta lista de strings de bullet para o bloco de dispensa."""
    bullets = []
    for item in itens:
        bullets.append(
            "• %s %s, %s;" % (item["artigo"], item["nome"], item["texto"])
        )
    return bullets


# ── Substituição do BLOCO_DISPENSA_ANUENCIA ───────────────────────────────────

def _substituir_bloco_dispensa(document, itens):
    """Substitui {{BLOCO_DISPENSA_ANUENCIA}} pela lista dinâmica de bullets."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _Oxm

    def _criar_para_bullet(texto, ref_elem):
        new_p = _Oxm("w:p")
        pPr = ref_elem.find(_qn("w:pPr"))
        if pPr is not None:
            new_p.append(copy.deepcopy(pPr))
        new_r = _Oxm("w:r")
        rPr_elem = _Oxm("w:rPr")
        rFonts = _Oxm("w:rFonts")
        rFonts.set(_qn("w:ascii"), _ARIAL)
        rFonts.set(_qn("w:hAnsi"), _ARIAL)
        rPr_elem.append(rFonts)
        sz = _Oxm("w:sz")
        sz.set(_qn("w:val"), str(int(_PT12.pt * 2)))
        rPr_elem.append(sz)
        new_r.append(rPr_elem)
        t = _Oxm("w:t")
        t.text = texto
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_r.append(t)
        new_p.append(new_r)
        return new_p

    for paragraph in all_document_paragraphs(document):
        if not paragraph.runs:
            continue
        full_text = "".join(r.text for r in paragraph.runs)
        if not any(
            normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_dispensa_anuencia"
            for m in PLACEHOLDER_PATTERN.finditer(full_text)
        ):
            continue

        bullets  = _montar_bullets_dispensa(itens)
        ref_elem = paragraph._element

        if not bullets:
            clear_paragraph_content(paragraph)
            r = paragraph.add_run("[NENHUM CONFRONTANTE SELECIONADO]")
            r.font.name = _ARIAL
            r.font.size = _PT12
            return

        # Primeiro bullet no próprio parágrafo placeholder
        clear_paragraph_content(paragraph)
        r = paragraph.add_run(bullets[0])
        r.font.name = _ARIAL
        r.font.size = _PT12

        # Demais bullets: novos parágrafos inseridos sequencialmente
        prev = ref_elem
        for bullet in bullets[1:]:
            new_p = _criar_para_bullet(bullet, ref_elem)
            prev.addnext(new_p)
            prev = new_p

        return


# ── Bloco de assinaturas dos declarantes ─────────────────────────────────────

def _substituir_bloco_assinaturas_declarantes(document, proprietarios):
    """Substitui {{BLOCO_ASSINATURAS_DECLARANTES}} com linhas de assinatura."""
    assinaturas = montar_assinaturas_pessoas(proprietarios)
    if not assinaturas:
        return

    for paragraph in all_document_paragraphs(document):
        if not paragraph.runs:
            continue
        full_text = "".join(r.text for r in paragraph.runs)
        if not any(
            normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_assinaturas_declarantes"
            for m in PLACEHOLDER_PATTERN.finditer(full_text)
        ):
            continue

        clear_paragraph_content(paragraph)
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
        paragraph.paragraph_format.keep_together = True

        for i, (linha1, linha2) in enumerate(assinaturas):
            if i:
                for _ in range(4):
                    paragraph.add_run().add_break()
            paragraph.add_run("________________________________________")
            paragraph.add_run().add_break()
            r = paragraph.add_run(linha1)
            r.bold = True
            if linha2:
                for parte in linha2.split("\n"):
                    paragraph.add_run().add_break()
                    paragraph.add_run(parte)

        return


# ── Função principal ───────────────────────────────────────────────────────────

def fill_declaracao_dispensa_anuencia_template(
    template_path, output_path, data,
    serv_layer=None, serv_feature=None,
    confinante_layer=None, confinante_features=None,
):
    """Gera a Declaração de Dispensa de Anuência de Confrontantes.

    Args:
        data:                 dict com dados do PDF + QGIS (_collect_memorial_data).
        serv_layer:           camada QGIS Serviços 2 (para DB queries).
        serv_feature:         feature do ID digitado (dados gerais do imóvel).
        confinante_layer:     camada Confinantes Principal.
        confinante_features:  lista de features selecionadas na camada de confinantes.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Biblioteca python-docx não encontrada. Instale python-docx no Python do QGIS."
        ) from exc

    template_path = os.path.abspath(template_path)
    output_path   = os.path.abspath(ensure_docx_path(output_path))
    if template_path == output_path:
        raise RuntimeError("Escolha um arquivo de saída diferente do modelo original.")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copyfile(template_path, output_path)

    document = Document(output_path)
    req_data = copy.copy(data)
    avisos   = []

    servico_id = req_data.get("__servico_id")

    # ── Dados do serviço via DB ─────────────────────────────────────────────
    servico_data = {}
    if serv_layer and servico_id:
        try:
            servico_data = buscar_dados_servico(serv_layer, servico_id)
        except Exception:
            pass

    def _sd(key):
        """Lê campo: DB → PDF data → feature field."""
        v = limpar_valor(servico_data.get(key, ""))
        if v and v.upper() not in ("NULL", "NONE"):
            return v
        v2 = limpar_valor(str(req_data.get(key, "")))
        if v2 and v2.upper() not in ("NULL", "NONE"):
            return v2
        return _fval(serv_layer, serv_feature, key)

    # ── Proprietários do serviço base ────────────────────────────────────────
    owners = req_data.get("__owners", [])
    if not owners and serv_layer and servico_id:
        try:
            owners = buscar_proprietarios_servico(serv_layer, servico_id)
        except Exception:
            pass
    declarantes = enriquecer_proprietarios(owners, serv_layer)

    # ── Campos básicos ────────────────────────────────────────────────────────
    denominacao   = _sd("denominacao")  or req_data.get("denominacao", "")
    matricula     = _sd("matricula")    or req_data.get("matricula", "")
    codigo_incra  = _sd("codigo_incra") or req_data.get("codigo_incra", "")
    folha         = _sd("folha")        or "-"
    livro         = _sd("livro")        or "-"
    ret_aver      = _sd("ret_aver")     or "-"

    nome_cartorio_raw = (
        _sd("nome_cartorio") or _sd("cartorio") or
        req_data.get("nome_cartorio") or req_data.get("cartorio", "")
    )

    trt_raw = _sd("trt") or req_data.get("trt", "")
    if trt_raw:
        trt_raw = re.sub(
            r"^\s*T\.?R\.?T\.?\s*:?\s*", "", trt_raw, flags=re.IGNORECASE
        ).strip()

    # Área: formata como "X,XXXX ha" com 4 casas decimais e vírgula decimal
    area_raw = _sd("area_ha") or _sd("area") or req_data.get("area_ha", req_data.get("area", ""))
    area_fmt = _formatar_area_ha(area_raw) if area_raw else "[ÁREA NÃO INFORMADA]"

    # Perímetro: formata como "X.XXX,XX m" com 2 casas decimais, ponto de milhar
    perim_raw = (
        _sd("perimetro") or _sd("perimetro_m") or
        req_data.get("perimetro", req_data.get("perimetro_m", ""))
    )
    perim_fmt = _formatar_perimetro_str(perim_raw) if perim_raw else "[PERÍMETRO NÃO INFORMADO]"

    # Município e UF separados
    municipio = _sd("municipio") or ""
    uf        = (_sd("uf") or "").upper()
    if not municipio:
        mun_uf_raw = req_data.get("municipio_uf", "")
        if "/" in mun_uf_raw:
            partes    = mun_uf_raw.rsplit("/", 1)
            municipio = partes[0].strip()
            if not uf:
                uf = partes[1].strip().upper()
        else:
            municipio = mun_uf_raw.strip()

    # ── Itens de dispensa (confrontantes selecionados) ───────────────────────
    itens = []
    if confinante_features and confinante_layer:
        itens = _montar_itens_dispensa(
            confinante_features, confinante_layer, municipio, uf
        )
    else:
        avisos.append(
            "Nenhum confrontante selecionado; "
            "{{BLOCO_DISPENSA_ANUENCIA}} não será substituído."
        )

    # ── Dict de substituição ─────────────────────────────────────────────────
    req_data["denominacao"]   = denominacao   or "[DENOMINAÇÃO NÃO INFORMADA]"
    req_data["municipio"]      = municipio     or "[MUNICÍPIO NÃO INFORMADO]"
    req_data["municipio_data"] = formatar_municipio_para_data(municipio) if municipio else ""
    req_data["uf"]             = uf            or "[UF NÃO INFORMADA]"
    req_data["nome_cartorio"]  = (
        _formatar_nome_proprio(nome_cartorio_raw)
        if nome_cartorio_raw
        else ""
    )
    req_data["matricula"]      = matricula     or ""
    req_data["folha"]          = folha
    req_data["livro"]          = livro
    req_data["ret_aver"]       = ret_aver
    req_data["codigo_incra"]   = codigo_incra  or "[CÓDIGO INCRA NÃO INFORMADO]"
    req_data["area"]           = area_fmt
    req_data["perimetro"]      = perim_fmt
    req_data["trt"]            = formatar_trt_sem_duplicar_uf(trt_raw, uf) or "[TRT NÃO INFORMADO]"

    # ── Substituições especializadas (antes de substituir_marcadores_docx) ───
    _sub_bloco_proprietarios(document, declarantes, avisos)
    _substituir_bloco_dispensa(document, itens)
    _substituir_bloco_assinaturas_declarantes(document, declarantes)

    # ── Demais placeholders simples ───────────────────────────────────────────
    substituir_marcadores_docx(document, req_data)

    # ── Manter bloco RT junto na mesma página ─────────────────────────────────
    paragraphs = list(all_document_paragraphs(document))
    marcar_blocos_assinatura_tecnica(paragraphs)

    document.save(output_path)
    return avisos
