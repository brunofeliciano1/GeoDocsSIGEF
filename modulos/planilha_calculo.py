import os
import shutil

from ..docx_utils import (
    ensure_docx_path,
    substituir_marcadores_docx,
    paragraph_is_single_placeholder,
    normalize_key,
    all_document_paragraphs,
    clear_paragraph_content,
    add_labeled_value,
    set_signature_run_font,
    marcar_blocos_assinatura_tecnica,
)
from .pessoas_utils import (
    enriquecer_proprietarios,
    normalizar_nome_destaque,
    limpar_valor,
)


def find_planilha_template(plugin_dir):
    path = os.path.join(plugin_dir, "models", "planilha_de_calculo.docx")
    return path if os.path.exists(path) else None


def fill_planilha_calculo_template(template_path, output_path, data, layer=None):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Biblioteca python-docx não encontrada. Instale python-docx no Python do QGIS."
        ) from exc

    template_path = os.path.abspath(template_path)
    output_path = os.path.abspath(ensure_docx_path(output_path))

    if template_path == output_path:
        raise RuntimeError("Escolha um arquivo de saída diferente do modelo original.")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copyfile(template_path, output_path)

    document = Document(output_path)
    segments  = data.get("__pdf_segments", [])

    # Enriquece owners com cônjuge via pessoas_utils
    owners      = data.get("__owners", [])
    declarantes = enriquecer_proprietarios(owners, layer)

    # Placeholders especiais são resolvidos ANTES da substituição normal para
    # não serem apagados como texto vazio.
    _replace_tabela_placeholder(document, segments)
    _replace_bloco_proprietarios_planilha(document, declarantes)
    substituir_marcadores_docx(document, data)
    # Correção de paginação: assinatura do RT não pode separar da imagem
    _manter_assinatura_junto(document)
    document.save(output_path)
    return len(segments)


def _replace_tabela_placeholder(doc, segments):
    to_replace = [
        p for p in doc.paragraphs
        if paragraph_is_single_placeholder(p.text, "tabela")
    ]
    for paragraph in to_replace:
        _insert_vertex_table_replacing_paragraph(doc, paragraph, segments)


def _replace_bloco_proprietarios_planilha(doc, declarantes):
    """Substitui {{BLOCO_PROPRIETARIOS}} com formato resumido, incluindo cônjuge.

    Formato por entrada (Arial 12):
      Proprietário(a): NOME   ← rótulo bold, nome não-bold
      CPF: xxx               ← rótulo bold, valor não-bold

    Cônjuge aparece como entrada separada imediatamente abaixo, sem linha extra.
    """
    to_replace = [
        p for p in all_document_paragraphs(doc)
        if "bloco_proprietarios" in normalize_key(p.text)
    ]
    for paragraph in to_replace:
        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = True

        first_entry = True
        for decl in declarantes:
            if not first_entry:
                paragraph.add_run().add_break()
            first_entry = False

            nome = normalizar_nome_destaque(decl.get("nome", ""))
            _add_owner_entry_planilha(paragraph, nome, decl.get("cpf"), decl.get("cnpj"))

            conjuge = decl.get("_conjuge")
            if conjuge and limpar_valor(conjuge.get("nome", "")):
                paragraph.add_run().add_break()
                cnome = normalizar_nome_destaque(conjuge["nome"])
                _add_owner_entry_planilha(paragraph, cnome, conjuge.get("cpf"), conjuge.get("cnpj"))


def _add_owner_entry_planilha(paragraph, nome, cpf=None, cnpj=None):
    """Escreve 'Proprietário(a): NOME' + 'CPF/CNPJ: xxx' com Arial 12."""
    add_labeled_value(paragraph, "Proprietário(a):", nome or "____________________",
                      label_bold=True, value_bold=False, signature_font=True)
    paragraph.add_run().add_break()
    cpf_val  = limpar_valor(cpf  or "")
    cnpj_val = limpar_valor(cnpj or "")
    if cpf_val:
        add_labeled_value(paragraph, "CPF:", cpf_val,
                          label_bold=True, value_bold=False, signature_font=True)
    elif cnpj_val:
        add_labeled_value(paragraph, "CNPJ:", cnpj_val,
                          label_bold=True, value_bold=False, signature_font=True)


def _manter_assinatura_junto(doc):
    """Evita que a imagem da assinatura se separe do bloco de identificação do RT.

    Usa all_document_paragraphs() e marcar_blocos_assinatura_tecnica() do docx_utils —
    mesma lógica robusta do Memorial.
    """
    paragraphs = list(all_document_paragraphs(doc))
    marcar_blocos_assinatura_tecnica(paragraphs)


def _set_table_autofit_to_window(table):
    """AutoAjustar à janela: tabela ocupa 100% da largura útil da página."""
    try:
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return

    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl   = table._tbl
    tblPr = tbl.tblPr

    for child in tblPr.findall(qn("w:tblW")):
        tblPr.remove(child)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")
    tblPr.append(tblW)

    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "autofit")

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcW  = tcPr.find(qn("w:tcW"))
            if tcW is not None:
                tcPr.remove(tcW)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_table_borders(table):
    """Aplica todas as bordas (linha simples preta) à tabela."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return

    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:%s" % side)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _format_vertex_table(table):
    """Aplica fonte Arial 7pt, compactação e alinhamento às células."""
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return

    CONFRONTACAO_COL = 7

    for row_i, row in enumerate(table.rows):
        for col_i, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(0)
                para.paragraph_format.line_spacing  = 1
                if col_i == CONFRONTACAO_COL:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7)
                    if row_i == 0:
                        run.bold = True


def _insert_vertex_table_replacing_paragraph(doc, paragraph, segments):
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError("Biblioteca python-docx não encontrada.") from exc

    HEADERS = ["Código", "Longitude", "Latitude", "Alt. (m)", "Vante", "Azimute", "Dist. (m)", "Confrontação"]
    KEYS    = ["codigo", "longitude", "latitude", "altitude", "vante", "azimute", "dist_m", "confrontacao"]

    table = doc.add_table(rows=1 + len(segments), cols=len(HEADERS))

    # Preenche células (texto simples; formatação aplicada depois)
    for col_i, header in enumerate(HEADERS):
        table.rows[0].cells[col_i].text = header

    for row_i, segment in enumerate(segments, 1):
        for col_i, key in enumerate(KEYS):
            table.rows[row_i].cells[col_i].text = str(segment.get(key, "") or "")

    # Aplica formatação (autofit, bordas, fonte, alinhamento)
    _set_table_autofit_to_window(table)
    _add_table_borders(table)
    _format_vertex_table(table)

    # Move a tabela para a posição do marcador e remove o parágrafo
    p_elem   = paragraph._p
    tbl_elem = table._tbl
    tbl_elem.getparent().remove(tbl_elem)
    p_elem.addprevious(tbl_elem)
    p_elem.getparent().remove(p_elem)
