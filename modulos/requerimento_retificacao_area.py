import os
import re
import shutil
import copy
import unicodedata
from datetime import datetime

from docx.shared import Pt

from ..docx_utils import (
    ensure_docx_path,
    substituir_marcadores_docx,
    normalize_key,
    all_document_paragraphs,
    PLACEHOLDER_PATTERN,
    clear_paragraph_content,
    marcar_blocos_assinatura_tecnica,
)
from .pessoas_utils import (
    enriquecer_proprietarios,
    qualificar_proprietarios_segs,
    montar_assinaturas_pessoas,
    buscar_proprietarios_servico,
    buscar_dados_servico,
    limpar_valor,
)
from .declaracao_respeito_limites import (
    _segmentos_com_formato,
    _adicionar_run_formatado,
)

_ARIAL = "Arial"
_PT12  = Pt(12)


# ── Template ───────────────────────────────────────────────────────────────────

def find_requerimento_retificacao_area_template(plugin_dir):
    path = os.path.join(plugin_dir, "models", "requerimento_retificacao_area.docx")
    return path if os.path.exists(path) else None


# ── Sanitização para nome de arquivo ──────────────────────────────────────────

def sanitizar_nome_denominacao(nome):
    """Remove acentos, converte para maiúsculas e troca espaços/símbolos por _."""
    nome = unicodedata.normalize("NFKD", str(nome or ""))
    nome = "".join(c for c in nome if not unicodedata.combining(c))
    nome = re.sub(r"[^\w\s-]", "", nome).strip()
    nome = re.sub(r"[\s-]+", "_", nome)
    nome = re.sub(r"_+", "_", nome)
    return nome[:80].upper()


# ── Helpers de leitura da camada QGIS ─────────────────────────────────────────

def _field_lookup(layer):
    return {normalize_key(f.name()): f.name() for f in layer.fields()}


def _fval(layer, feature, key):
    """Lê campo da feature QGIS com segurança. Retorna string ou ''."""
    if layer is None or feature is None:
        return ""
    fname = _field_lookup(layer).get(normalize_key(key))
    if not fname:
        return ""
    try:
        v = feature[fname]
        s = str(v).strip() if v is not None else ""
        return "" if s.upper() in ("NULL", "NONE") else s
    except Exception:
        return ""


# ── Formatação de nome próprio / endereço ─────────────────────────────────────

_PARTICULAS = frozenset({"de", "da", "do", "das", "dos", "e"})


def _formatar_nome_proprio(texto):
    """Primeira letra maiúscula, partículas em minúsculo, S/N sempre maiúsculo."""
    if not texto:
        return ""
    palavras = str(texto).lower().split()
    resultado = []
    for i, palavra in enumerate(palavras):
        raiz = palavra.rstrip(",.;:")
        if raiz in ("s/n", "sn"):
            resultado.append(palavra.upper())
        elif i == 0 or raiz not in _PARTICULAS:
            resultado.append(palavra.capitalize())
        else:
            resultado.append(palavra)
    return " ".join(resultado)


def _formatar_endereco_cartorio(texto):
    """Nome próprio + sigla de UF após barra em maiúsculo + CEP em maiúsculo.

    Exemplos:
      RUA EUCLIDES ONOFRE, 126 - CENTRO - ASSARÉ/CE - CEP: 63140-000
      → Rua Euclides Onofre, 126 - Centro - Assaré/CE - CEP: 63140-000
    """
    if not texto:
        return ""
    resultado = _formatar_nome_proprio(texto)
    # /UF (exatamente 2 letras após barra) → maiúsculo
    resultado = re.sub(
        r"/([a-zA-Z]{2})\b",
        lambda m: "/" + m.group(1).upper(),
        resultado,
    )
    # CEP (qualquer caixa) → maiúsculo
    resultado = re.sub(r"\bcep\b", "CEP", resultado, flags=re.IGNORECASE)
    return resultado


# ── Inferência de gênero pelo primeiro nome ───────────────────────────────────

_NOMES_FEMININOS = frozenset({
    "ana", "maria", "francisca", "antonia", "josefa", "raimunda",
    "rosangela", "lucia", "luciana", "luiza", "luzia", "marcia",
    "marcela", "marilia", "marta", "monica", "patricia", "paula",
    "paulina", "renata", "regina", "rita", "roberta", "rosa", "rosana",
    "rosiane", "rosilene", "rosimeire", "rosilane", "rosinete", "sandra",
    "silvia", "simone", "sonia", "sueli", "suely", "tereza", "teresa",
    "tais", "tatiana", "valeria", "vanessa", "veronica", "viviane",
    "adriana", "aline", "amanda", "andrea", "angela", "aparecida",
    "beatriz", "bruna", "camila", "carla", "carolina", "cassia",
    "catarina", "cecilia", "cilene", "claudia", "cristiane", "cristina",
    "daniela", "debora", "denise", "edna", "elaine", "eliane",
    "elisangela", "elizabeth", "eliza", "eloisa", "emanuela", "fabiana",
    "fatima", "fernanda", "flavia", "gabriela", "geovana", "giovana",
    "gisele", "graciela", "helena", "ilda", "ines", "iracema", "isabel",
    "isabela", "isadora", "janaina", "jessica", "joana", "joelma",
    "julia", "juliana", "karina", "katia", "larissa", "laura", "leticia",
    "lidiane", "lilian", "liliane", "livia", "lorena", "lourdes", "magda",
    "manuela", "mariana", "marina", "michele", "michelle", "natalia",
    "neide", "neusa", "neuza", "nilza", "olivia", "raquel", "rebecca",
    "rebeca", "sabrina", "sara", "sarah", "solange", "stefany", "stephany",
    "suzana", "taina", "tania", "thais", "valquiria", "vera", "vitoria",
    "yara", "yasmin",
})

_NOMES_MASCULINOS = frozenset({
    "joao", "jose", "antonio", "francisco", "raimundo", "carlos", "paulo",
    "pedro", "lucas", "luiz", "luis", "marcos", "marco", "marcelo",
    "bruno", "mateus", "matheus", "gabriel", "rafael", "daniel", "davi",
    "david", "felipe", "filipe", "fernando", "andre", "alexandre", "alex",
    "alberto", "albert", "adriano", "alan", "alano", "alessandro",
    "alfredo", "alisson", "alysson", "amilton", "anderson", "angelo",
    "anselmo", "ari", "ariel", "arnaldo", "arthur", "artur", "augusto",
    "benedito", "benicio", "bernardo", "caio", "caique", "cesar", "cicero",
    "claudio", "cleber", "cristiano", "danilo", "diego", "diogo",
    "douglas", "edilson", "edson", "eduardo", "elias", "emerson",
    "emiliano", "emanuel", "enrique", "henrique", "eron", "everton",
    "fabio", "fabricio", "flavio", "george", "gilberto", "gilmar",
    "gustavo", "heitor", "helio", "igor", "israel", "ivan", "ivo",
    "jacob", "jaime", "jair", "jefferson", "jeronimo", "joaquim", "jonas",
    "jorge", "julio", "junior", "leandro", "leonardo", "levi", "luciano",
    "manoel", "manuel", "marciano", "mario", "mauricio", "max", "miguel",
    "moises", "natan", "nathan", "nelson", "nilson", "osmar", "osvaldo",
    "otavio", "patricio", "raul", "renan", "renato", "ricardo", "roberto",
    "rodrigo", "rogerio", "romario", "ronaldo", "rubens", "samuel",
    "sergio", "silas", "tiago", "thiago", "tulio", "valdir", "valmir",
    "vanderlei", "vanderley", "vicente", "victor", "vitor", "vinicius",
    "wagner", "walter", "wanderson", "wilson", "wellington", "wesley",
})


def inferir_genero_primeiro_nome(nome_completo):
    """Infere gênero pelo primeiro nome. Retorna 'FEMININO', 'MASCULINO' ou 'DESCONHECIDO'."""
    if not nome_completo:
        return "DESCONHECIDO"
    normalizado = unicodedata.normalize("NFD", str(nome_completo))
    normalizado = "".join(c for c in normalizado if unicodedata.category(c) != "Mn")
    normalizado = normalizado.strip().lower()
    partes = normalizado.split()
    if not partes:
        return "DESCONHECIDO"
    primeiro = partes[0]
    if primeiro in _NOMES_FEMININOS:
        return "FEMININO"
    if primeiro in _NOMES_MASCULINOS:
        return "MASCULINO"
    return "DESCONHECIDO"


def _ajustar_tratamento_oficial(document, oficial_cartorio_raw):
    """Substitui 'Ao Ilmo. Sr.' por 'À Ilma. Sra.' no parágrafo do cartório quando feminino."""
    if inferir_genero_primeiro_nome(oficial_cartorio_raw) != "FEMININO":
        return

    _MASC = "Ao Ilmo. Sr."
    _FEM  = "À Ilma. Sra."  # "À Ilma. Sra."

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        if _MASC not in full_text:
            continue

        # Tenta trocar dentro de um único run
        for run in runs:
            if _MASC in run.text:
                run.text = run.text.replace(_MASC, _FEM, 1)
                return

        # Texto dividido entre runs: reconstrói o parágrafo
        idx = full_text.index(_MASC)
        segs_antes  = _segmentos_com_formato(runs, 0, idx)
        segs_depois = _segmentos_com_formato(runs, idx + len(_MASC), len(full_text))
        clear_paragraph_content(paragraph)
        for texto, rPr in segs_antes:
            _adicionar_run_formatado(paragraph, texto, rPr)
        r = paragraph.add_run(_FEM)
        r.font.name = _ARIAL
        r.font.size = _PT12
        for texto, rPr in segs_depois:
            _adicionar_run_formatado(paragraph, texto, rPr)
        return


# ── Conversão de área para extenso ────────────────────────────────────────────

_UNIDADES = [
    "", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove",
    "dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete",
    "dezoito", "dezenove",
]
_DEZENAS = [
    "", "", "vinte", "trinta", "quarenta", "cinquenta",
    "sessenta", "setenta", "oitenta", "noventa",
]
_CENTENAS = [
    "", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos",
    "seiscentos", "setecentos", "oitocentos", "novecentos",
]


def _extenso_ate_99(n):
    if n < 20:
        return _UNIDADES[n]
    dez, uni = n // 10, n % 10
    if uni == 0:
        return _DEZENAS[dez]
    return _DEZENAS[dez] + " e " + _UNIDADES[uni]


def _extenso_ate_999(n):
    if n == 0:
        return ""
    if n == 100:
        return "cem"
    if n < 100:
        return _extenso_ate_99(n)
    cent = n // 100
    resto = n % 100
    if resto == 0:
        return _CENTENAS[cent]
    return _CENTENAS[cent] + " e " + _extenso_ate_99(resto)


def _numero_para_extenso(n):
    """Converte inteiro 0-9999 para extenso em português."""
    if n == 0:
        return "zero"
    if n < 1000:
        return _extenso_ate_999(n)
    mil = n // 1000
    resto = n % 1000
    mil_str = "mil" if mil == 1 else _extenso_ate_999(mil) + " mil"
    if resto == 0:
        return mil_str
    # "e" antes do restante quando < 100 (ex.: "mil e cinco"); sem "e" quando >= 100
    conector = " e " if resto < 100 else " "
    return mil_str + conector + _extenso_ate_999(resto)


def _area_para_extenso(area_str):
    """Converte área em hectares (string) para texto por extenso.

    A parte inteira = hectares.
    1ª e 2ª casas decimais = ares.
    3ª e 4ª casas decimais = centiares.
    """
    area_str = str(area_str or "").strip()
    if not area_str:
        return "[ÁREA POR EXTENSO NÃO INFORMADA]"

    # Remove sufixo "ha" se presente (ex.: "5,1214 ha" → "5,1214")
    area_str = re.sub(r"\s*ha\s*$", "", area_str, flags=re.IGNORECASE).strip()
    if not area_str:
        return "[ÁREA POR EXTENSO NÃO INFORMADA]"

    area_str = area_str.replace(",", ".")
    try:
        if "." in area_str:
            int_str, dec_str = area_str.split(".", 1)
        else:
            int_str, dec_str = area_str, ""
        hectares = int(int_str) if int_str else 0
        dec_str  = dec_str.ljust(4, "0")[:4]
        ares      = int(dec_str[:2])
        centiares = int(dec_str[2:4])
    except (ValueError, TypeError):
        return "[ÁREA POR EXTENSO NÃO INFORMADA]"

    ha_label = "hectare" if hectares == 1 else "hectares"
    partes_principal = "%s %s" % (_numero_para_extenso(hectares), ha_label)

    if ares == 0 and centiares == 0:
        return partes_principal

    partes_decimal = []
    if ares > 0:
        are_label = "are" if ares == 1 else "ares"
        partes_decimal.append("%s %s" % (_numero_para_extenso(ares), are_label))
    if centiares > 0:
        ca_label = "centiare" if centiares == 1 else "centiares"
        partes_decimal.append("%s %s" % (_numero_para_extenso(centiares), ca_label))

    return "%s, %s" % (partes_principal, " e ".join(partes_decimal))


# ── Substituição do BLOCO_DECLARANTES ─────────────────────────────────────────

def _substituir_bloco_declarantes(document, declarantes, avisos_out):
    """Substitui {{BLOCO_DECLARANTES}} com qualificação jurídica completa."""
    if not declarantes:
        avisos_out.append(
            "Aviso: Nenhum proprietário encontrado. "
            "{{BLOCO_DECLARANTES}} não foi substituído."
        )
        return

    todos_segs = qualificar_proprietarios_segs(declarantes, avisos_out)

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        matches = [
            m for m in PLACEHOLDER_PATTERN.finditer(full_text)
            if normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_declarantes"
        ]
        if not matches:
            continue
        match = matches[0]
        segs_antes  = _segmentos_com_formato(runs, 0, match.start())
        segs_depois = _segmentos_com_formato(runs, match.end(), len(full_text))
        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = True
        for texto, rPr_clone in segs_antes:
            _adicionar_run_formatado(paragraph, texto, rPr_clone)
        for texto, bold in todos_segs:
            if texto:
                r = paragraph.add_run(texto)
                r.bold = bold
                r.font.name = _ARIAL
                r.font.size = _PT12
        for texto, rPr_clone in segs_depois:
            _adicionar_run_formatado(paragraph, texto, rPr_clone)


# ── Substituição do BLOCO_ASSINATURAS_PROPRIETARIOS ───────────────────────────

def _substituir_bloco_assinaturas(document, declarantes):
    """Substitui {{BLOCO_ASSINATURAS_PROPRIETARIOS}} com assinaturas centralizadas."""
    assinaturas = montar_assinaturas_pessoas(declarantes)
    if not assinaturas:
        return

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        if not any(
            normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_assinaturas_proprietarios"
            for m in PLACEHOLDER_PATTERN.finditer(full_text)
        ):
            continue

        clear_paragraph_content(paragraph)
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
        paragraph.paragraph_format.keep_together = True

        for i, (linha1, linha2) in enumerate(assinaturas):
            if i:
                for _ in range(4):
                    paragraph.add_run().add_break()
            sig = paragraph.add_run("________________________________________")
            sig.font.name = _ARIAL
            sig.font.size = _PT12
            paragraph.add_run().add_break()
            r = paragraph.add_run(linha1)
            r.bold = True
            r.font.name = _ARIAL
            r.font.size = _PT12
            if linha2:
                for parte in linha2.split("\n"):
                    paragraph.add_run().add_break()
                    pr = paragraph.add_run(parte)
                    pr.font.name = _ARIAL
                    pr.font.size = _PT12


# ── Substituição da PRIMEIRA ocorrência de {{ESCREVER_EXTENSO}} ──────────────

def _substituir_primeiro_extenso(document, valor):
    """Substitui SOMENTE a primeira ocorrência de {{ESCREVER_EXTENSO}} no documento.

    O modelo usa {{ESCREVER_EXTENSO}} duas vezes no mesmo parágrafo:
    1ª → extenso da área da matrícula (substituído aqui)
    2ª → extenso da área real/certificada (deixado para substituir_marcadores_docx)
    """
    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        matches = [
            m for m in PLACEHOLDER_PATTERN.finditer(full_text)
            if normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "escrever_extenso"
        ]
        if not matches:
            continue
        match = matches[0]
        segs_antes  = _segmentos_com_formato(runs, 0, match.start())
        segs_depois = _segmentos_com_formato(runs, match.end(), len(full_text))
        clear_paragraph_content(paragraph)
        for texto, rPr in segs_antes:
            _adicionar_run_formatado(paragraph, texto, rPr)
        r = paragraph.add_run(valor)
        r.font.name = _ARIAL
        r.font.size = _PT12
        for texto, rPr in segs_depois:
            _adicionar_run_formatado(paragraph, texto, rPr)
        return  # para após a primeira substituição


# ── Função principal ───────────────────────────────────────────────────────────

def fill_requerimento_retificacao_area_template(
    template_path, output_path, data, layer=None, feature=None
):
    """Gera o Requerimento de Retificação de Área.

    Args:
        data:    dict com dados do PDF e QGIS (vindo de _collect_memorial_data).
        layer:   camada QGIS Serviços 2 (para campos extras não capturados pelo PDF).
        feature: feature QGIS do serviço atual (para folha, livro, ret_aver, etc.).
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Biblioteca python-docx não encontrada. Instale python-docx no Python do QGIS."
        ) from exc

    template_path = os.path.abspath(template_path)
    output_path   = os.path.abspath(ensure_docx_path(output_path))
    if template_path == output_path:
        raise RuntimeError("Escolha um arquivo de saída diferente do modelo original.")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copyfile(template_path, output_path)

    document  = Document(output_path)
    req_data  = copy.copy(data)
    avisos    = []

    servico_id = req_data.get("__servico_id")

    # ── Dados do serviço via DB (fonte primária para campos extras) ──────────
    # buscar_dados_servico faz SELECT * da tabela servicos.servicos, garantindo
    # que folha, livro, ret_aver, oficial_cartorio, area_matricula, etc. venham
    # independentemente do que a camada QGIS expõe.
    servico_data = {}
    if layer and servico_id:
        try:
            servico_data = buscar_dados_servico(layer, servico_id)
        except Exception:
            pass

    def _sd(key):
        """Busca campo no DB primeiro, fallback na feature QGIS."""
        v = limpar_valor(servico_data.get(key, ""))
        if v and v.upper() not in ("NULL", "NONE"):
            return v
        return _fval(layer, feature, key)

    # ── Proprietários: servico_id → busca direta no DB ───────────────────────
    owners = req_data.get("__owners", [])
    if not owners and layer and servico_id:
        try:
            owners = buscar_proprietarios_servico(layer, servico_id)
        except Exception:
            pass
    declarantes = enriquecer_proprietarios(owners, layer)

    # ── Campos da camada Serviços 2 ──────────────────────────────────────────
    folha_raw             = _sd("folha")
    livro_raw             = _sd("livro")
    ret_aver_raw          = _sd("ret_aver")
    oficial_cartorio_raw  = _sd("oficial_cartorio")
    endereco_cartorio_raw = _sd("endereco_cartorio")
    area_matricula_raw    = _sd("area_matricula")

    # ── Área por extenso: duas versões distintas ─────────────────────────────
    # 1ª ocorrência de {{ESCREVER_EXTENSO}} → extenso da área da matrícula
    # 2ª ocorrência de {{ESCREVER_EXTENSO}} → extenso da área real/certificada (PDF)
    area_matricula_extenso = _area_para_extenso(area_matricula_raw)
    area_real_raw          = req_data.get("area") or req_data.get("area_ha", "")
    area_real_extenso      = _area_para_extenso(area_real_raw)

    # ── Montagem do dict de substituição ────────────────────────────────────
    req_data["folha"]             = folha_raw   or "-"
    req_data["livro"]             = livro_raw   or "-"
    req_data["ret_aver"]          = ret_aver_raw or "-"
    req_data["ano_atual"]         = str(datetime.now().year)
    req_data["oficial_cartorio"]  = (
        _formatar_nome_proprio(oficial_cartorio_raw)
        if oficial_cartorio_raw
        else "[OFICIAL DO CARTÓRIO NÃO INFORMADO]"
    )
    req_data["endereco_cartorio"] = (
        _formatar_endereco_cartorio(endereco_cartorio_raw)
        if endereco_cartorio_raw
        else "[ENDEREÇO DO CARTÓRIO NÃO INFORMADO]"
    )
    req_data["area_matricula"]    = area_matricula_raw or "[ÁREA DA MATRÍCULA NÃO INFORMADA]"

    # nome_cartorio / cartorio: prioridade DB, depois data existente, formatado
    _nc_raw = _sd("nome_cartorio") or _sd("cartorio")
    if not _nc_raw:
        _nc_raw = req_data.get("nome_cartorio") or req_data.get("cartorio", "")
    if _nc_raw and _nc_raw not in ("[CARTÓRIO NÃO INFORMADO]",):
        _nc_fmt = _formatar_nome_proprio(_nc_raw)
        req_data["nome_cartorio"] = _nc_fmt
        req_data["cartorio"]      = _nc_fmt

    # 2ª ocorrência de {{ESCREVER_EXTENSO}} → área real; será resolvida pelo
    # substituir_marcadores_docx após a primeira ter sido tratada abaixo
    req_data["escrever_extenso"]  = area_real_extenso

    # ── Substituições especializadas (antes de substituir_marcadores_docx) ──
    _substituir_bloco_declarantes(document, declarantes, avisos)
    _substituir_bloco_assinaturas(document, declarantes)

    # Ajusta tratamento (Ao Ilmo. Sr. / À Ilma. Sra.) conforme gênero inferido
    _ajustar_tratamento_oficial(document, oficial_cartorio_raw)

    # 1ª ocorrência de {{ESCREVER_EXTENSO}} → extenso da matrícula (antes do marcadores_docx)
    _substituir_primeiro_extenso(document, area_matricula_extenso)

    # ── Demais placeholders simples ──────────────────────────────────────────
    substituir_marcadores_docx(document, req_data)

    # ── Mantém bloco do RT junto na mesma página ────────────────────────────
    paragraphs = list(all_document_paragraphs(document))
    marcar_blocos_assinatura_tecnica(paragraphs)

    document.save(output_path)
    return avisos
