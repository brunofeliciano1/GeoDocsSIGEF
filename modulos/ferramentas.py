import os
import re
from collections import OrderedDict

from qgis.PyQt.QtWidgets import (
    QFileDialog,
    QFormLayout,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)


class DropFileLineEdit(QLineEdit):

    def __init__(self, extensions, parent=None):
        super().__init__(parent)
        self.extensions = tuple(extension.lower() for extension in extensions)
        self.setAcceptDrops(True)
        self.setPlaceholderText("Arraste o arquivo aqui ou clique em Procurar")

    def dragEnterEvent(self, event):
        if self._accepted_event(event):
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if self._accepted_event(event):
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event):
        urls = event.mimeData().urls()
        if not urls:
            return
        path = urls[0].toLocalFile()
        if self._accepted_path(path):
            self.setText(path)
            event.acceptProposedAction()
        else:
            event.ignore()

    def _accepted_event(self, event):
        urls = event.mimeData().urls()
        return bool(urls and self._accepted_path(urls[0].toLocalFile()))

    def _accepted_path(self, path):
        return bool(path and os.path.splitext(path)[1].lower() in self.extensions)


def default_output_directory():
    downloads = os.path.join(os.path.expanduser("~"), "Downloads")
    return downloads if os.path.isdir(downloads) else os.path.expanduser("~")


def build_tools_tab(dialog):
    tab = QWidget()
    layout = QVBoxLayout(tab)

    intro = QLabel(
        "Ferramenta auxiliar baseada no script antigo do SIGEF. "
        "Ela lê o PDF, extrai a tabela de vértices e gera MEMORIAL_FERRAMENTA.docx, "
        "PLANILHA.xlsx e CARTAS_DE_ANUENCIAS.docx na pasta escolhida, sem gerar ZIP. "
        "Esta aba é independente e não altera a lógica da aba Memorial."
    )
    intro.setObjectName("introLabel")
    intro.setWordWrap(True)
    layout.addWidget(intro)

    form = QFormLayout()
    dialog.tools_pdf_path = DropFileLineEdit((".pdf",))
    dialog.tools_output_path = QLineEdit()
    dialog.tools_output_path.setText(default_output_directory())

    form.addRow("PDF SIGEF:", _file_picker_row(
        dialog,
        dialog.tools_pdf_path,
        "Selecionar PDF SIGEF",
        "Arquivos PDF (*.pdf)"
    ))
    form.addRow("Pasta de saída:", _tools_output_picker_row(dialog))
    layout.addLayout(form)

    buttons = QHBoxLayout()
    dialog.tools_generate_button = QPushButton("Gerar documentos da ferramenta")
    dialog.tools_generate_button.setObjectName("primaryButton")
    dialog.tools_generate_button.clicked.connect(lambda: _generate_tools_documents(dialog))
    buttons.addStretch()
    buttons.addWidget(dialog.tools_generate_button)
    layout.addLayout(buttons)

    label = QLabel("Log da Ferramenta")
    label.setObjectName("sectionLabel")
    layout.addWidget(label)

    dialog.tools_log = QTextEdit()
    dialog.tools_log.setReadOnly(True)
    dialog.tools_log.setPlaceholderText("O resultado da ferramenta aparece aqui.")
    layout.addWidget(dialog.tools_log, 1)
    return tab


def _file_picker_row(dialog, line_edit, title, file_filter):
    row = QWidget()
    layout = QHBoxLayout(row)
    layout.setContentsMargins(0, 0, 0, 0)
    button = QPushButton("Procurar")
    button.clicked.connect(lambda: _select_input_file(dialog, line_edit, title, file_filter))
    layout.addWidget(line_edit)
    layout.addWidget(button)
    return row


def _select_input_file(dialog, line_edit, title, file_filter):
    path, _ = QFileDialog.getOpenFileName(dialog, title, "", file_filter)
    if path:
        line_edit.setText(path)


def _tools_output_picker_row(dialog):
    row = QWidget()
    layout = QHBoxLayout(row)
    layout.setContentsMargins(0, 0, 0, 0)
    button = QPushButton("Selecionar pasta")
    button.clicked.connect(lambda: _select_tools_output_folder(dialog))
    layout.addWidget(dialog.tools_output_path)
    layout.addWidget(button)
    return row


def _select_tools_output_folder(dialog):
    path = QFileDialog.getExistingDirectory(
        dialog,
        "Selecionar pasta de saída",
        dialog.tools_output_path.text().strip() or default_output_directory()
    )
    if path:
        dialog.tools_output_path.setText(path)


def _generate_tools_documents(dialog):
    try:
        pdf_path = dialog.tools_pdf_path.text().strip()
        output_folder = dialog.tools_output_path.text().strip() or default_output_directory()
        if not pdf_path or not os.path.exists(pdf_path):
            raise RuntimeError("Selecione um PDF SIGEF válido na aba Ferramentas.")
        if not os.path.isdir(output_folder):
            os.makedirs(output_folder, exist_ok=True)

        result = gerar_documentos_ferramenta_sigef(pdf_path, output_folder)
        dialog.tools_log.setPlainText(
            "Ferramenta finalizada com sucesso.\n\n"
            "Vértices encontrados: {vertices}\n"
            "Memorial: {memorial}\n"
            "Planilha: {planilha}\n"
            "Anuências: {anuencias}".format(**result)
        )
        QMessageBox.information(dialog, "GeoDocs SIGEF", "Documentos da ferramenta gerados com sucesso.")
    except Exception as exc:
        dialog.tools_log.setPlainText(str(exc))
        QMessageBox.warning(dialog, "GeoDocs SIGEF", str(exc))


# =========================================================
# FERRAMENTAS - SCRIPT ANTIGO SIGEF ADAPTADO
# Esta seção é independente da aba Memorial.
# =========================================================

def gerar_documentos_ferramenta_sigef(pdf_path, output_folder):
    texto_completo, vertices, candidatas = ferramenta_ler_pdf(pdf_path)
    if not vertices:
        exemplos = "\n".join(candidatas[:5])
        raise RuntimeError(
            "Não foi possível localizar a tabela de vértices no PDF.\n"
            + ("Exemplos de linhas candidatas:\n%s" % exemplos if exemplos else "")
        )

    cabecalho = ferramenta_extrair_cabecalho(texto_completo)

    memorial = os.path.join(output_folder, "MEMORIAL_FERRAMENTA.docx")
    planilha = os.path.join(output_folder, "PLANILHA.xlsx")
    anuencias = os.path.join(output_folder, "CARTAS_DE_ANUENCIAS.docx")
    ferramenta_gerar_memorial_narrativo(memorial, cabecalho, vertices)
    ferramenta_gerar_planilha(planilha, vertices)
    ferramenta_gerar_anuencias(anuencias, cabecalho, vertices)


    return {
        "vertices": len(vertices),
        "memorial": memorial,
        "planilha": planilha,
        "anuencias": anuencias,
    }


def ferramenta_ler_pdf(pdf_path):
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError(
            "Biblioteca 'pdfplumber' não encontrada. Instale no Python do QGIS com: python -m pip install pdfplumber"
        ) from exc

    texto_paginas = []
    vertices = []
    linhas_candidatas = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            texto = page.extract_text() or ""
            texto_paginas.append(texto)
            for linha in texto.splitlines():
                linha = " ".join(linha.strip().split())
                if not linha or ferramenta_ignorar_linha(linha):
                    continue
                if "°" in linha and '"' in linha:
                    linhas_candidatas.append(linha)
                registro = ferramenta_parse_linha_vertice(linha)
                if registro:
                    vertices.append(registro)

    vertices_unicos = []
    vistos = set()
    for item in vertices:
        chave = (
            item["codigo"], item["longitude"], item["latitude"], item["altitude"],
            item["vante"], item["azimute"], item["dist"], item["confrontacao"]
        )
        if chave not in vistos:
            vistos.add(chave)
            vertices_unicos.append(item)

    return "\n".join(texto_paginas), vertices_unicos, linhas_candidatas


def ferramenta_ignorar_linha(linha):
    ignorar = [
        "MINISTÉRIO DA AGRICULTURA",
        "INSTITUTO NACIONAL DE COLONIZAÇÃO",
        "MEMORIAL DESCRITIVO",
        "Página ",
        "CERTIFICAÇÃO:",
        "Data Certificação:",
        "Data da Geração:",
        "DESCRIÇÃO DA PARCELA",
        "VÉRTICE SEGMENTO VANTE",
        "Código Longitude Latitude Altitude",
        "Cartório (CNS):",
        "Este Memorial Descritivo foi gerado automaticamente",
        "Em atendimento ao § 5° do art. 176",
        "A autenticidade deste documento pode ser verificada",
        "Parcela certificada pelo SIGEF",
        "Certificada - Sem Confirmação de Registro em Cartório",
        "nenhuma outra poligonal constante do cadastro georreferenciado do INCRA",
    ]
    return any(item in linha for item in ignorar)


def ferramenta_eh_codigo_vertice(token):
    if len(token) < 3:
        return False
    if "°" in token or "," in token or '"' in token:
        return False
    return bool(re.match(r"^[A-Z0-9][A-Z0-9\-]*$", token, re.IGNORECASE))


def ferramenta_eh_coord_dms(token):
    return bool(re.match(r"^-?\d+°\d+'\d+[.,]\d+\"$", token))


def ferramenta_eh_numero_br(token):
    return bool(re.match(r"^-?\d+(?:[.,]\d+)?$", token))


def ferramenta_eh_azimute(token):
    return bool(re.match(r"^\d+°\d+'$", token))


def ferramenta_parse_linha_vertice(linha):
    partes = linha.split()
    if len(partes) < 8:
        return None
    try:
        codigo = partes[0]
        longitude = partes[1]
        latitude = partes[2]
        altitude = partes[3]
        vante = partes[4]
        azimute = partes[5]
        dist = partes[6]
    except IndexError:
        return None

    if not ferramenta_eh_codigo_vertice(codigo):
        return None
    if not ferramenta_eh_coord_dms(longitude):
        return None
    if not ferramenta_eh_coord_dms(latitude):
        return None
    if not ferramenta_eh_numero_br(altitude):
        return None
    if not ferramenta_eh_codigo_vertice(vante):
        return None
    if not ferramenta_eh_azimute(azimute):
        return None
    if not ferramenta_eh_numero_br(dist):
        return None

    confrontacao = " ".join(partes[7:]).strip()
    if not confrontacao:
        return None

    return {
        "codigo": codigo,
        "longitude": longitude,
        "latitude": latitude,
        "altitude": altitude,
        "vante": vante,
        "azimute": azimute,
        "dist": dist,
        "confrontacao": confrontacao,
    }


def ferramenta_extrair_cabecalho(texto):
    def pegar(padrao, default=""):
        m = re.search(padrao, texto, re.MULTILINE)
        return m.group(1).strip() if m else default

    sistema = pegar(r"Sistema Geodésico de referência:\s*(.+?)\s+Documento de RT:")
    if not sistema:
        sistema = pegar(r"Sistema Geodésico de referência:\s*(.+)")

    conselho = pegar(r"Conselho Profissional:\s*(.+?)(?:Código de credenciamento:|$)")
    cred = pegar(r"Código de credenciamento:\s*(.+?)(?:\s+Conselho Profissional:|$)")
    if not cred:
        cred = pegar(r"Código de credenciamento:\s*(.+)")

    area = pegar(r"Área \(Sistema Geodésico Local\):\s*([\d,\.]+)\s*ha")
    perimetro = pegar(r"Perímetro \(m\):\s*([\d\.,]+)\s*m")

    return {
        "denominacao": ferramenta_limpar_campo_cabecalho(pegar(r"Denominação:\s*(.+)"), "Natureza da Área"),
        "proprietario": pegar(r"Proprietário\(a\):\s*(.+)"),
        "matricula": ferramenta_limpar_campo_cabecalho(pegar(r"Matrícula do imóvel:\s*(.+)"), "Código INCRA/SNCR"),
        "municipio_uf": pegar(r"Município/UF:\s*(.+)"),
        "natureza_area": pegar(r"Natureza da Área:\s*(.+)"),
        "cpf_cnpj": pegar(r"CPF:\s*(.+)") or pegar(r"CNPJ:\s*(.+)"),
        "codigo_incra": pegar(r"Código INCRA/SNCR:\s*(.+)"),
        "area_ha": area,
        "perimetro_m": perimetro,
        "sistema": sistema,
        "responsavel_tecnico": pegar(r"Responsável Técnico\(a\):\s*(.+)"),
        "formacao": pegar(r"Formação:\s*(.+)"),
        "conselho_profissional": conselho,
        "credenciamento": cred,
    }


def ferramenta_limpar_campo_cabecalho(valor, proximo_rotulo):
    valor = " ".join(str(valor or "").split())
    if not valor:
        return ""
    valor = re.split(r"\s+%s\s*:" % re.escape(proximo_rotulo), valor, maxsplit=1, flags=re.IGNORECASE)[0]
    return valor.strip(" ,.;:-")


def ferramenta_normalizar_confrontacao(confrontacao):
    texto = " ".join(str(confrontacao or "").split())
    maiusculo = texto.upper()

    if "RIO " in maiusculo:
        m = re.search(r"(RIO\s+[^|]+?)(?:,|$)", texto, re.IGNORECASE)
        return m.group(1).upper().strip() if m else texto.upper()
    if "RIACHO" in maiusculo:
        return texto.upper()
    if "CÓRREGO" in maiusculo or "CORREGO" in maiusculo:
        return texto.upper()
    if "BREJO" in maiusculo:
        return texto.upper()
    if "ESTRADA" in maiusculo or "RODOVIA" in maiusculo or "CE-060" in maiusculo:
        m = re.search(r"(ESTRADA[^|]*|RODOVIA[^|]*|CE-060)", texto, re.IGNORECASE)
        return m.group(1).upper().strip() if m else texto.upper()
    if " de " in texto:
        nome = texto.split(" de ", 1)[1].strip(" .")
        return "propriedade de %s" % nome
    if "|" in texto:
        partes = [p.strip(" .") for p in texto.split("|") if p.strip()]
        if partes:
            nome = partes[-1]
            return "propriedade de %s" % nome
    return "propriedade de %s" % texto


def ferramenta_gerar_memorial_narrativo(out_path, cabecalho, vertices):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Biblioteca python-docx não encontrada. Instale python-docx no Python do QGIS.") from exc

    doc = Document()
    p = doc.add_paragraph()
    mapa_coords = {item["codigo"]: (item["latitude"], item["longitude"], item["altitude"]) for item in vertices}
    primeiro = vertices[0]

    p.add_run("LIMITES E CONFRONTAÇÕES: Inicia-se a descrição deste perímetro no vértice ")
    r = p.add_run(primeiro["codigo"]); r.bold = True
    p.add_run(" (Latitude ")
    r = p.add_run(primeiro["latitude"]); r.bold = True
    p.add_run(", Longitude ")
    r = p.add_run(primeiro["longitude"]); r.bold = True
    p.add_run(", Altitude ")
    r = p.add_run(primeiro["altitude"]); r.bold = True
    p.add_run("), georreferenciado no Sistema Geodésico Brasileiro, DATUM - ")
    r = p.add_run("SIRGAS2000"); r.bold = True
    p.add_run(", MC-39°W; ")

    for item in vertices:
        lat2, lon2, alt2 = mapa_coords.get(item["vante"], ("", "", ""))
        confrontante = ferramenta_normalizar_confrontacao(item["confrontacao"])
        p.add_run("deste segue confrontando com ")
        r = p.add_run(confrontante); r.bold = True
        p.add_run(", do ponto ")
        r = p.add_run(item["codigo"]); r.bold = True
        p.add_run(" (Latitude ")
        r = p.add_run(item["latitude"]); r.bold = True
        p.add_run(", Longitude ")
        r = p.add_run(item["longitude"]); r.bold = True
        p.add_run(", Altitude ")
        r = p.add_run(item["altitude"]); r.bold = True
        p.add_run(") ao ponto ")
        r = p.add_run(item["vante"]); r.bold = True
        if lat2 and lon2 and alt2:
            p.add_run(" (Latitude ")
            r = p.add_run(lat2); r.bold = True
            p.add_run(", Longitude ")
            r = p.add_run(lon2); r.bold = True
            p.add_run(", Altitude ")
            r = p.add_run(alt2); r.bold = True
            p.add_run(")")
        p.add_run("; azimute ")
        r = p.add_run(item["azimute"]); r.bold = True
        p.add_run(" – distância de ")
        r = p.add_run('%s m' % item["dist"]); r.bold = True
        p.add_run("; ")

    p.add_run("chegando novamente ao ponto inicial, perfazendo um total de ")
    r = p.add_run('%s m' % cabecalho.get("perimetro_m", "")); r.bold = True
    p.add_run(".")
    doc.save(out_path)


def ferramenta_gerar_planilha(out_path, vertices):
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font
    except ImportError as exc:
        raise RuntimeError("Biblioteca openpyxl não encontrada. Instale openpyxl no Python do QGIS.") from exc

    wb = Workbook()
    ws = wb.active
    ws.title = "Vertices"
    cabecalhos = ["Código", "Longitude", "Latitude", "Altitude (m)", "Vante", "Azimute", "Dist. (m)", "Confrontação"]
    ws.append(cabecalhos)
    for col in ws[1]:
        col.font = Font(bold=True)
    for item in vertices:
        ws.append([item["codigo"], item["longitude"], item["latitude"], item["altitude"], item["vante"], item["azimute"], item["dist"], item["confrontacao"]])
    larguras = {"A": 20, "B": 18, "C": 18, "D": 14, "E": 20, "F": 12, "G": 12, "H": 60}
    for coluna, largura in larguras.items():
        ws.column_dimensions[coluna].width = largura
    wb.save(out_path)


def ferramenta_gerar_anuencias(out_path, cabecalho, vertices):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError("Biblioteca python-docx não encontrada. Instale python-docx no Python do QGIS.") from exc

    doc = Document()
    mapa_coords = {item["codigo"]: (item["latitude"], item["longitude"]) for item in vertices}
    agrupados = OrderedDict()
    for item in vertices:
        chave = ferramenta_normalizar_confrontacao(item["confrontacao"])
        agrupados.setdefault(chave, []).append(item)

    total = len(agrupados)
    atual = 0
    for confrontante, segmentos in agrupados.items():
        atual += 1
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = titulo.add_run("DECLARAÇÃO DE RECONHECIMENTO DE LIMITES")
        r.bold = True

        doc.add_paragraph(
            "Eu, %s, confrontante do imóvel de %s, CPF %s, matrícula %s, situado no município de %s, DECLARO, para os devidos fins, que reconheço como corretos e aceitos os limites comuns entre nossas propriedades, descritos a seguir:"
            % (confrontante, cabecalho.get("proprietario", ""), cabecalho.get("cpf_cnpj", ""), cabecalho.get("matricula", ""), cabecalho.get("municipio_uf", ""))
        )

        for item in segmentos:
            lat1, lon1 = mapa_coords.get(item["codigo"], ("", ""))
            lat2, lon2 = mapa_coords.get(item["vante"], ("", ""))
            p = doc.add_paragraph()
            p.add_run("do ponto ")
            r = p.add_run(item["codigo"]); r.bold = True
            p.add_run(" (Latitude: ")
            r = p.add_run(lat1); r.bold = True
            p.add_run(", Longitude: ")
            r = p.add_run(lon1); r.bold = True
            p.add_run(") ao ponto ")
            r = p.add_run(item["vante"]); r.bold = True
            p.add_run(" (Latitude: ")
            r = p.add_run(lat2); r.bold = True
            p.add_run(", Longitude: ")
            r = p.add_run(lon2); r.bold = True
            p.add_run("); azimute ")
            r = p.add_run(item["azimute"]); r.bold = True
            p.add_run(" – distância de ")
            r = p.add_run('%s m' % item["dist"]); r.bold = True
            p.add_run(";")

        doc.add_paragraph("Declaramos não existir qualquer disputa ou discordância sobre os limites acima descritos.")
        doc.add_paragraph("Várzea Alegre – CE, [DATA].")
        doc.add_paragraph("")
        doc.add_paragraph("Confrontante: _________________________________")
        doc.add_paragraph("Proprietário(a): ______________________________")
        doc.add_paragraph("Responsável Técnico: _________________________")
        if atual < total:
            doc.add_page_break()

    doc.save(out_path)
