import os
import re
import shutil
import copy
import unicodedata
from difflib import SequenceMatcher
from datetime import date

from ..docx_utils import (
    ensure_docx_path,
    substituir_marcadores_docx,
    normalize_key,
    all_document_paragraphs,
    PLACEHOLDER_PATTERN,
    clear_paragraph_content,
    marcar_blocos_assinatura_tecnica,
    formatar_trt_sem_duplicar_uf,
    formatar_municipio_para_data,
)

# Módulo central de pessoas — qualificação, busca, formatação
from .pessoas_utils import (
    limpar_valor                    as _limpar,
    flexionar                       as _flex,
    detectar_tipo_pessoa            as _detectar_tipo,
    buscar_vinculados               as _buscar_vinculados,
    buscar_vinculado                as _buscar_vinculado,
    qualificar_pessoa_fisica_segs   as _qualificar_pf_segs,
    qualificar_pessoa_juridica_segs as _qualificar_pj_segs,
    qualificar_proprietarios_segs,
    montar_assinaturas_pessoas      as _montar_assinaturas,
    buscar_dados_servico            as _buscar_dados_servico,
    buscar_proprietarios_servico    as _buscar_proprietarios_servico,
    buscar_proprietarios_confinante as _buscar_confinante_pessoas,
    enriquecer_confinante           as _enriquecer_confinante,
    enriquecer_proprietarios        as _enriquecer_proprietarios,
    expandir_com_conjuges           as _flatten_owners_com_conjuges,
    montar_bloco_proprietarios_simples as _build_proprietarios_inline,
)

# Helpers DOCX (preservação de formatação de runs)
from .declaracao_respeito_limites import (
    _segmentos_com_formato,
    _adicionar_run_formatado,
)


# ── Helpers locais ─────────────────────────────────────────────────────────────

def _field_lookup(layer):
    return {normalize_key(f.name()): f.name() for f in layer.fields()}


def _fval(layer, feature, key):
    fname = _field_lookup(layer).get(normalize_key(key))
    if not fname:
        return ""
    v = feature[fname]
    s = str(v).strip() if v is not None else ""
    return "" if s.upper() in ("NULL", "NONE") else s


# ── Template ───────────────────────────────────────────────────────────────────

def find_declaracao_confrontantes_template(plugin_dir):
    path = os.path.join(plugin_dir, "models", "declaracao_de_confrontantes.docx")
    return path if os.path.exists(path) else None


# ── Sanitização de nome de arquivo ────────────────────────────────────────────

def sanitizar_nome_arquivo(nome):
    """Remove acentos, converte para maiúsculas, troca espaços por _ e limpa símbolos."""
    nome = unicodedata.normalize("NFKD", str(nome or ""))
    nome = "".join(c for c in nome if not unicodedata.combining(c))
    nome = re.sub(r"[^\w\s-]", "", nome).strip()
    nome = re.sub(r"[\s-]+", "_", nome)
    nome = re.sub(r"_+", "_", nome)   # underlines duplicados
    return nome[:80].upper()


def nome_confinante_para_arquivo(layer, confinante_feature):
    """Retorna o nome base sanitizado para o arquivo da Declaração de Confrontante.

    Fonte de dados (em ordem de prioridade):
    1. Nome da primeira pessoa vinculada via confinante_pessoas → pessoas.
    2. Campo nome_propriedade da feature do confinante.
    3. Fallback: "CONFINANTE".

    Nunca usa servicos.confinante.nome.
    """
    # Tenta buscar pessoas vinculadas pelo id do confinante
    nome_base = ""
    try:
        cid = int(float(_fval(layer, confinante_feature, "id")))
    except (ValueError, TypeError):
        cid = None

    if cid and layer:
        try:
            pessoas = _buscar_confinante_pessoas(layer, cid)
            if pessoas:
                nome_base = str(pessoas[0].get("nome", "") or "").strip()
        except Exception:
            pass

    # Fallback: nome_propriedade
    if not nome_base:
        try:
            nome_base = _fval(layer, confinante_feature, "nome_propriedade").strip()
        except Exception:
            pass

    return sanitizar_nome_arquivo(nome_base) if nome_base else "CONFINANTE"


# ── Normalização para similaridade (remove palavras fracas) ───────────────────

_PALAVRAS_FRACAS = frozenset({
    "de", "da", "do", "dos", "das", "e", "em", "a", "o", "as", "os",
    "imovel", "rural", "confrontante", "proprietario", "area",
})


def normalizar_nome(texto):
    if not texto:
        return ""
    t = unicodedata.normalize("NFKD", str(texto).lower())
    t = "".join(c for c in t if not unicodedata.combining(c))
    t = re.sub(r"[^\w\s]", " ", t)
    return " ".join(w for w in t.split() if w and w not in _PALAVRAS_FRACAS)


def _similaridade(a, b):
    na, nb = normalizar_nome(a), normalizar_nome(b)
    if not na or not nb:
        return 0.0
    return SequenceMatcher(None, na, nb).ratio()


# ── Normalização para detecção de infraestrutura (mantém todas as palavras) ───

def _normalizar_infra(texto):
    """Normaliza sem remover palavras fracas — necessário para 'area de servidao publica'."""
    if not texto:
        return ""
    t = unicodedata.normalize("NFKD", str(texto).lower())
    t = "".join(c for c in t if not unicodedata.combining(c))
    t = re.sub(r"[^\w\s]", " ", t)   # hífen e pontuação → espaço (BR-402 → br 402)
    return " ".join(t.split())


# Padrões que caracterizam infraestrutura, área pública ou recurso natural
_RE_INFRA = re.compile(
    r"area\s+de\s+servidao"
    r"|servidao\s+publica"
    r"|servidao\s+administrativa"
    r"|area\s+publica"
    r"|faixa\s+de\s+dominio"
    r"|faixa\s+de\s+passagem"
    r"|\brodovia\b"
    r"|\bbr\s+\d"     # BR-402 → br 402
    r"|\bce\s+\d"     # CE-085 → ce 085
    r"|\bba\s+\d|\bma\s+\d|\bpi\s+\d|\bto\s+\d"
    r"|\bgo\s+\d|\bmg\s+\d|\bsp\s+\d|\bpr\s+\d|\bsc\s+\d"
    r"|\bestrada\b|\brua\b|\bavenida\b"
    r"|linha\s+de\s+transmissao"
    r"|linha\s+de\s+distribuicao"
    r"|linha\s+ferrea"
    r"|\bferrovia\b"
    r"|\briacho\b"
    r"|\brio\b"
    r"|\bcorrego\b"
    r"|\bacude\b|\blagoa\b|\bcanal\b|\blago\b",
)


def is_confrontante_infraestrutura_ou_recurso(nome, nome_propriedade=None, tipo_imovel=None):
    """Retorna True se o confrontante parece ser infraestrutura, área pública ou recurso natural."""
    for texto in (nome, nome_propriedade, tipo_imovel):
        if texto and _RE_INFRA.search(_normalizar_infra(texto)):
            return True
    return False


# ── Artigo preposicionado e sufixo para o nome do confrontante infra ──────────

_PREFIXOS_FEMININO = frozenset({
    "area", "rodovia", "estrada", "rua", "avenida",
    "linha", "lagoa", "faixa", "reserva", "ferrovia",
})
_PREFIXOS_MASCULINO = frozenset({
    "rio", "riacho", "corrego", "acude", "canal", "lago",
})

# Nomes hídricos masculinos — para gerar "o qual confronta com"
_MASCULINOS_HIDRO = frozenset({
    "rio", "riacho", "acude", "corrego", "canal", "lago", "ribeirao",
})


def _artigo_de(nome_infra):
    """Retorna 'da ', 'do ' ou 'de ' conforme a primeira palavra do nome."""
    if not nome_infra:
        return "de "
    primeiro = _normalizar_infra(nome_infra).split()[0] if nome_infra.strip() else ""
    if primeiro in _PREFIXOS_FEMININO:
        return "da "
    if primeiro in _PREFIXOS_MASCULINO:
        return "do "
    return "de "


def _sufixo_qual_confronta(nome_infra):
    """Retorna ', o qual confronta com' (hídrico) ou ', a qual confronta com'."""
    primeiro = _normalizar_infra(nome_infra or "").split()[0] if nome_infra else ""
    if primeiro in _MASCULINOS_HIDRO:
        return ", o qual confronta com"
    return ", a qual confronta com"


# ── Busca de representantes de entidade de infra ──────────────────────────────

# Prioridade diferente da geral: PROCURADOR antes de REPRESENTANTE_LEGAL
_TIPOS_REPRESENTACAO_INFRA = (
    "PROCURADOR", "REPRESENTANTE", "REPRESENTANTE_LEGAL", "RESPONSAVEL"
)


def _buscar_representantes_infra(declarantes, layer):
    """Busca PROCURADOR/REPRESENTANTE ignorando a classificação PF/PJ da entidade.

    Entidades de infra frequentemente não têm CPF/CNPJ/tipo_pessoa e são
    classificadas como PF, o que impediria a busca normal por representantes.
    """
    if not layer:
        return [], None
    for decl in declarantes:
        pid = decl.get("pessoa_id")
        if not pid:
            continue
        for tv in _TIPOS_REPRESENTACAO_INFRA:
            try:
                reps = _buscar_vinculados(layer, pid, tv)
                if reps:
                    return reps, tv
            except Exception:
                continue
    return [], None


# ── Extração de trechos do PDF ─────────────────────────────────────────────────

def _extrair_trechos_confrontante(pdf_segments, nome_confinante):
    if not pdf_segments or not nome_confinante:
        return []
    return [
        seg for seg in pdf_segments
        if seg.get("confrontacao")
        and _similaridade(seg["confrontacao"], nome_confinante) >= 0.75
    ]


# ── Bloco de coordenadas ───────────────────────────────────────────────────────

def _build_bloco_coordenadas(trechos, all_segments=None):
    """Formata os trechos do confrontante como bullets.

    Quando vante_latitude/vante_longitude estiverem vazios, busca as coordenadas
    do ponto final em all_segments pelo código do vante.
    """
    if not trechos:
        return "[TRECHOS DE CONFRONTAÇÃO NÃO ENCONTRADOS]"

    coords_by_vertex = {}
    if all_segments:
        for s in all_segments:
            codigo = s.get("codigo", "")
            lat = str(s.get("latitude", "") or "").strip()
            lon = str(s.get("longitude", "") or "").strip()
            if codigo and lat and lon and lat.upper() not in ("NONE", "NULL") and lon.upper() not in ("NONE", "NULL"):
                coords_by_vertex[codigo] = (lat, lon)

    def _nv(v):
        s = str(v or "").strip()
        return s if s and s.upper() not in ("NONE", "NULL") else ""

    linhas = []
    for seg in trechos:
        pini    = seg.get("codigo", "")
        lat_i   = _nv(seg.get("latitude", ""))
        lon_i   = _nv(seg.get("longitude", ""))
        pvante  = seg.get("vante", "")
        lat_v   = _nv(seg.get("vante_latitude", ""))
        lon_v   = _nv(seg.get("vante_longitude", ""))
        azimute = _nv(seg.get("azimute", ""))
        dist    = _nv(seg.get("dist_m", ""))

        if (not lat_v or not lon_v) and pvante and pvante in coords_by_vertex:
            lat_v, lon_v = coords_by_vertex[pvante]

        pt_ini = ("%s (Latitude: %s, Longitude: %s)" % (pini, lat_i, lon_i)
                  if lat_i and lon_i else pini)
        pt_fim = ("%s (Latitude: %s, Longitude: %s)" % (pvante, lat_v, lon_v)
                  if lat_v and lon_v else pvante)

        trecho = "• do ponto %s ao ponto %s" % (pt_ini, pt_fim)
        extras = []
        if azimute:
            extras.append("azimute %s" % azimute)
        if dist:
            extras.append("distância de %s m" % dist)
        if extras:
            trecho += "; %s;" % " – ".join(extras)
        linhas.append(trecho)

    return "\n".join(linhas)


# ── Qualificação do tipo de proprietário (placeholder simples) ─────────────────

def _qualificacao_confrontante(declarantes_confinante):
    """Retorna 'proprietário', 'proprietária' ou 'proprietário(a)'."""
    if not declarantes_confinante:
        return "proprietário(a)"
    decl0 = declarantes_confinante[0]
    if decl0.get("_tipo") == "PJ":
        return "proprietária"
    sexo = str(decl0.get("sexo") or "").strip().upper()
    if sexo == "MASCULINO":
        return "proprietário"
    if sexo == "FEMININO":
        return "proprietária"
    return "proprietário(a)"


# ── Segmentos para BLOCO_CONFRONTANTE quando confrontante é infra/recurso ──────

def _build_segs_infra(nome_infra, declarantes, layer, avisos_out):
    """Monta segmentos de qualificação para confrontante de infraestrutura/recurso."""
    nome_infra_up = (nome_infra or "[CONFRONTANTE]").upper()

    representantes, tipo_vinculo_infra = _buscar_representantes_infra(declarantes, layer)

    qualidade = "procurador" if str(tipo_vinculo_infra or "").upper() == "PROCURADOR" else "representante"
    artigo = _artigo_de(nome_infra)
    sufixo = _sufixo_qual_confronta(nome_infra)

    if not representantes:
        avisos_out.append(
            "Confrontante '%s' parece ser infraestrutura/recurso/área pública, "
            "mas não possui representante/procurador vinculado em pessoas_vinculos. "
            "Preencha manualmente o texto da declaração." % nome_infra
        )
        return [
            ("____________________", True),   # nome do representante (bold, para preenchimento)
            (", na qualidade de %s %s" % (qualidade, artigo), False),
            (nome_infra_up, True),
            (sufixo, False),
        ]

    todos_segs = []
    for i, rep in enumerate(representantes):
        tipo_rep = _detectar_tipo(rep)
        entry = dict(rep)
        entry["_tipo"] = tipo_rep
        entry["_conjuge"] = None
        entry["_representantes"] = []
        pid = rep.get("pessoa_id")
        if pid and layer:
            if tipo_rep == "PF":
                entry["_conjuge"] = _buscar_vinculado(layer, pid, "CONJUGE")
            else:
                entry["_representantes"] = _buscar_vinculados(layer, pid, "REPRESENTANTE_LEGAL")
        if tipo_rep == "PJ":
            segs, avs = _qualificar_pj_segs(entry, entry.get("_representantes", []))
        else:
            segs, avs = _qualificar_pf_segs(entry, entry.get("_conjuge"))
        avisos_out.extend(avs)
        if todos_segs:
            todos_segs.append(("; ", False))
        todos_segs.extend(segs)

    todos_segs.append((", na qualidade de %s %s" % (qualidade, artigo), False))
    todos_segs.append((nome_infra_up, True))
    todos_segs.append((sufixo, False))
    return todos_segs


# ── Assinaturas para confrontante de infra/recurso ────────────────────────────

def _montar_assinaturas_infra(nome_infra, representantes, tipo_vinculo):
    """Monta assinaturas com 'Procurador de: NOME_INFRA' ou 'Representante de: ...'."""
    nome_infra_up = (nome_infra or "[CONFRONTANTE]").upper()
    tv = str(tipo_vinculo or "REPRESENTANTE").upper()
    tv_label = "Procurador de" if tv == "PROCURADOR" else "Representante de"

    if not representantes:
        return [
            (
                "____________________",
                "%s: %s" % (tv_label, nome_infra_up),
            )
        ]

    assinaturas = []
    for rep in representantes:
        nome_rep = rep.get("nome", "").upper()
        cpf  = rep.get("cpf", "")
        cnpj = rep.get("cnpj", "")
        partes = ["%s: %s" % (tv_label, nome_infra_up)]
        if cpf:
            partes.append("CPF: %s" % cpf)
        elif cnpj:
            partes.append("CNPJ: %s" % cnpj)
        assinaturas.append((nome_rep, "\n".join(partes)))
    return assinaturas


# ── Substituições especializadas ──────────────────────────────────────────────

def _substituir_placeholder_negrito(document, placeholder_key, valor):
    """Substitui {{placeholder_key}} em todo o documento com o valor sempre em negrito.

    Processa cada parágrafo independentemente, preservando a formatação dos
    trechos antes e depois do placeholder.  Deve ser chamada antes de
    substituir_marcadores_docx para garantir que o valor seja bold
    independentemente do que o template define para aquele run.
    """
    if not valor:
        return
    key_norm = normalize_key(placeholder_key)
    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        match = next(
            (
                m for m in PLACEHOLDER_PATTERN.finditer(full_text)
                if normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == key_norm
            ),
            None,
        )
        if not match:
            continue
        segs_antes  = _segmentos_com_formato(runs, 0, match.start())
        segs_depois = _segmentos_com_formato(runs, match.end(), len(full_text))
        clear_paragraph_content(paragraph)
        for texto, rPr in segs_antes:
            _adicionar_run_formatado(paragraph, texto, rPr)
        r = paragraph.add_run(valor)
        r.bold = True
        for texto, rPr in segs_depois:
            _adicionar_run_formatado(paragraph, texto, rPr)


def _substituir_bloco_confrontante(
    document, declarantes, fallback_nome, avisos_out,
    is_infra=False, layer=None, proprietarios=None,
):
    """Substitui {{BLOCO_CONFRONTANTE}}.

    Para infra: qualifica o representante/procurador vinculado, acrescenta
    ', na qualidade de ... da/do NOME_INFRA, a qual confronta com o imóvel
    rural pertencente à [qualificação completa dos proprietários]...'.
    Os placeholders {{denominacao}}, {{nome_cartorio}}, etc. permanecem para
    serem resolvidos por substituir_marcadores_docx.
    """
    if is_infra:
        todos_segs = _build_segs_infra(fallback_nome, declarantes, layer, avisos_out)
        # Qualificação completa dos proprietários do imóvel principal (run a run)
        prop_segs = qualificar_proprietarios_segs(proprietarios or [], avisos_out)
        segs_continuacao = (
            [(" o imóvel rural pertencente à ", False)]
            + prop_segs
            + [
                (", proprietário do imóvel rural denominado \"", False),
                ("{{denominacao}}", True),
                ("\", registrado no ", False),
                ("{{nome_cartorio}}", True),
                (", {{municipio}}-{{uf}}, com matrícula sob nº.: ", False),
                ("{{matricula}}", True),
                (", cadastrado no INCRA sob o Código n.º ", False),
                ("{{codigo_incra}}", True),
                (".", False),
            ]
        )
    else:
        todos_segs = []
        for i, decl in enumerate(declarantes):
            if decl.get("_tipo") == "PJ":
                segs, avs = _qualificar_pj_segs(decl, decl.get("_representantes", []))
            else:
                segs, avs = _qualificar_pf_segs(
                    decl,
                    conjuge=decl.get("_conjuge"),
                    representantes=decl.get("_representantes"),
                    tipo_representacao=decl.get("_tipo_representacao"),
                )
            avisos_out.extend(avs)
            if i:
                todos_segs.append(("; ", False))
            todos_segs.extend(segs)

        if not todos_segs:
            nome_up = (fallback_nome or "[CONFRONTANTE NÃO INFORMADO]").upper()
            todos_segs = [
                (nome_up, True),
                (
                    ", _____________________, _____________________, "
                    "portador(a) do RG nº _________, "
                    "inscrito(a) no CPF/MF sob o nº _________________",
                    False,
                ),
            ]
            avisos_out.append(
                "Nenhuma pessoa vinculada ao confinante '%s'. "
                "{{BLOCO_CONFRONTANTE}} preenchido para completar manualmente." % fallback_nome
            )
        segs_continuacao = None  # preserva segs_depois do template

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        matches = [
            m for m in PLACEHOLDER_PATTERN.finditer(full_text)
            if normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_confrontante"
        ]
        if not matches:
            continue
        match = matches[0]
        segs_antes  = _segmentos_com_formato(runs, 0, match.start())
        segs_depois = _segmentos_com_formato(runs, match.end(), len(full_text))
        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = True
        for texto, rPr in segs_antes:
            _adicionar_run_formatado(paragraph, texto, rPr)
        for texto, bold in todos_segs:
            if texto:
                r = paragraph.add_run(texto)
                r.bold = bold
        if is_infra:
            for texto, bold in segs_continuacao:
                if texto:
                    r = paragraph.add_run(texto)
                    r.bold = bold
        else:
            for texto, rPr in segs_depois:
                _adicionar_run_formatado(paragraph, texto, rPr)


def _substituir_bloco_proprietarios_inline(document, proprietarios, avisos_out):
    """Substitui {{bloco_proprietarios}} com qualificação completa dos proprietários.

    Funciona tanto para placeholder inline (preserva texto antes/depois)
    quanto para placeholder standalone (parágrafo inteiro).
    """
    todos_segs = qualificar_proprietarios_segs(proprietarios, avisos_out)

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        matches = [
            m for m in PLACEHOLDER_PATTERN.finditer(full_text)
            if normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_proprietarios"
        ]
        if not matches:
            continue
        match = matches[0]
        segs_antes  = _segmentos_com_formato(runs, 0, match.start())
        segs_depois = _segmentos_com_formato(runs, match.end(), len(full_text))
        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = True
        for texto, rPr in segs_antes:
            _adicionar_run_formatado(paragraph, texto, rPr)
        for texto, bold in todos_segs:
            if texto:
                r = paragraph.add_run(texto)
                r.bold = bold
        for texto, rPr in segs_depois:
            _adicionar_run_formatado(paragraph, texto, rPr)


def _substituir_bloco_assinatura_confrontante(
    document, declarantes, fallback_nome, avisos_out,
    is_infra=False, layer=None
):
    """Substitui {{BLOCO_ASSINATURA_CONFRONTANTE}}."""
    if is_infra:
        representantes, tipo_vinculo = _buscar_representantes_infra(declarantes, layer)
        assinaturas = _montar_assinaturas_infra(fallback_nome, representantes, tipo_vinculo)
    else:
        if declarantes:
            assinaturas = _montar_assinaturas(declarantes)
        else:
            nome_up = (fallback_nome or "[NOME DO CONFINANTE]").upper()
            assinaturas = [(nome_up, "CPF: __________________")]

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        if not any(
            normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_assinatura_confrontante"
            for m in PLACEHOLDER_PATTERN.finditer(full_text)
        ):
            continue
        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = True
        for i, (linha1, linha2) in enumerate(assinaturas):
            if i:
                for _ in range(4):
                    paragraph.add_run().add_break()
            paragraph.add_run("________________________________________")
            paragraph.add_run().add_break()
            r = paragraph.add_run(linha1)
            r.bold = True
            if linha2:
                for parte in linha2.split("\n"):
                    paragraph.add_run().add_break()
                    paragraph.add_run(parte)


def _substituir_bloco_assinaturas_proprietarios(document, proprietarios, avisos_out):
    """Substitui {{BLOCO_ASSINATURAS_PROPRIETARIOS}} com assinaturas corretas.

    - PF sem representante: próprio e cônjuge assinam
    - PF com representante/procurador: representante assina com rótulo
    - PJ: representante legal assina
    Entre assinaturas: 4 quebras de linha.
    """
    assinaturas = _montar_assinaturas(proprietarios)
    if not assinaturas:
        return

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        if not any(
            normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_assinaturas_proprietarios"
            for m in PLACEHOLDER_PATTERN.finditer(full_text)
        ):
            continue
        clear_paragraph_content(paragraph)
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
        paragraph.paragraph_format.keep_together = True
        for i, (linha1, linha2) in enumerate(assinaturas):
            if i:
                for _ in range(4):
                    paragraph.add_run().add_break()
            paragraph.add_run("________________________________________")
            paragraph.add_run().add_break()
            r = paragraph.add_run(linha1)
            r.bold = True
            if linha2:
                for parte in linha2.split("\n"):
                    paragraph.add_run().add_break()
                    paragraph.add_run(parte)


def _substituir_bloco_coordenadas(document, trechos, avisos_out, all_segments=None):
    """Substitui {{BLOCO_COORDENADAS_CONFRONTANTE}} e insere quebra de página ao final."""
    linhas = _build_bloco_coordenadas(trechos, all_segments).split("\n")

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue
        full_text = "".join(r.text for r in runs)
        if not any(
            normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_coordenadas_confrontante"
            for m in PLACEHOLDER_PATTERN.finditer(full_text)
        ):
            continue

        rPr = None
        if runs[0]._r.rPr is not None:
            rPr = copy.deepcopy(runs[0]._r.rPr)

        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = False

        for i, linha in enumerate(linhas):
            if i:
                paragraph.add_run().add_break()
            r = paragraph.add_run(linha)
            if rPr is not None:
                r_elem = r._r
                existing = r_elem.rPr
                if existing is not None:
                    r_elem.remove(existing)
                r_elem.insert(0, copy.deepcopy(rPr))

        try:
            from docx.enum.text import WD_BREAK
            paragraph.add_run().add_break(WD_BREAK.PAGE)
        except Exception:
            pass


def _aplicar_quebra_segunda_pagina(document):
    """Garante quebra antes do bloco 'Declaramos ainda que o Profissional...'."""
    for paragraph in all_document_paragraphs(document):
        texto_lower = paragraph.text.lower()
        if "declaramos ainda" in texto_lower and "profissional" in texto_lower:
            paragraph.paragraph_format.page_break_before = True
            break


# ── Data atual por extenso ─────────────────────────────────────────────────────

_MESES_PT = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
    5: "maio", 6: "junho", 7: "julho", 8: "agosto",
    9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro",
}


def _data_atual_extenso():
    hoje = date.today()
    return "%d de %s de %d" % (hoje.day, _MESES_PT[hoje.month], hoje.year)


# ── Função principal ───────────────────────────────────────────────────────────

def fill_declaracao_confrontantes_template(
    template_path, output_path, layer, confinante_feature, pdf_segments, pdf_data
):
    """Gera a Declaração de Confrontantes para o confinante selecionado."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Biblioteca python-docx não encontrada. Instale python-docx no Python do QGIS."
        ) from exc

    template_path = os.path.abspath(template_path)
    output_path   = os.path.abspath(ensure_docx_path(output_path))
    if template_path == output_path:
        raise RuntimeError("Escolha um arquivo de saída diferente do modelo original.")

    avisos = []

    # ── Campos da feature do confinante ─────────────────────────────────
    fv = lambda key: _fval(layer, confinante_feature, key) if layer else ""

    confinante_id_raw = fv("id")
    codigo_raw        = fv("codigo")
    # confinante.nome NÃO é fonte oficial — dados vêm exclusivamente de
    # confinante_pessoas → pessoas (ver regras abaixo).

    if not codigo_raw:
        raise RuntimeError(
            "O campo 'codigo' do confinante está vazio. "
            "Confirme se o confinante selecionado está corretamente preenchido."
        )
    try:
        servico_id = int(float(codigo_raw))
    except (ValueError, TypeError):
        raise RuntimeError(
            "O campo 'codigo' do confinante não é um número válido: '%s'." % codigo_raw
        )

    # ── Dados do serviço principal ───────────────────────────────────────
    servico_data = {}
    if layer:
        try:
            servico_data = _buscar_dados_servico(layer, servico_id)
        except Exception:
            avisos.append(
                "Aviso: não foi possível consultar servicos.servicos. "
                "Os dados do imóvel principal virão apenas do PDF."
            )
    if not servico_data:
        avisos.append(
            "Aviso: nenhum serviço encontrado com id=%s. "
            "Os dados do imóvel principal virão apenas do PDF." % servico_id
        )

    # ── Proprietários do serviço principal (com cônjuge e representantes) ─
    proprietarios = []
    if layer:
        try:
            proprietarios = _buscar_proprietarios_servico(layer, servico_id)
            proprietarios = _enriquecer_proprietarios(proprietarios, layer)
        except Exception:
            pass
    if not proprietarios:
        avisos.append(
            "Aviso: nenhum proprietário encontrado para o serviço id=%s." % servico_id
        )

    # ── Pessoas do confinante ────────────────────────────────────────────
    confinante_pessoas = []
    if layer and confinante_id_raw:
        try:
            confinante_pessoas = _buscar_confinante_pessoas(layer, int(float(confinante_id_raw)))
        except Exception:
            pass

    # Bloqueia a geração se não há pessoa vinculada — confinante.nome não é fallback
    if not confinante_pessoas:
        raise RuntimeError(
            "O confinante (id=%s) não possui pessoa vinculada em confinante_pessoas. "
            "Cadastre ou vincule uma pessoa antes de gerar a declaração." % confinante_id_raw
        )

    declarantes_confinante = _enriquecer_confinante(confinante_pessoas, layer)

    # Nome representativo derivado das pessoas vinculadas (diagnóstico / mensagens)
    nome_confinante_display = " / ".join(
        p.get("nome", "").strip()
        for p in confinante_pessoas
        if p.get("nome", "").strip()
    ) or ("CONFINANTE_%s" % confinante_id_raw)

    # ── Nome do imóvel confrontante ──────────────────────────────────────
    # Fonte: confinante.nome_propriedade.  Se vazio, traço para preenchimento manual.
    # confinante.nome NÃO é usado como nome do imóvel.
    _nome_prop_raw = fv("nome_propriedade")
    nome_imovel_confrontante = (
        _nome_prop_raw
        if (_nome_prop_raw and _nome_prop_raw != "0")
        else "________________________________________"
    )

    # ── Detecção de infraestrutura/recurso/área pública ──────────────────
    # Usa o nome da pessoa vinculada (e campos da camada) — não confinante.nome
    is_infra = is_confrontante_infraestrutura_ou_recurso(
        nome_confinante_display,
        _nome_prop_raw or None,
        fv("tipo_imovel") if layer else None,
    )
    if is_infra:
        avisos.append(
            "Confrontante '%s' identificado como infraestrutura/recurso/área pública. "
            "Usando qualificação do representante/procurador vinculado." % nome_confinante_display
        )

    # ── Trechos do PDF ───────────────────────────────────────────────────
    # Busca pelos nomes das pessoas vinculadas e por nome_propriedade;
    # confinante.nome NÃO é usado na comparação com o PDF.
    _candidatos_pdf = [
        p.get("nome", "").strip()
        for p in confinante_pessoas
        if p.get("nome", "").strip()
    ]
    if _nome_prop_raw and _nome_prop_raw != "0":
        _candidatos_pdf.append(_nome_prop_raw)

    trechos = []
    _nome_busca_pdf = nome_confinante_display
    for _cand in _candidatos_pdf:
        _t = _extrair_trechos_confrontante(pdf_segments or [], _cand)
        if _t:
            trechos = _t
            _nome_busca_pdf = _cand
            break

    if not trechos:
        avisos.append(
            "Nenhum trecho de confrontação encontrado no PDF para '%s'. "
            "Verifique o nome da pessoa vinculada e o PDF selecionado." % _nome_busca_pdf
        )

    # ── Helper: dados do serviço com fallback para PDF ───────────────────
    def _sd(*keys):
        for k in keys:
            v = _limpar(servico_data.get(k) or pdf_data.get(k, ""))
            if v and v.upper() not in ("NULL", "NONE"):
                return v
        return ""

    denominacao   = _sd("denominacao")
    matricula     = _sd("matricula")
    nome_cartorio = _sd("nome_cartorio", "cartorio") or ""
    municipio     = _sd("municipio")
    uf            = _sd("uf")
    codigo_incra  = _sd("codigo_incra", "codigo_incra_sncr")

    trt = _limpar(pdf_data.get("trt", "")) or _limpar(servico_data.get("trt", ""))
    if trt:
        trt = re.sub(r"^\s*T\.?R\.?T\.?\s*:?\s*", "", trt, flags=re.IGNORECASE).strip()
        trt = formatar_trt_sem_duplicar_uf(trt, uf)

    # ── Qualificação do confrontante para placeholder simples ────────────
    if is_infra:
        _infra_reps, _infra_tv = _buscar_representantes_infra(declarantes_confinante, layer)
        qualificacao = "procurador" if str(_infra_tv or "").upper() == "PROCURADOR" else "representante"
    else:
        qualificacao = _qualificacao_confrontante(declarantes_confinante)

    # ── Dicionário de substituições simples ──────────────────────────────
    data = {
        "denominacao":               denominacao,
        "matricula":                 matricula,
        "nome_cartorio":             nome_cartorio,
        "municipio":                 municipio,
        "municipio_data":            formatar_municipio_para_data(municipio),
        "uf":                        (uf or "").upper(),
        "codigo_incra":              codigo_incra,
        "trt":                       trt,
        "t_r_t":                     trt,
        "nome_imovel_confrontante":  nome_imovel_confrontante,
        "data_atual":                _data_atual_extenso(),
        "qualificacao_confrontante": qualificacao,
        # __owners: lista plana para fallback do docx_utils (write_owner_block_paragraph)
        "__owners":                  _flatten_owners_com_conjuges(proprietarios),
    }

    # ── Gerar documento ──────────────────────────────────────────────────
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copyfile(template_path, output_path)
    document = Document(output_path)

    # 1. Substituições especializadas de blocos (devem ocorrer antes de substituir_marcadores_docx)
    _substituir_bloco_confrontante(
        document, declarantes_confinante, nome_confinante_display, avisos,
        is_infra=is_infra, layer=layer, proprietarios=proprietarios,
    )
    _substituir_bloco_proprietarios_inline(document, proprietarios, avisos)
    _substituir_bloco_assinatura_confrontante(
        document, declarantes_confinante, nome_confinante_display, avisos,
        is_infra=is_infra, layer=layer,
    )
    _substituir_bloco_assinaturas_proprietarios(document, proprietarios, avisos)
    _substituir_bloco_coordenadas(document, trechos, avisos, all_segments=pdf_segments)

    # 2. Placeholders que devem sair sempre em negrito (antes de substituir_marcadores_docx
    #    para sobrepor qualquer formatação que o template possa ter nos runs desses placeholders)
    for _key, _val in [
        ("denominacao",   denominacao),
        ("nome_cartorio", nome_cartorio),
        ("matricula",     matricula),
        ("codigo_incra",  codigo_incra),
    ]:
        _substituir_placeholder_negrito(document, _key, _val)

    # 3. Demais placeholders simples (data, município, TRT, etc.)
    substituir_marcadores_docx(document, data)

    # 4. Quebra de página antes do bloco de assinaturas do RT
    _aplicar_quebra_segunda_pagina(document)

    # 5. Mantém assinatura do RT junto na mesma página
    paragraphs = list(all_document_paragraphs(document))
    marcar_blocos_assinatura_tecnica(paragraphs)

    document.save(output_path)
    return avisos
