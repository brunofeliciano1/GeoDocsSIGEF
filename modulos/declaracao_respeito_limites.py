import os
import re
import shutil
import copy
import unicodedata

from ..docx_utils import (
    ensure_docx_path,
    substituir_marcadores_docx,
    normalize_key,
    all_document_paragraphs,
    PLACEHOLDER_PATTERN,
    clear_paragraph_content,
    marcar_blocos_assinatura_tecnica,
    paragraph_is_single_placeholder,
)
from .pessoas_utils import (
    enriquecer_proprietarios,
    qualificar_proprietarios_segs,
    montar_assinaturas_pessoas,
)


# ═══════════════════════════════════════════════════════════════
#  Classificação do tipo de imóvel do confrontante
# ═══════════════════════════════════════════════════════════════

# ── Bem da União: faixa de marinha, praias
_BEM_UNIAO_RE = re.compile(
    r'\b(?:AREA\s+DE\s+MARINHA|TERRENOS?\s+DE\s+MARINHA|TERRAS\s+DE\s+MARINHA'
    r'|PRAIA|MANGUE)\b',
)

# ── Ferrovia: linha férrea, trilhos
_FERROVIA_RE = re.compile(
    r'\b(?:LINHA\s+FERREA|FERROVIA|TRILHO)\b',
)

# ── Infraestrutura Pública: linhas de transmissão/distribuição, rede elétrica
_INFRA_PUBLICA_RE = re.compile(
    r'\b(?:LINHA\s+DE\s+TRANSMISSAO|LINHA\s+DE\s+DISTRIBUICAO'
    r'|REDE\s+ELETRICA)\b',
)

# ── Área Pública: servidões, faixas de domínio (verificar ANTES de Via Pública
# para que "SERVIDAO DE PASSAGEM" não seja capturado por "PASSAGEM")
_AREA_PUBLICA_RE = re.compile(
    r'AREA\s+DE\s+SERVIDAO\s+PUBLICA'
    r'|SERVIDAO\s+PUBLICA'
    r'|FAIXA\s+DE\s+DOMINIO'
    r'|\bAREA\s+PUBLICA\b'
    r'|SERVIDAO\s+DE\s+PASSAGEM',
)

# ── Imóvel Rural: propriedades privadas (verificar ANTES de Via Pública/Hídrico
# para evitar falso positivo em nomes como "Sítio Caminho Novo" ou "Fazenda Açude")
_RURAL_RE = re.compile(
    r'\b(?:SITIO|FAZENDA|IMOVEL|GLEBA|LOTE|PROPRIEDADE|TERRENO)\b'
    r'|\bDE\s+PROPRIEDADE\s+DE\b',
)

# ── Via Pública: estradas, rodovias, ruas, rodovias federais/estaduais
_VIA_PUBLICA_RE = re.compile(
    r'\b(?:ESTRADA|RODOVIA|CAMINHO|PASSAGEM|RUA|AVENIDA|TRAVESSA|VIELA)\b'
    r'|\bBR-\d|\bCE-\d',
)

# ── Recurso Hídrico: rios, riachos, açudes (só quando isolado, após checar rural)
_HIDRICO_RE = re.compile(
    r'\b(?:RIO|RIACHO|CORREGO|ACUDE|LAGO|LAGOA|BARRAGEM|CANAL)\b',
)

# ── Separador de sub-atributos internos do confinante enriquecido
# Ex.: "FAZENDA X, MATRÍCULA N° 123, CNS: Y" → split em "MATRICULA"
_SUB_ATTR_SPLIT_RE = re.compile(
    r',\s*(?:MATRICULA|CNS|CCIR|CIB|CAR)\b',
    re.IGNORECASE,
)


def classificar_tipo_imovel_confrontante(texto_confrontante):
    """Classifica o tipo de imóvel do confrontante.

    Categorias (em ordem de precedência):
      Bem da União > Ferrovia > Infraestrutura Pública > Área Pública
      > Imóvel Rural > Via Pública > Recurso Hídrico > Imóvel Rural (fallback)

    Imóvel Rural é verificado antes de Via Pública/Hídrico para evitar falsos
    positivos em nomes de propriedades que contenham "Caminho", "Açude", etc.
    """
    t = unicodedata.normalize("NFKD", str(texto_confrontante or "").upper())
    t = "".join(c for c in t if not unicodedata.combining(c))
    t = re.sub(r"\s+", " ", t).strip()
    if not t:
        return "Imóvel Rural"
    if _BEM_UNIAO_RE.search(t):
        return "Bem da União"
    if _FERROVIA_RE.search(t):
        return "Ferrovia"
    if _INFRA_PUBLICA_RE.search(t):
        return "Infraestrutura Pública"
    if _AREA_PUBLICA_RE.search(t):
        return "Área Pública"
    if _RURAL_RE.search(t):
        return "Imóvel Rural"
    if _VIA_PUBLICA_RE.search(t):
        return "Via Pública"
    if _HIDRICO_RE.search(t):
        return "Recurso Hídrico"
    return "Imóvel Rural"


# ═══════════════════════════════════════════════════════════════
#  Template
# ═══════════════════════════════════════════════════════════════

def find_declaracao_template(plugin_dir):
    path = os.path.join(plugin_dir, "models", "declaracao_respeito_limites.docx")
    return path if os.path.exists(path) else None


# ═══════════════════════════════════════════════════════════════
#  Geração principal
# ═══════════════════════════════════════════════════════════════

def fill_declaracao_template(template_path, output_path, data, layer=None):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Biblioteca python-docx não encontrada. Instale python-docx no Python do QGIS."
        ) from exc

    template_path = os.path.abspath(template_path)
    output_path   = os.path.abspath(ensure_docx_path(output_path))
    if template_path == output_path:
        raise RuntimeError("Escolha um arquivo de saída diferente do modelo original.")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copyfile(template_path, output_path)

    document  = Document(output_path)
    decl_data = copy.copy(data)
    owners    = decl_data.get("__owners", [])

    declarantes = enriquecer_proprietarios(owners, layer)

    avisos = []

    # {{BLOCO_DECLARANTES}} — qualificação jurídica completa com bold seletivo
    _substituir_bloco_declarantes(document, declarantes, avisos)

    # {{BLOCO_ASSINATURAS_DECLARANTES}}
    _substituir_bloco_assinaturas(document, declarantes)

    # Prepara confrontantes: converte para \n e calcula tipo_imovel_*
    _preparar_confrontantes_decl(decl_data)

    # Substitui confrontantes/tipo_imovel na tabela com quebras de linha reais
    _substituir_confrontantes_tabela_decl(document, decl_data)

    # Demais placeholders (MUNICIPIO, UF, TRT, etc.)
    substituir_marcadores_docx(document, decl_data)

    # Mantém bloco do RT junto na mesma página
    paragraphs = list(all_document_paragraphs(document))
    marcar_blocos_assinatura_tecnica(paragraphs)

    document.save(output_path)
    return avisos


# ═══════════════════════════════════════════════════════════════
#  Substituição {{BLOCO_DECLARANTES}} com bold parcial
# ═══════════════════════════════════════════════════════════════

def _substituir_bloco_declarantes(document, declarantes, avisos_out):
    """Substitui {{BLOCO_DECLARANTES}} preservando formatação original antes e depois.

    Estratégia:
    1. Coletar segmentos de texto + rPr clonado dos runs antes e depois do marcador.
    2. Limpar o parágrafo.
    3. Recriar runs antes (com rPr original) + bloco declarantes (negrito seletivo)
       + runs depois (com rPr original).
    """
    if not declarantes:
        avisos_out.append(
            "Aviso: Nenhum proprietário encontrado. "
            "{{BLOCO_DECLARANTES}} não foi substituído. "
            "Selecione a camada QGIS correta e tente novamente."
        )
        return

    todos_segs = qualificar_proprietarios_segs(declarantes, avisos_out)

    for paragraph in all_document_paragraphs(document):
        runs = paragraph.runs
        if not runs:
            continue

        full_text = "".join(run.text for run in runs)
        matches = [
            m for m in PLACEHOLDER_PATTERN.finditer(full_text)
            if normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_declarantes"
        ]
        if not matches:
            continue

        match = matches[0]

        segs_antes  = _segmentos_com_formato(runs, 0, match.start())
        segs_depois = _segmentos_com_formato(runs, match.end(), len(full_text))

        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = True

        for texto, rPr_clone in segs_antes:
            _adicionar_run_formatado(paragraph, texto, rPr_clone)

        for texto, bold in todos_segs:
            if texto:
                r = paragraph.add_run(texto)
                r.bold = bold

        for texto, rPr_clone in segs_depois:
            _adicionar_run_formatado(paragraph, texto, rPr_clone)


def _segmentos_com_formato(runs, text_start, text_end):
    """Retorna [(texto, rPr_clone)] para o intervalo [text_start, text_end) dos runs."""
    from copy import deepcopy
    resultado = []
    pos = 0
    for run in runs:
        run_start = pos
        run_end   = pos + len(run.text)
        pos       = run_end

        seg_start = max(run_start, text_start)
        seg_end   = min(run_end,   text_end)
        if seg_start >= seg_end:
            continue

        texto = run.text[seg_start - run_start : seg_end - run_start]
        if not texto:
            continue

        rPr = run._r.rPr
        resultado.append((texto, deepcopy(rPr) if rPr is not None else None))
    return resultado


def _adicionar_run_formatado(paragraph, texto, rPr_clone):
    """Adiciona um run com o rPr original clonado (preserva bold, fonte, tamanho, cor, etc.)."""
    if not texto:
        return
    run = paragraph.add_run(texto)
    if rPr_clone is None:
        return
    r_elem = run._r
    existing = r_elem.rPr
    if existing is not None:
        r_elem.remove(existing)
    r_elem.insert(0, rPr_clone)


# ═══════════════════════════════════════════════════════════════
#  Bloco de assinaturas
# ═══════════════════════════════════════════════════════════════

def _substituir_bloco_assinaturas(document, declarantes):
    assinaturas = montar_assinaturas_pessoas(declarantes)

    to_replace = [
        p for p in all_document_paragraphs(document)
        if "bloco_assinaturas_declarantes" in normalize_key(p.text)
    ]
    for paragraph in to_replace:
        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = True
        for i, (linha1, linha2) in enumerate(assinaturas):
            if i:
                paragraph.add_run().add_break()
                paragraph.add_run().add_break()
                paragraph.add_run().add_break()
                paragraph.add_run().add_break()
            paragraph.add_run("________________________________________")
            paragraph.add_run().add_break()
            r = paragraph.add_run(linha1)
            r.bold = True
            if linha2:
                for parte in linha2.split("\n"):
                    paragraph.add_run().add_break()
                    paragraph.add_run(parte)


# ═══════════════════════════════════════════════════════════════
#  Confrontantes: preparação e substituição na tabela
# ═══════════════════════════════════════════════════════════════

_CONFRONTANTE_DECL_KEYS = frozenset({
    "confrontante_norte", "confrontante_leste", "confrontante_sul", "confrontante_oeste",
    "tipo_imovel_norte", "tipo_imovel_leste", "tipo_imovel_sul", "tipo_imovel_oeste",
})


def _preparar_confrontantes_decl(decl_data):
    """Prepara confrontantes e tipo_imovel_* em decl_data.

    Confrontantes múltiplos por lado são separados por '\\n'.
    Sub-atributos internos (MATRÍCULA, CNS, CCIR, CIB) fazem parte do mesmo
    confrontante e não geram linhas extras na coluna Tipo de Imóvel.
    """
    for direction in ("norte", "leste", "sul", "oeste"):
        raw = str(decl_data.get("confrontante_%s" % direction) or "").strip()
        # Múltiplos confrontantes por lado → separados por \n
        # Sub-atributos (MATRÍCULA, CNS, etc.) ficam inline no mesmo item
        confrontantes = [c.strip() for c in raw.split("\n") if c.strip()] if raw else []
        decl_data["confrontante_%s" % direction] = "\n".join(confrontantes)
        # Uma linha de tipo por confrontante, baseada no nome-raiz (antes dos sub-atributos)
        tipos = []
        for c in confrontantes:
            nome_raiz = _SUB_ATTR_SPLIT_RE.split(c)[0].strip()
            tipos.append(classificar_tipo_imovel_confrontante(nome_raiz))
        decl_data["tipo_imovel_%s" % direction] = "\n".join(tipos)


def _substituir_confrontantes_tabela_decl(document, decl_data):
    """Percorre as células das tabelas e substitui placeholders de confrontante/tipo_imovel.

    Usa run.add_break() para múltiplos valores (separados por \\n), preservando a
    formatação original do run. Valores já substituídos são removidos do docx
    (o parágrafo não terá mais o marcador), então substituir_marcadores_docx
    não os processa novamente.
    """
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _substituir_paragraph_confrontante_decl(paragraph, decl_data)


def _substituir_paragraph_confrontante_decl(paragraph, decl_data):
    runs = paragraph.runs
    if not runs:
        return
    full_text = "".join(run.text for run in runs)
    if not full_text.strip():
        return
    matches = list(PLACEHOLDER_PATTERN.finditer(full_text))
    if not matches:
        return

    for match in matches:
        raw_key = match.group(2) or match.group(3) or match.group(4) or ""
        key = normalize_key(raw_key)
        if key not in _CONFRONTANTE_DECL_KEYS:
            continue
        if not paragraph_is_single_placeholder(full_text, key):
            break  # conteúdo misto: deixa substituir_marcadores_docx tratar

        value = str(decl_data.get(key, "") or "")
        lines = value.split("\n") if value else [""]

        rPr = None
        if runs[0]._r.rPr is not None:
            rPr = copy.deepcopy(runs[0]._r.rPr)

        clear_paragraph_content(paragraph)
        for i, line in enumerate(lines):
            if i > 0:
                paragraph.add_run().add_break()
            r = paragraph.add_run(line)
            if rPr is not None:
                r_elem = r._r
                existing = r_elem.rPr
                if existing is not None:
                    r_elem.remove(existing)
                r_elem.insert(0, copy.deepcopy(rPr))
        break
