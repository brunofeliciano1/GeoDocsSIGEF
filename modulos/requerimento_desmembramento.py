import os
import re
import shutil
import copy
from datetime import datetime

from docx.shared import Pt

from ..docx_utils import (
    ensure_docx_path,
    substituir_marcadores_docx,
    normalize_key,
    all_document_paragraphs,
    PLACEHOLDER_PATTERN,
    clear_paragraph_content,
    marcar_blocos_assinatura_tecnica,
)
from .pessoas_utils import (
    enriquecer_proprietarios,
    buscar_proprietarios_servico,
    buscar_dados_servico,
    limpar_valor,
)
from .requerimento_retificacao_area import (
    sanitizar_nome_denominacao,
    _fval,
    _formatar_nome_proprio,
    _formatar_endereco_cartorio,
    _ajustar_tratamento_oficial,
    _area_para_extenso,
    _substituir_bloco_declarantes,
    _substituir_bloco_assinaturas,
)

_ARIAL = "Arial"
_PT12  = Pt(12)


def find_requerimento_desmembramento_template(plugin_dir):
    path = os.path.join(plugin_dir, "models", "requerimento_desmembramento.docx")
    return path if os.path.exists(path) else None


# ── Leitura de campos das feições selecionadas ────────────────────────────────

def _ler_area_ha_feature(layer, feat):
    """Lê área em ha: prioridade area_ha_sigef (oficial SIGEF), fallback area_ha/area."""
    for nome in ("area_ha_sigef", "area_ha", "area", "area_certificada"):
        v = _fval(layer, feat, nome)
        if v:
            return v
    return ""


def _ler_perimetro_feature(layer, feat):
    """Lê perímetro em m: prioridade perimetro_sigef (oficial SIGEF), fallback perimetro."""
    for nome in ("perimetro_sigef", "perimetro", "perimetro_m"):
        v = _fval(layer, feat, nome)
        if v:
            return v
    return ""


# ── Formatação de área e perímetro ────────────────────────────────────────────

def _parse_float(val_str):
    val = str(val_str or "").strip()
    val = re.sub(r"\s*(ha|hectares?|m2|m²|metros?)\s*$", "", val, flags=re.IGNORECASE).strip()
    # Tira sufixo " m" de valores de perímetro (ex.: "63,33 m")
    val = re.sub(r"(?<=\d)\s+m\s*$", "", val, flags=re.IGNORECASE).strip()
    val = val.replace(",", ".")
    try:
        return float(val) if val else None
    except (ValueError, TypeError):
        return None


def _formatar_area_ha(area_ha_str):
    """Formata área em ha com 4 casas decimais e vírgula decimal.

    Exemplo: '53.8254' → '53,8254 ha'
    """
    val = _parse_float(area_ha_str)
    if val is None:
        return "[ÁREA NÃO INFORMADA]"
    return "{:.4f}".format(val).replace(".", ",") + " ha"


def _formatar_perimetro_str(perim_str):
    """Formata perímetro com 2 casas decimais, vírgula decimal e ponto de milhar.

    Exemplo: '6466.3' → '6.466,30 m'
    """
    val = _parse_float(perim_str)
    if val is None:
        return "[PERÍMETRO NÃO INFORMADO]"
    # "{:,.2f}" usa vírgula como separador de milhar e ponto decimal (padrão en-US)
    # Troca para padrão pt-BR: ponto de milhar, vírgula decimal
    s = "{:,.2f}".format(val)          # "6,466.30"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")  # "6.466,30"
    return s + " m"


# ── Construção dos itens de área ──────────────────────────────────────────────

def _construir_itens_areas(selected_features, layer):
    """Ordena features por desm_ordem e constrói lista de dicts por área."""
    def _ordem(feat):
        v = _fval(layer, feat, "desm_ordem")
        try:
            return float(str(v or "").replace(",", "."))
        except (ValueError, TypeError):
            return 9999.0

    feats_ord = sorted(selected_features, key=lambda f: (_ordem(f), f.id()))

    itens = []
    gleba_num = 0
    for feat in feats_ord:
        tipo = (_fval(layer, feat, "desm_tipo") or "").strip().upper()
        area_ha = _ler_area_ha_feature(layer, feat)
        perim   = _ler_perimetro_feature(layer, feat)

        if tipo == "REMANESCENTE":
            nome = "Área Remanescente"
        elif tipo == "DESMEMBRADA":
            gleba_num += 1
            nome = "Gleba %02d" % gleba_num
        else:
            nome = ("Área (%s)" % tipo) if tipo else "Área Desconhecida"

        itens.append({
            "tipo":      tipo,
            "nome":      nome,
            "area_str":  _formatar_area_ha(area_ha),
            "perim_str": _formatar_perimetro_str(perim),
        })
    return itens


def _montar_textos_bullets(itens_areas, matricula):
    """Monta lista de strings de bullet para cada área."""
    bullets = []
    for item in itens_areas:
        if item["tipo"] == "REMANESCENTE":
            desc = (
                "Área Remanescente (correspondente ao restante da matrícula "
                "nº %s deste CRI)" % matricula
            )
        else:
            desc = (
                "%s - Área Desmembrada "
                "(matrícula a ser aberta perante este CRI)" % item["nome"]
            )
        bullets.append(
            "• %s – com área de %s e perímetro de %s;" % (
                desc, item["area_str"], item["perim_str"]
            )
        )
    return bullets


def _montar_lista_areas(itens_areas):
    """'Área Remanescente, Gleba 01 e Gleba 02'."""
    nomes = [item["nome"] for item in itens_areas]
    if not nomes:
        return "[ÁREAS NÃO INFORMADAS]"
    if len(nomes) == 1:
        return nomes[0]
    return ", ".join(nomes[:-1]) + " e " + nomes[-1]


# ── Substituição do BLOCO_AREAS_DESMEMBRAMENTO ────────────────────────────────

def _substituir_bloco_areas_desmembramento(document, itens_areas, matricula):
    """Substitui {{BLOCO_AREAS_DESMEMBRAMENTO}} pela lista dinâmica de bullets.

    O primeiro bullet reutiliza o parágrafo placeholder; os demais são
    inseridos como novos parágrafos via XML, preservando a formatação.
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _Oxm

    def _criar_para_bullet(texto, ref_elem):
        """Cria elemento XML de parágrafo com texto em Arial 12, copiando pPr."""
        new_p = _Oxm("w:p")
        pPr = ref_elem.find(_qn("w:pPr"))
        if pPr is not None:
            new_p.append(copy.deepcopy(pPr))
        new_r = _Oxm("w:r")
        rPr = _Oxm("w:rPr")
        rFonts = _Oxm("w:rFonts")
        rFonts.set(_qn("w:ascii"), _ARIAL)
        rFonts.set(_qn("w:hAnsi"), _ARIAL)
        rPr.append(rFonts)
        sz = _Oxm("w:sz")
        sz.set(_qn("w:val"), str(int(_PT12.pt * 2)))
        rPr.append(sz)
        new_r.append(rPr)
        t = _Oxm("w:t")
        t.text = texto
        new_r.append(t)
        new_p.append(new_r)
        return new_p

    for paragraph in all_document_paragraphs(document):
        if not paragraph.runs:
            continue
        full_text = "".join(r.text for r in paragraph.runs)
        if not any(
            normalize_key(m.group(2) or m.group(3) or m.group(4) or "") == "bloco_areas_desmembramento"
            for m in PLACEHOLDER_PATTERN.finditer(full_text)
        ):
            continue

        bullets = _montar_textos_bullets(itens_areas, matricula)
        ref_elem = paragraph._element

        if not bullets:
            clear_paragraph_content(paragraph)
            r = paragraph.add_run("[ÁREAS DO DESMEMBRAMENTO NÃO INFORMADAS]")
            r.font.name = _ARIAL
            r.font.size = _PT12
            return

        # Primeiro bullet no próprio parágrafo placeholder
        clear_paragraph_content(paragraph)
        r = paragraph.add_run(bullets[0])
        r.font.name = _ARIAL
        r.font.size = _PT12

        # Demais bullets: novos parágrafos inseridos sequencialmente
        prev = ref_elem
        for bullet in bullets[1:]:
            new_p = _criar_para_bullet(bullet, ref_elem)
            prev.addnext(new_p)
            prev = new_p

        return


# ── Função principal ───────────────────────────────────────────────────────────

def fill_requerimento_desmembramento_template(
    template_path, output_path, data, layer=None, feature=None,
    selected_features=None
):
    """Gera o Requerimento de Desmembramento.

    Args:
        data:              dict com dados do PDF + QGIS (_collect_memorial_data).
        layer:             camada QGIS Serviços 2.
        feature:           feature do ID digitado no plugin (dados gerais do documento).
        selected_features: feições selecionadas na camada (REMANESCENTE + DESMEMBRADA),
                           usadas SOMENTE para montar os blocos de áreas.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Biblioteca python-docx não encontrada. Instale python-docx no Python do QGIS."
        ) from exc

    template_path = os.path.abspath(template_path)
    output_path   = os.path.abspath(ensure_docx_path(output_path))
    if template_path == output_path:
        raise RuntimeError("Escolha um arquivo de saída diferente do modelo original.")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copyfile(template_path, output_path)

    document  = Document(output_path)
    req_data  = copy.copy(data)
    avisos    = []

    servico_id = req_data.get("__servico_id")

    # ── Diagnóstico: serviço base e feições selecionadas ─────────────────────
    diag = ["[DESM] Serviço base (ID digitado): %s" % (servico_id or "(não identificado)")]
    if selected_features and layer:
        diag.append("[DESM] Feições selecionadas para o desmembramento:")
        for sf in selected_features:
            sf_id    = _fval(layer, sf, "id") or str(sf.id())
            sf_tipo  = (_fval(layer, sf, "desm_tipo") or "?").strip().upper()
            sf_ordem = _fval(layer, sf, "desm_ordem") or "?"
            diag.append("  ID %s | desm_tipo = %s | ordem = %s" % (sf_id, sf_tipo, sf_ordem))
    else:
        diag.append("[DESM] Nenhuma feição selecionada para o desmembramento.")
    avisos.extend(diag)

    # ── Dados do serviço via DB (ID digitado = serviço base) ─────────────────
    servico_data = {}
    if layer and servico_id:
        try:
            servico_data = buscar_dados_servico(layer, servico_id)
        except Exception:
            pass

    def _sd(key):
        v = limpar_valor(servico_data.get(key, ""))
        if v and v.upper() not in ("NULL", "NONE"):
            return v
        return _fval(layer, feature, key)

    # ── Proprietários (serviço base = ID digitado) ────────────────────────────
    owners = req_data.get("__owners", [])
    if not owners and layer and servico_id:
        try:
            owners = buscar_proprietarios_servico(layer, servico_id)
        except Exception:
            pass
    declarantes = enriquecer_proprietarios(owners, layer)

    # ── Campos gerais (camada Serviços 2 / DB) ────────────────────────────────
    folha_raw             = _sd("folha")
    livro_raw             = _sd("livro")
    ret_aver_raw          = _sd("ret_aver")
    oficial_cartorio_raw  = _sd("oficial_cartorio")
    endereco_cartorio_raw = _sd("endereco_cartorio")
    area_matricula_raw    = _sd("area_matricula")

    # Ambas as ocorrências de {{ESCREVER_EXTENSO}} referem-se à area_matricula
    area_matricula_extenso = _area_para_extenso(area_matricula_raw)

    # ── Dict de substituição ─────────────────────────────────────────────────
    req_data["folha"]             = folha_raw    or "-"
    req_data["livro"]             = livro_raw    or "-"
    req_data["ret_aver"]          = ret_aver_raw or "-"
    req_data["ano_atual"]         = str(datetime.now().year)
    req_data["oficial_cartorio"]  = (
        _formatar_nome_proprio(oficial_cartorio_raw)
        if oficial_cartorio_raw
        else "[OFICIAL DO CARTÓRIO NÃO INFORMADO]"
    )
    req_data["endereco_cartorio"] = (
        _formatar_endereco_cartorio(endereco_cartorio_raw)
        if endereco_cartorio_raw
        else "[ENDEREÇO DO CARTÓRIO NÃO INFORMADO]"
    )
    req_data["area_matricula"]   = (
        area_matricula_raw.replace(".", ",") if area_matricula_raw
        else "[ÁREA DA MATRÍCULA NÃO INFORMADA]"
    )
    req_data["escrever_extenso"] = area_matricula_extenso

    _nc_raw = _sd("nome_cartorio") or _sd("cartorio")
    if not _nc_raw:
        _nc_raw = req_data.get("nome_cartorio") or req_data.get("cartorio", "")
    if _nc_raw and _nc_raw not in ("[CARTÓRIO NÃO INFORMADO]",):
        _nc_fmt = _formatar_nome_proprio(_nc_raw)
        req_data["nome_cartorio"] = _nc_fmt
        req_data["cartorio"]      = _nc_fmt

    # ── Blocos de áreas ───────────────────────────────────────────────────────
    matricula = str(req_data.get("matricula", "")).strip() or "[MATRÍCULA]"

    if selected_features and layer:
        itens_areas = _construir_itens_areas(selected_features, layer)
        lista_areas = _montar_lista_areas(itens_areas)
    else:
        itens_areas = []
        lista_areas = "[ÁREAS DO DESMEMBRAMENTO NÃO INFORMADAS]"
        avisos.append(
            "Nenhuma feição selecionada; "
            "{{BLOCO_AREAS_DESMEMBRAMENTO}} e {{LISTA_AREAS_DESMEMBRAMENTO}} não substituídos."
        )

    req_data["lista_areas_desmembramento"] = lista_areas

    # ── Substituições especializadas (antes de substituir_marcadores_docx) ───
    _substituir_bloco_declarantes(document, declarantes, avisos)
    _substituir_bloco_assinaturas(document, declarantes)
    _ajustar_tratamento_oficial(document, oficial_cartorio_raw)
    _substituir_bloco_areas_desmembramento(document, itens_areas, matricula)

    # ── Demais placeholders simples ───────────────────────────────────────────
    substituir_marcadores_docx(document, req_data)

    # ── Mantém bloco do RT junto na mesma página ──────────────────────────────
    paragraphs = list(all_document_paragraphs(document))
    marcar_blocos_assinatura_tecnica(paragraphs)

    document.save(output_path)
    return avisos
