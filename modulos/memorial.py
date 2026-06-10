import os
import re
import shutil

from ..docx_utils import (
    PLACEHOLDER_PATTERN,
    normalize_key,
    public_data,
    ensure_docx_path,
    all_document_paragraphs,
    clear_paragraph_content,
    add_labeled_value,
    set_signature_run_font,
    signature_line_for_owner,
    add_blank_signature_lines,
    paragraph_is_single_placeholder,
    value_for_placeholder,
    replace_match_across_runs,
    replace_in_paragraphs,
    replace_in_tables,
    replace_placeholders_in_paragraph,
    write_owner_block_paragraph,
    write_owner_signature_block_paragraph,
    substituir_marcadores_docx,
    remover_paragrafos_placeholder_vazio,
    remover_paragrafos_rotulo_vazio,
    paragraph_has_image,
    is_responsible_signature_paragraph,
    is_short_signature_line,
    is_signature_separator,
    find_signature_block_start,
    find_signature_block_end,
    find_signature_block_end_after,
    mark_keep_block,
    marcar_blocos_assinatura_tecnica,
)
from .pessoas_utils import (
    enriquecer_proprietarios,
    montar_assinaturas_pessoas,
    normalizar_nome_destaque,
    limpar_valor,
)


def find_memorial_template(plugin_dir):
    path = os.path.join(plugin_dir, "models", "memorial.docx")
    return path if os.path.exists(path) else None


def default_memorial_output_path():
    downloads = os.path.join(os.path.expanduser("~"), "Downloads")
    downloads = downloads if os.path.isdir(downloads) else os.path.expanduser("~")
    return os.path.join(downloads, "MEMORIAL.docx")


def resolve_memorial_output_path(path):
    if not path:
        return default_memorial_output_path()
    if os.path.splitext(path)[1].lower() == ".docx":
        return path
    return os.path.join(path, "MEMORIAL.docx")


def fill_memorial_template(template_path, output_path, data, layer=None):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Biblioteca python-docx não encontrada. Instale python-docx no Python do QGIS."
        ) from exc

    template_path = os.path.abspath(template_path)
    output_path = os.path.abspath(ensure_docx_path(output_path))

    if template_path == output_path:
        raise RuntimeError("Escolha um arquivo de saída diferente do modelo original.")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copyfile(template_path, output_path)

    document = Document(output_path)

    # Enriquece owners com cônjuge e representantes via pessoas_utils
    owners = data.get("__owners", [])
    declarantes = enriquecer_proprietarios(owners, layer)

    # Substitui os blocos de proprietários ANTES de substituir_marcadores_docx
    # para que os placeholders desapareçam e não sejam processados novamente
    if declarantes:
        _substituir_bloco_proprietarios_memorial(document, declarantes)
        _substituir_bloco_assinaturas_memorial(document, declarantes)

    remover_paragrafos_placeholder_vazio(document, data)
    replaced_count = substituir_marcadores_docx(document, data)
    remover_paragrafos_rotulo_vazio(document)
    aplicar_regras_de_paginacao(document)
    document.save(output_path)
    return replaced_count


# ── Bloco de proprietários do memorial ───────────────────────

def _substituir_bloco_proprietarios_memorial(document, declarantes):
    """Substitui {{BLOCO_PROPRIETARIOS}} com formato resumido do memorial.

    Formato por entrada:
      Proprietário(a): NOME   ← rótulo bold, nome não-bold
      CPF: xxx               ← rótulo bold, valor não-bold

    Cônjuge aparece como entrada separada imediatamente abaixo.
    Entre proprietários: uma quebra de linha (sem linha em branco extra).
    """
    for paragraph in all_document_paragraphs(document):
        if not paragraph_is_single_placeholder(paragraph.text, "bloco_proprietarios"):
            continue

        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = True

        first_entry = True
        for decl in declarantes:
            if not first_entry:
                paragraph.add_run().add_break()
            first_entry = False

            nome = normalizar_nome_destaque(decl.get("nome", ""))
            _add_owner_entry_memorial(paragraph, nome, decl.get("cpf"), decl.get("cnpj"))

            conjuge = decl.get("_conjuge")
            if conjuge and limpar_valor(conjuge.get("nome", "")):
                paragraph.add_run().add_break()
                cnome = normalizar_nome_destaque(conjuge["nome"])
                _add_owner_entry_memorial(paragraph, cnome, conjuge.get("cpf"), conjuge.get("cnpj"))

        break  # único placeholder por documento


def _add_owner_entry_memorial(paragraph, nome, cpf=None, cnpj=None):
    """Adiciona 'Proprietário(a): NOME' + 'CPF/CNPJ: xxx' em dois runs por linha."""
    add_labeled_value(paragraph, "Proprietário(a):", nome or "____________________",
                      label_bold=True, value_bold=False, signature_font=True)
    paragraph.add_run().add_break()
    cpf_val  = limpar_valor(cpf  or "")
    cnpj_val = limpar_valor(cnpj or "")
    if cpf_val:
        add_labeled_value(paragraph, "CPF:", cpf_val,
                          label_bold=True, value_bold=False, signature_font=True)
    elif cnpj_val:
        add_labeled_value(paragraph, "CNPJ:", cnpj_val,
                          label_bold=True, value_bold=False, signature_font=True)


# ── Bloco de assinaturas do memorial ─────────────────────────

def _substituir_bloco_assinaturas_memorial(document, declarantes):
    """Substitui {{BLOCO_ASSINATURAS_PROPRIETARIOS}} usando montar_assinaturas_pessoas.

    Formato por assinante:
      ________________________________________
      Proprietário(a): NOME   ← rótulo não-bold, nome bold
      CPF: xxx               ← rótulo não-bold, valor bold
      [linhas adicionais de linha2 separadas por \\n]

    Representantes/procuradores seguem o padrão de montar_assinaturas_pessoas:
      ________________________________________
      Proprietário(a): REP_NOME
      Representante de: PROP_NOME
      CPF: xxx
    """
    assinaturas = montar_assinaturas_pessoas(declarantes)
    if not assinaturas:
        return

    for paragraph in all_document_paragraphs(document):
        if not paragraph_is_single_placeholder(paragraph.text, "bloco_assinaturas_proprietarios"):
            continue

        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

        clear_paragraph_content(paragraph)
        paragraph.paragraph_format.keep_together = True

        for i, (linha1, linha2) in enumerate(assinaturas):
            if i:
                add_blank_signature_lines(paragraph, 6)

            # Linha de assinatura proporcional ao nome
            ref_len = len("Proprietário(a): %s" % linha1)
            sig_len = max(45, min(ref_len, 70))
            sig_run = paragraph.add_run("_" * sig_len)
            set_signature_run_font(sig_run)
            paragraph.add_run().add_break()

            # "Proprietário(a): NOME" — rótulo não-bold, nome bold
            lr = paragraph.add_run("Proprietário(a): ")
            set_signature_run_font(lr, bold=False)
            nr = paragraph.add_run(linha1)
            set_signature_run_font(nr, bold=True)

            # Linhas de linha2 ("CPF: xxx", "Representante de: X", etc.)
            for parte in (linha2.split("\n") if linha2 else []):
                paragraph.add_run().add_break()
                if ": " in parte:
                    label_part, _, val_part = parte.partition(": ")
                    llr = paragraph.add_run(label_part + ": ")
                    set_signature_run_font(llr, bold=False)
                    vlr = paragraph.add_run(val_part)
                    set_signature_run_font(vlr, bold=True)
                else:
                    pr = paragraph.add_run(parte)
                    set_signature_run_font(pr, bold=False)

        break  # único placeholder por documento


def build_memorial_preview(data):
    parts = [
        "MEMORIAL DESCRITIVO",
        "CERTIFICAÇÃO: %s" % data.get("codigo_da_certificacao", data.get("certificacao", "")),
        "",
        "Denominação: %s" % data.get("denominacao", ""),
        data.get("bloco_proprietarios", ""),
        "Código INCRA/SNCR: %s" % data.get("codigo_incra", ""),
        "Município/UF: %s" % data.get("municipio_uf", ""),
        "Matrícula: %s" % data.get("matricula", ""),
        "Cartório: %s" % data.get("cartorio", ""),
        "CNS: %s" % data.get("cns", ""),
        "Comarca: %s" % data.get("comarca", ""),
        "TRT: %s" % data.get("trt", ""),
        "Área: %s" % data.get("area", data.get("area_ha", "")),
        "Perímetro: %s" % data.get("perimetro", data.get("perimetro_m", "")),
        "",
        "DESCRIÇÃO DO PERÍMETRO",
        data.get("descricao_perimetro", ""),
        "",
        "CONFRONTANTES DO IMÓVEL",
        "AO NORTE: %s" % data.get("confrontante_norte", ""),
        "AO LESTE: %s" % data.get("confrontante_leste", ""),
        "AO SUL: %s" % data.get("confrontante_sul", ""),
        "AO OESTE: %s" % data.get("confrontante_oeste", ""),
        "",
        "DECLARAÇÃO",
        "O proprietário/interessado e o profissional habilitado declaram estar cientes da previsão contida no art.213, § 14, da Lei n° 6.015/73.",
        "",
        "ASSINATURAS",
        "Resp. Téc.: BRUNO FELICIANO DE LIMA ALVES",
        data.get("bloco_assinaturas_proprietarios", ""),
    ]
    return "\n".join(part for part in parts if part is not None)


def aplicar_regras_de_paginacao(doc):
    paragraphs = list(all_document_paragraphs(doc))
    marcar_blocos_assinatura_tecnica(paragraphs)
    marcar_bloco_declaracao(paragraphs)

    for paragraph in paragraphs:
        text_key = normalize_key(paragraph.text)
        if "proprietario_interessado" in text_key or "art_213" in text_key or "lei_n_6_015_73" in text_key:
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = True


def marcar_bloco_declaracao(paragraphs):
    for index, paragraph in enumerate(paragraphs):
        text_key = normalize_key(paragraph.text)
        if text_key != "declaracao" and not text_key.startswith("declaracao"):
            continue

        paragraph.paragraph_format.page_break_before = True
        end_index = find_signature_block_end_after(paragraphs, index)
        mark_keep_block(paragraphs, index, end_index)
